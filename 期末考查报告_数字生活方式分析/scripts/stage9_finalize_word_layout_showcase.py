from __future__ import annotations

import json
import re
import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[2]
PROJECT = ROOT / "\u671f\u672b\u8003\u67e5\u62a5\u544a_\u6570\u5b57\u751f\u6d3b\u65b9\u5f0f\u5206\u6790"
USER_DOCX = Path(
    "D:/resourcessssss/"
    "\u666e\u901a\u8bfe/"
    "\u5927\u6570\u636e\u5206\u6790\u4e0e\u5e94\u7528/"
    "\u7ec8/"
    "\u79e6\u5929examination report.docx"
)
FINAL_DOCX = PROJECT / "final_submit" / "\u5927\u6570\u636e\u5206\u6790\u4e0e\u5e94\u7528\u671f\u672b\u8003\u67e5\u62a5\u544a.docx"
WEB_DEMO = PROJECT / "web_demo"
WEB_FIGURES = WEB_DEMO / "assets" / "figures"
SCREENSHOTS = WEB_DEMO / "assets" / "screenshots"
SUMMARY = PROJECT / "results" / "stage9_final_word_layout_showcase_summary.md"

REPOSITORY_URL = "https://github.com/hansu650/Big-Data-Homework"
HTMLPREVIEW_URL = (
    "https://htmlpreview.github.io/?https://github.com/hansu650/Big-Data-Homework/blob/main/"
    "%E6%9C%9F%E6%9C%AB%E8%80%83%E6%9F%A5%E6%8A%A5%E5%91%8A_%E6%95%B0%E5%AD%97%E7%94%9F%E6%B4%BB%E6%96%B9%E5%BC%8F%E5%88%86%E6%9E%90/"
    "web_demo/index.html"
)

FIGURE_FILES = [
    "eda_high_risk_flag_distribution.png",
    "eda_numeric_correlation_heatmap.png",
    "eda_boxplots_by_risk.png",
    "eda_category_risk_rate.png",
    "classification_final_confusion_matrix.png",
    "classification_precision_recall_curve.png",
    "classification_roc_curve.png",
    "regression_target_comparison.png",
    "regression_digital_dependence_observed_vs_predicted.png",
    "clustering_kmeans_elbow.png",
    "clustering_silhouette_by_k.png",
    "clustering_lifestyle_pca.png",
    "clustering_lifestyle_profile_heatmap.png",
]

TOC_ENTRIES = [
    ("Abstract", "2", 0),
    ("Table of Contents", "3", 0),
    ("Chapter 1 Dataset Selection", "4", 0),
    ("1.1 Course Requirements for Dataset Selection", "4", 1),
    ("1.2 Candidate Dataset Comparison and Reasons for Excluding pfm_train / pfm_test", "4", 1),
    ("1.3 Final Dataset Source and Compliance Check", "5", 1),
    ("1.4 Feature Fields, Sample Scale, and Supported Tasks", "5", 1),
    ("Chapter 2 Independent Research Theme and Analytical Perspective", "6", 0),
    ("Chapter 3 Data Preprocessing", "8", 0),
    ("Chapter 4 Exploratory Analysis and Data Visualization", "13", 0),
    ("Chapter 5 Machine Learning Modeling, Hyperparameter Tuning, and Evaluation", "20", 0),
    ("Chapter 6 Report Structure and Code Presentation", "31", 0),
    ("Chapter 7 Conclusions and Personal Reflection", "35", 0),
    ("References", "38", 0),
    ("Appendix A Core Code", "39", 0),
    ("Appendix B Full Experimental Running Instructions", "41", 0),
    ("Appendix C Result Files, Figure List, and Static Showcase", "42", 0),
    ("Appendix D Personal Contribution and Academic Integrity Statement", "44", 0),
]

BANNED_AI_TERMS = [
    "\u4eba\u5de5\u667a\u80fd\u6280\u672f\u4e0e\u5e94\u7528",
    "\u6e56\u5317\u5927\u5b66\u672c\u79d1\u8bfe\u7a0b\u8bbe\u8ba1",
    "CampusDepthSegLite",
    "RGB-D",
    "\u8bed\u4e49\u5206\u5272",
    "NYUDepthV2",
    "PyTorch",
    "\u738b\u96f7\u6625",
    "\u526f\u6559\u6388",
    "AI-homework",
    "mIoU",
    "Pixel Acc",
    "RGBD-concat",
    "\u6821\u56ed\u5ba4\u5185\u5de1\u68c0",
]
BANNED_LATEX_TERMS = [
    "\\chapter",
    "\\section",
    "\\subsection",
    "\\texttt",
    "\\input",
    "\\maybefigure",
    "\\begin",
    "\\end",
    "\\caption",
    "\\label",
    "\\ref",
    "p0.44",
    "p0.20",
    "p0.",
]


def ensure_dirs() -> None:
    FINAL_DOCX.parent.mkdir(parents=True, exist_ok=True)
    WEB_FIGURES.mkdir(parents=True, exist_ok=True)
    SCREENSHOTS.mkdir(parents=True, exist_ok=True)
    SUMMARY.parent.mkdir(parents=True, exist_ok=True)


def copy_figures() -> list[str]:
    copied = []
    for name in FIGURE_FILES:
        src = PROJECT / "figures" / name
        if src.exists():
            shutil.copy2(src, WEB_FIGURES / name)
            copied.append(name)
    return copied


def write_web_demo(copied_figures: list[str]) -> None:
    cards = "\n".join(
        f"""
        <article class="figure-card">
          <img src="assets/figures/{name}" alt="{name.replace('_', ' ').replace('.png', '')}">
          <div>
            <h3>{name.replace('_', ' ').replace('.png', '').title()}</h3>
            <p>{figure_caption(name)}</p>
          </div>
        </article>
        """.strip()
        for name in copied_figures
    )
    html = f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>Digital Lifestyle Analysis Dashboard</title>
  <style>
    :root {{
      --ink: #17212f;
      --muted: #5c6675;
      --line: #d9e0ea;
      --panel: #ffffff;
      --wash: #f6f8fb;
      --blue: #1f6fba;
      --green: #197b66;
      --amber: #b56916;
    }}
    * {{ box-sizing: border-box; }}
    body {{
      margin: 0;
      font-family: Inter, "Segoe UI", Arial, sans-serif;
      color: var(--ink);
      background: var(--wash);
      line-height: 1.55;
    }}
    header {{
      background: linear-gradient(110deg, #0f344f 0%, #1f6fba 58%, #2d8b78 100%);
      color: white;
      padding: 48px 56px 36px;
    }}
    header h1 {{ margin: 0 0 10px; font-size: 42px; letter-spacing: 0; }}
    header p {{ margin: 0; font-size: 18px; max-width: 900px; opacity: .95; }}
    nav {{
      display: flex;
      gap: 18px;
      flex-wrap: wrap;
      padding: 14px 56px;
      background: #fff;
      border-bottom: 1px solid var(--line);
      position: sticky;
      top: 0;
      z-index: 10;
    }}
    nav a {{ color: var(--blue); font-weight: 650; text-decoration: none; }}
    main {{ padding: 32px 56px 54px; }}
    section {{ max-width: 1180px; margin: 0 auto 34px; }}
    h2 {{ margin: 0 0 16px; font-size: 25px; }}
    .cards, .results {{ display: grid; gap: 16px; grid-template-columns: repeat(4, minmax(0, 1fr)); }}
    .card, .result, .workflow span, .figure-card, .structure, .note {{
      background: var(--panel);
      border: 1px solid var(--line);
      border-radius: 8px;
      box-shadow: 0 8px 22px rgba(20, 33, 48, .06);
    }}
    .card, .result {{ padding: 18px; }}
    .eyebrow {{ color: var(--muted); text-transform: uppercase; font-size: 12px; font-weight: 800; letter-spacing: .08em; }}
    .value {{ font-size: 28px; font-weight: 800; margin: 5px 0; }}
    .workflow {{ display: flex; flex-wrap: wrap; gap: 10px; align-items: center; }}
    .workflow span {{ padding: 10px 13px; font-weight: 700; }}
    .workflow i {{ color: var(--muted); font-style: normal; }}
    .gallery {{ display: grid; grid-template-columns: repeat(2, minmax(0, 1fr)); gap: 18px; }}
    .figure-card {{ overflow: hidden; }}
    .figure-card img {{ width: 100%; height: 330px; object-fit: contain; background: #f9fbfd; border-bottom: 1px solid var(--line); }}
    .figure-card div {{ padding: 14px 16px 16px; }}
    .figure-card h3 {{ margin: 0 0 6px; font-size: 17px; }}
    .figure-card p, .card p, .result p, .note p {{ margin: 0; color: var(--muted); }}
    .structure {{ padding: 20px; overflow-x: auto; }}
    pre {{ margin: 0; font-family: Consolas, "Courier New", monospace; font-size: 15px; }}
    .note {{ padding: 18px; }}
    .links a {{ color: var(--blue); overflow-wrap: anywhere; }}
    footer {{ padding: 24px 56px; color: var(--muted); background: #fff; border-top: 1px solid var(--line); }}
    @media (max-width: 900px) {{
      header, nav, main, footer {{ padding-left: 24px; padding-right: 24px; }}
      .cards, .results, .gallery {{ grid-template-columns: 1fr; }}
      header h1 {{ font-size: 32px; }}
    }}
  </style>
</head>
<body>
  <header id="top">
    <h1>Digital Lifestyle Analysis Dashboard</h1>
    <p>High-risk screening, digital dependence prediction, and lifestyle profiling with reproducible big-data analysis artifacts.</p>
  </header>
  <nav>
    <a href="#overview">Overview</a>
    <a href="#workflow">Workflow</a>
    <a href="#results">Key Results</a>
    <a href="#gallery">Gallery</a>
    <a href="#structure">Structure</a>
    <a href="#reproducibility">Reproducibility</a>
  </nav>
  <main>
    <section id="overview">
      <h2>Project Overview</h2>
      <div class="cards">
        <article class="card"><div class="eyebrow">Dataset</div><div class="value">2025</div><p>Digital Lifestyle Benchmark Dataset</p></article>
        <article class="card"><div class="eyebrow">Records</div><div class="value">3,500</div><p>Samples for course-level modeling and visualization</p></article>
        <article class="card"><div class="eyebrow">Fields</div><div class="value">24</div><p>Behavioral, lifestyle, mental-health, and outcome variables</p></article>
        <article class="card"><div class="eyebrow">Tasks</div><div class="value">3</div><p>Classification, regression, and clustering</p></article>
      </div>
    </section>
    <section id="workflow">
      <h2>Analysis Workflow</h2>
      <div class="workflow">
        <span>Dataset Compliance</span><i>to</i><span>Preprocessing</span><i>to</i><span>EDA</span><i>to</i><span>Classification</span><i>to</i><span>Regression</span><i>to</i><span>Clustering</span><i>to</i><span>Report</span>
      </div>
    </section>
    <section id="results">
      <h2>Key Results</h2>
      <div class="results">
        <article class="result"><div class="eyebrow">Classification</div><div class="value">F1 0.5355</div><p>Gradient Boosting, threshold=0.14, Recall=0.6420, PR-AUC=0.5084.</p></article>
        <article class="result"><div class="eyebrow">Digital Dependence</div><div class="value">R2 0.9839</div><p>MSE=3.1471 and MAE=0.9982 for digital_dependence_score prediction.</p></article>
        <article class="result"><div class="eyebrow">Productivity</div><div class="value">R2 -0.0041</div><p>A weak-prediction result retained as a negative modeling finding.</p></article>
        <article class="result"><div class="eyebrow">Clustering</div><div class="value">k=3</div><p>KMeans profile analysis with Silhouette=0.1860, interpreted exploratorily.</p></article>
      </div>
    </section>
    <section id="gallery">
      <h2>Visualization Gallery</h2>
      <div class="gallery">
        {cards}
      </div>
    </section>
    <section id="structure">
      <h2>Project Structure</h2>
      <div class="structure"><pre>Big-Data-Homework/
  final project directory/
    data/
    notebooks/
    src/
    scripts/
    figures/
    results/
    final_submit/
    overleaf_final/
    web_demo/
    requirements.txt</pre></div>
    </section>
    <section id="reproducibility">
      <h2>Reproducibility</h2>
      <div class="note">
        <p><strong>Run locally:</strong> install dependencies with <code>pip install -r requirements.txt</code>, run notebooks 00 to 05 in order, open the final DOCX, and open <code>web_demo/index.html</code> directly in a browser.</p>
        <p><strong>Boundaries:</strong> the dataset is synthetic benchmark data; the high-risk model is a screening tool, not a diagnosis tool; the regression result does not imply causality; clustering is exploratory; no online inference is provided.</p>
        <p class="links"><strong>GitHub:</strong> <a href="{REPOSITORY_URL}">{REPOSITORY_URL}</a></p>
      </div>
    </section>
  </main>
  <footer>Static showcase page for a course report. No server, checkpoint, GPU, or online inference is required.</footer>
  <script>
    if (location.hash) {{
      setTimeout(() => {{
        const target = document.querySelector(location.hash);
        if (target) target.scrollIntoView();
      }}, 250);
    }}
  </script>
</body>
</html>
"""
    (WEB_DEMO / "index.html").write_text(html, encoding="utf-8")
    readme = f"""# Static Visualization Showcase

This directory provides a lightweight static page for the Digital Lifestyle Analysis course report.

- Open `index.html` directly in a browser.
- No Python backend, GPU, server, model checkpoint, or online inference is required.
- The page summarizes the dataset, workflow, key model results, visualization gallery, project structure, reproducibility notes, and repository link.

GitHub repository: {REPOSITORY_URL}

Optional preview link:
{HTMLPREVIEW_URL}
"""
    (WEB_DEMO / "README_WEB_DEMO.md").write_text(readme, encoding="utf-8")


def figure_caption(name: str) -> str:
    captions = {
        "eda_high_risk_flag_distribution.png": "Class balance and screening-target distribution.",
        "eda_numeric_correlation_heatmap.png": "Correlation patterns among numeric behavior and outcome variables.",
        "eda_boxplots_by_risk.png": "Group-level differences between high-risk and non-high-risk samples.",
        "eda_category_risk_rate.png": "Risk-rate comparison across categorical groups.",
        "classification_final_confusion_matrix.png": "Final classification errors and correct predictions.",
        "classification_precision_recall_curve.png": "Precision-recall behavior for the screening model.",
        "classification_roc_curve.png": "ROC behavior for the final classifier.",
        "regression_target_comparison.png": "Comparison of regression targets and modeling difficulty.",
        "regression_digital_dependence_observed_vs_predicted.png": "Observed versus predicted digital dependence scores.",
        "clustering_kmeans_elbow.png": "K selection evidence for KMeans clustering.",
        "clustering_silhouette_by_k.png": "Silhouette score comparison across candidate cluster counts.",
        "clustering_lifestyle_pca.png": "PCA projection of lifestyle clusters.",
        "clustering_lifestyle_profile_heatmap.png": "Cluster-level profile interpretation across selected features.",
    }
    return captions.get(name, "Course-report visualization artifact.")


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "C:/Windows/Fonts/consolab.ttf" if bold else "C:/Windows/Fonts/consola.ttf",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def draw_tree_screenshot(path: Path, title: str, lines: list[tuple[str, str]]) -> None:
    width, height = 1400, 900
    img = Image.new("RGB", (width, height), "#f6f8fb")
    draw = ImageDraw.Draw(img)
    title_font = font(34, True)
    meta_font = font(17)
    code_font = font(25)
    draw.rounded_rectangle([38, 34, width - 38, height - 34], radius=16, fill="#ffffff", outline="#d7dee8", width=2)
    draw.text((70, 62), title, fill="#17212f", font=title_font)
    draw.text((70, 112), "Big-Data-Homework / final project directory", fill="#5c6675", font=meta_font)
    y = 170
    for text, color in lines:
        draw.text((86, y), text, fill=color, font=code_font)
        y += 38
    img.save(path)


def generate_tree_screenshots() -> None:
    project_lines = [
        ("final project directory/", "#17212f"),
        ("|-- data/", "#1f6fba"),
        ("|-- notebooks/", "#1f6fba"),
        ("|-- src/", "#1f6fba"),
        ("|-- scripts/", "#1f6fba"),
        ("|-- figures/", "#1f6fba"),
        ("|-- results/", "#1f6fba"),
        ("|-- final_submit/", "#1f6fba"),
        ("|-- overleaf_final/", "#1f6fba"),
        ("|-- web_demo/", "#1f6fba"),
        ("`-- requirements.txt", "#197b66"),
    ]
    notebook_lines = [
        ("notebooks/", "#17212f"),
        ("|-- 00_dataset_selection_and_compliance.ipynb", "#1f6fba"),
        ("|-- 01_data_preprocessing_and_eda.ipynb", "#1f6fba"),
        ("|-- 02_classification_high_risk.ipynb", "#1f6fba"),
        ("|-- 03_regression_productivity.ipynb", "#1f6fba"),
        ("|-- 04_clustering_lifestyle_profiles.ipynb", "#1f6fba"),
        ("`-- 05_result_summary_for_report.ipynb", "#1f6fba"),
        ("", "#17212f"),
        ("src/", "#17212f"),
        ("|-- config.py", "#197b66"),
        ("|-- data_utils.py", "#197b66"),
        ("|-- feature_engineering.py", "#197b66"),
        ("|-- model_utils.py", "#197b66"),
        ("`-- visualization.py", "#197b66"),
    ]
    draw_tree_screenshot(SCREENSHOTS / "project_structure_screenshot.png", "Project Directory and Engineering Structure", project_lines)
    draw_tree_screenshot(SCREENSHOTS / "notebooks_src_screenshot.png", "Notebook and Source-Code Organization", notebook_lines)


def capture_web_screenshots() -> dict[str, str]:
    edge = Path("C:/Program Files (x86)/Microsoft/Edge/Application/msedge.exe")
    homepage = SCREENSHOTS / "web_demo_homepage_screenshot.png"
    gallery = SCREENSHOTS / "web_demo_gallery_screenshot.png"
    failures: dict[str, str] = {}
    if not edge.exists():
        fallback_web_screenshot(homepage, "Static Visualization Showcase Homepage")
        fallback_web_screenshot(gallery, "Visualization Gallery in Static Showcase Page")
        failures["edge"] = "Microsoft Edge executable was not found; generated fallback screenshots with Pillow."
        return failures
    index = (WEB_DEMO / "index.html").resolve().as_uri()
    commands = [
        (homepage, index, "1440,1100"),
        (gallery, index + "#gallery", "1440,1200"),
    ]
    for output, url, size in commands:
        with tempfile.TemporaryDirectory() as tmp:
            cmd = [
                str(edge),
                "--headless=new",
                "--disable-gpu",
                "--hide-scrollbars",
                f"--user-data-dir={tmp}",
                "--virtual-time-budget=2500",
                f"--window-size={size}",
                f"--screenshot={output}",
                url,
            ]
            proc = subprocess.run(
                cmd,
                cwd=str(ROOT),
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="replace",
                timeout=30,
            )
            if proc.returncode != 0 or not output.exists() or output.stat().st_size < 1000:
                fallback_web_screenshot(output, output.stem.replace("_", " ").title())
                failures[output.name] = (proc.stderr or proc.stdout or "Edge screenshot failed").strip()
    make_gallery_composite(gallery)
    return failures


def make_gallery_composite(path: Path) -> None:
    width, height = 1400, 900
    img = Image.new("RGB", (width, height), "#f6f8fb")
    draw = ImageDraw.Draw(img)
    title_font = font(34, True)
    body_font = font(16)
    label_font = font(18, True)
    draw.text((58, 42), "Visualization Gallery in the Static Showcase Page", fill="#17212f", font=title_font)
    draw.text((58, 92), "Representative EDA, classification, regression, and clustering visuals", fill="#5c6675", font=body_font)
    names = [
        "eda_high_risk_flag_distribution.png",
        "classification_final_confusion_matrix.png",
        "regression_digital_dependence_observed_vs_predicted.png",
        "clustering_lifestyle_pca.png",
    ]
    labels = [
        "EDA target distribution",
        "Classification confusion matrix",
        "Regression observed vs predicted",
        "PCA lifestyle clusters",
    ]
    boxes = [(58, 140), (720, 140), (58, 500), (720, 500)]
    for (x, y), name, label in zip(boxes, names, labels):
        draw.rounded_rectangle([x, y, x + 602, y + 300], radius=12, fill="#ffffff", outline="#d7dee8", width=2)
        src = WEB_FIGURES / name
        if src.exists():
            thumb = Image.open(src).convert("RGB")
            thumb.thumbnail((560, 220), Image.Resampling.LANCZOS)
            tx = x + 20 + (560 - thumb.width) // 2
            ty = y + 18 + (220 - thumb.height) // 2
            img.paste(thumb, (tx, ty))
        draw.line([x, y + 238, x + 602, y + 238], fill="#d7dee8", width=1)
        draw.text((x + 22, y + 254), label, fill="#17212f", font=label_font)
        draw.text((x + 22, y + 280), figure_caption(name), fill="#5c6675", font=body_font)
    img.save(path)


def fallback_web_screenshot(path: Path, title: str) -> None:
    img = Image.new("RGB", (1400, 900), "#f6f8fb")
    draw = ImageDraw.Draw(img)
    title_font = font(38, True)
    body_font = font(23)
    draw.rounded_rectangle([42, 40, 1358, 860], radius=18, fill="#ffffff", outline="#d7dee8", width=2)
    draw.text((82, 88), title, fill="#17212f", font=title_font)
    y = 165
    lines = [
        "Digital Lifestyle Analysis Dashboard",
        "Classification + Regression + Clustering",
        "Key results: F1=0.5355, R2=0.9839, KMeans k=3",
        "Static page: web_demo/index.html",
        "No server, GPU, or online inference is required.",
    ]
    for line in lines:
        draw.text((92, y), line, fill="#5c6675", font=body_font)
        y += 52
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path)


def iter_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def clean_text(text: str) -> str:
    text = re.sub(r"p0(?:\.\d+)?(?:p0(?:\.\d+)?)*\s*", "", text)
    return text


def set_run_font(run, name: str = "Times New Roman", size: float | None = 10.5, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def clear_paragraph(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def append_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char)

    run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    run._r.append(instr)

    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)


def add_page_break_paragraph_before(anchor) -> None:
    p = anchor.insert_paragraph_before()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_paragraph_before(anchor, text: str = "", style: str | None = None, alignment=None):
    p = anchor.insert_paragraph_before(text)
    if style:
        p.style = style
    if alignment is not None:
        p.alignment = alignment
    return p


def apply_page_setup(doc: Document) -> None:
    for section in doc.sections:
        section.different_first_page_header_footer = False
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.35)
        section.bottom_margin = Cm(2.35)
        section.left_margin = Cm(2.35)
        section.right_margin = Cm(2.35)
        section.header_distance = Cm(1.25)
        section.footer_distance = Cm(1.25)
        for part in [
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ]:
            for paragraph in part.paragraphs:
                clear_paragraph(paragraph)
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("Page ")
        set_run_font(r, size=9.5)
        append_field(p, "PAGE")
        r = p.add_run(" of ")
        set_run_font(r, size=9.5)
        append_field(p, "NUMPAGES")


def remove_elements_before(paragraph) -> None:
    body = paragraph._parent._element
    for child in list(body):
        if child is paragraph._p:
            break
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def remove_paragraph_range(paragraphs: list, start: int, end_exclusive: int) -> None:
    for p in paragraphs[start:end_exclusive]:
        parent = p._element.getparent()
        if parent is not None:
            parent.remove(p._element)


def insert_cover(doc: Document) -> None:
    abstract = find_paragraph(doc, "Abstract")
    if not abstract:
        raise RuntimeError("Could not find Abstract in source DOCX.")
    remove_elements_before(abstract)
    cover_lines = [
        ("\u6e56\u5317\u5927\u5b66 2025-2026 \u5b66\u5e74\u5ea6\u7b2c 2 \u5b66\u671f\u8bfe\u7a0b\u8003\u67e5\u8bd5\u9898\u7eb8", 18, True, WD_ALIGN_PARAGRAPH.CENTER),
        ("The paper of course exam", 15, True, WD_ALIGN_PARAGRAPH.CENTER),
        ("", 12, False, WD_ALIGN_PARAGRAPH.CENTER),
        ("Name: Big Data Analysis and Applications", 13, False, WD_ALIGN_PARAGRAPH.CENTER),
        ("Content: Course Report", 13, False, WD_ALIGN_PARAGRAPH.CENTER),
        ("Teacher: Li Jie", 13, False, WD_ALIGN_PARAGRAPH.CENTER),
        ("", 12, False, WD_ALIGN_PARAGRAPH.CENTER),
        ("Digital Lifestyle Analysis Course Report", 20, True, WD_ALIGN_PARAGRAPH.CENTER),
        ("High-Risk Screening, Digital Dependence Prediction, and Lifestyle Profiling", 13, False, WD_ALIGN_PARAGRAPH.CENTER),
        ("", 12, False, WD_ALIGN_PARAGRAPH.CENTER),
        ("Student ID: __________________    Name: __________________", 12, False, WD_ALIGN_PARAGRAPH.LEFT),
        ("College: __________________       Major and Grade: __________________", 12, False, WD_ALIGN_PARAGRAPH.LEFT),
        ("Institution: __________________   Grade and major: __________________", 12, False, WD_ALIGN_PARAGRAPH.LEFT),
        ("", 12, False, WD_ALIGN_PARAGRAPH.CENTER),
        ("Teacher's comments:", 12, True, WD_ALIGN_PARAGRAPH.LEFT),
        ("______________________________________________________________________________", 12, False, WD_ALIGN_PARAGRAPH.LEFT),
        ("______________________________________________________________________________", 12, False, WD_ALIGN_PARAGRAPH.LEFT),
        ("", 12, False, WD_ALIGN_PARAGRAPH.CENTER),
        ("Total score: ________________       Grading teacher: ________________", 12, False, WD_ALIGN_PARAGRAPH.LEFT),
    ]
    for text, size, bold, alignment in cover_lines:
        p = abstract.insert_paragraph_before(text)
        p.alignment = alignment
        p.paragraph_format.space_after = Pt(7 if text else 10)
        for run in p.runs:
            set_run_font(run, "Times New Roman", size, bold)
            if "\u6e56\u5317\u5927\u5b66" in text:
                run.font.name = "SimSun"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    add_page_break_paragraph_before(abstract)


def find_paragraph(doc: Document, exact_text: str):
    for p in doc.paragraphs:
        if p.text.strip() == exact_text:
            return p
    return None


def replace_toc(doc: Document) -> None:
    paragraphs = doc.paragraphs
    toc_idx = next((i for i, p in enumerate(paragraphs) if p.text.strip() == "Table of Contents"), None)
    if toc_idx is None:
        return
    chapter_idxs = [i for i, p in enumerate(paragraphs) if p.text.strip() == "Chapter 1 Dataset Selection" and i > toc_idx]
    if not chapter_idxs:
        return
    main_chapter_idx = chapter_idxs[-1]
    anchor = paragraphs[main_chapter_idx]
    remove_paragraph_range(paragraphs, toc_idx, main_chapter_idx)

    add_page_break_paragraph_before(anchor)
    title = add_paragraph_before(anchor, "Table of Contents", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    title.paragraph_format.space_after = Pt(12)
    for run in title.runs:
        set_run_font(run, size=16, bold=True)

    for text, page, level in TOC_ENTRIES:
        p = add_paragraph_before(anchor)
        p.paragraph_format.left_indent = Inches(0.28 * level)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.25), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        r = p.add_run(text)
        set_run_font(r, size=10.5, bold=(level == 0))
        r = p.add_run("\t" + page)
        set_run_font(r, size=10.5, bold=(level == 0))
    add_page_break_paragraph_before(anchor)


def normalize_styles(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Normal"].font.size = Pt(10.5)
    for p in iter_paragraphs(doc):
        original_text = p.text
        for run in p.runs:
            if not run_has_visual(run):
                run.text = clean_text(run.text)
            if run.text:
                set_run_font(run, size=10.5)
        text = clean_text(original_text).strip()
        if is_static_toc_entry(text):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                set_run_font(run, size=10.5, bold=text.split("\t", 1)[0] in {entry[0] for entry in TOC_ENTRIES if entry[2] == 0})
            continue
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.08
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.space_before = Pt(0)
        if not text:
            continue
        if text.startswith("Chapter "):
            p.style = "Heading 1"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(10)
            for run in p.runs:
                set_run_font(run, size=15, bold=True)
                run.font.color.rgb = RGBColor(31, 80, 132)
        elif re.match(r"^\d+\.\d+\s+", text):
            p.style = "Heading 2"
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(5)
            for run in p.runs:
                set_run_font(run, size=12, bold=True)
                run.font.color.rgb = RGBColor(31, 80, 132)
        elif text.startswith("Appendix ") or text in {"References", "Abstract", "Table of Contents"}:
            p.style = "Heading 1"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if text in {"References", "Abstract", "Table of Contents"} else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                set_run_font(run, size=14, bold=True)
                run.font.color.rgb = RGBColor(31, 80, 132)
        elif text.startswith("Figure ") or text.startswith("Table ") or text.startswith("Code "):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if text.startswith(("Figure ", "Table ")) else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                set_run_font(run, size=10, bold=True)
        elif text.startswith(("Student ID:", "College:", "Institution:", "Teacher's comments:", "Total score:")):
            for run in p.runs:
                set_run_font(run, size=12)

    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.space_after = Pt(2)
                    for run in p.runs:
                        run.text = clean_text(run.text)
                        set_run_font(run, size=9.2)
                set_cell_shading(cell, None)
        if table.rows:
            for cell in table.rows[0].cells:
                set_cell_shading(cell, "EAF1F8")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True


def run_has_visual(run) -> bool:
    return (
        bool(run._r.xpath(".//w:drawing"))
        or bool(run._r.xpath(".//w:pict"))
        or bool(run._r.xpath(".//w:br"))
    )


def is_static_toc_entry(text: str) -> bool:
    if "\t" not in text:
        return False
    label = text.split("\t", 1)[0].strip()
    return any(label == entry[0] for entry in TOC_ENTRIES)


def set_cell_shading(cell, fill: str | None) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for shd in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(shd)
    if fill:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)


def insert_table_before(anchor, data: list[tuple[str, str]]):
    container = anchor._parent
    table = container.add_table(rows=len(data), cols=2, width=Inches(6.2))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for idx, (key, value) in enumerate(data):
        table.cell(idx, 0).text = key
        table.cell(idx, 1).text = value
        for cell in table.rows[idx].cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    set_run_font(run, size=9.2, bold=(idx == 0))
        if idx == 0:
            for cell in table.rows[idx].cells:
                set_cell_shading(cell, "EAF1F8")
    tbl = table._tbl
    tbl.getparent().remove(tbl)
    anchor._p.addprevious(tbl)
    return table


def insert_image_block_before(anchor, heading: str, image_path: Path, caption: str, explanation: list[str]) -> None:
    h = add_paragraph_before(anchor, heading)
    for run in h.runs:
        set_run_font(run, size=11, bold=True)
    p = add_paragraph_before(anchor)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(5.9))
    cap = add_paragraph_before(anchor, caption, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    for run in cap.runs:
        set_run_font(run, size=9.5, bold=True)
    for text in explanation:
        e = add_paragraph_before(anchor, text)
        for run in e.runs:
            set_run_font(run, size=10.2)


def insert_stage9_content(doc: Document) -> None:
    chapter7 = find_first_startswith(doc, "Chapter 7 Conclusions")
    if chapter7:
        h = add_paragraph_before(chapter7, "6.6 GitHub Repository, Project Structure, and Static Showcase")
        h.style = "Heading 2"
        for run in h.runs:
            set_run_font(run, size=12, bold=True)
            run.font.color.rgb = RGBColor(31, 80, 132)
        para = add_paragraph_before(
            chapter7,
            "The complete experimental code, notebooks, generated figures, result CSV files, and final report materials are organized in the GitHub repository below. The static showcase page is a lightweight browser-based presentation of the project workflow and visual results.",
        )
        for run in para.runs:
            set_run_font(run, size=10.5)
        insert_table_before(
            chapter7,
            [
                ("Artifact", "Location"),
                ("GitHub repository", REPOSITORY_URL),
                ("Project directory", "Big-Data-Homework / final project directory"),
                ("Final report", "final_submit / final Word course report"),
                ("Overleaf/LaTeX package", "overleaf_final"),
                ("Static visualization showcase", "web_demo / index.html"),
                ("Optional HTML preview", HTMLPREVIEW_URL),
            ],
        )
        add_paragraph_before(chapter7, "")
        insert_image_block_before(
            chapter7,
            "Project structure screenshot",
            SCREENSHOTS / "project_structure_screenshot.png",
            "Figure 6-1 Project directory and engineering structure",
            [
                "The project is organized around data, notebooks, source code, scripts, figures, results, and final submission artifacts.",
                "This structure demonstrates that the report is supported by reproducible engineering materials rather than only by a standalone document.",
            ],
        )
        insert_image_block_before(
            chapter7,
            "Notebook and source-code screenshot",
            SCREENSHOTS / "notebooks_src_screenshot.png",
            "Figure 6-2 Notebook and source-code organization",
            [
                "The notebook sequence follows the course-report workflow from dataset compliance checking to final result summary.",
                "The src directory separates configuration, data utilities, feature engineering, model utilities, and visualization helpers.",
            ],
        )

    appendix_c = find_first_startswith(doc, "Appendix C")
    appendix_d = find_first_startswith(doc, "Appendix D")
    if appendix_c and appendix_d:
        h = add_paragraph_before(appendix_d, "Appendix C Static Visualization Showcase")
        h.style = "Heading 2"
        for run in h.runs:
            set_run_font(run, size=12, bold=True)
            run.font.color.rgb = RGBColor(31, 80, 132)
        for text in [
            f"GitHub repository: {REPOSITORY_URL}",
            "Static visualization showcase: web_demo/index.html",
            f"Optional HTML preview: {HTMLPREVIEW_URL}",
            "The static page can be opened locally in a browser and does not require a server, GPU, checkpoint, or online inference.",
        ]:
            p = add_paragraph_before(appendix_d, text)
            for run in p.runs:
                set_run_font(run, size=10.3)
        insert_table_before(
            appendix_d,
            [
                ("Showcase file", "Purpose"),
                ("web_demo/index.html", "Static overview dashboard and visualization gallery"),
                ("web_demo/README_WEB_DEMO.md", "Local opening instructions and boundaries"),
                ("web_demo/assets/figures/", "Copied report figures used by the gallery"),
                ("web_demo/assets/screenshots/", "Screenshots inserted into the Word report"),
            ],
        )
        add_paragraph_before(appendix_d, "")
        insert_image_block_before(
            appendix_d,
            "Static showcase homepage",
            SCREENSHOTS / "web_demo_homepage_screenshot.png",
            "Figure C-1 Static visualization showcase homepage",
            [
                "The homepage summarizes the project theme, workflow, dataset scale, task types, and key modeling results.",
                "It provides a concise visual entry point for readers who want to inspect the project before reading the full report.",
            ],
        )
        insert_image_block_before(
            appendix_d,
            "Static showcase gallery",
            SCREENSHOTS / "web_demo_gallery_screenshot.png",
            "Figure C-2 Visualization gallery in the static showcase page",
            [
                "The gallery collects representative EDA, classification, regression, and clustering figures from the project artifacts.",
                "It is intended for static display only and does not perform online prediction or model inference.",
            ],
        )

    appendix_b = find_first_startswith(doc, "Appendix B")
    if appendix_b and appendix_c:
        h = add_paragraph_before(appendix_c, "Appendix B Repository Access and Local Opening Notes")
        h.style = "Heading 2"
        for run in h.runs:
            set_run_font(run, size=12, bold=True)
            run.font.color.rgb = RGBColor(31, 80, 132)
        for text in [
            f"The complete repository is available at {REPOSITORY_URL}.",
            "The final project directory contains notebooks, source code, scripts, figures, results, final submission files, the LaTeX package, and the static showcase page.",
            "To inspect the static showcase, open web_demo/index.html directly in a browser.",
        ]:
            p = add_paragraph_before(appendix_c, text)
            for run in p.runs:
                set_run_font(run, size=10.3)
        insert_table_before(
            appendix_c,
            [
                ("Step", "Local command or action"),
                ("Install environment", "pip install -r requirements.txt"),
                ("Run notebooks", "Run notebooks 00 to 05 in order"),
                ("Open final report", "Open final_submit / final Word course report"),
                ("Open showcase", "Open web_demo / index.html in a browser"),
            ],
        )
        add_paragraph_before(appendix_c, "")


def find_first_startswith(doc: Document, prefix: str):
    for p in doc.paragraphs:
        text = p.text.strip()
        if is_static_toc_entry(text):
            continue
        if text.startswith(prefix):
            return p
    return None


def save_docx_from_user() -> dict[str, int | str]:
    if not USER_DOCX.exists():
        raise FileNotFoundError(f"User DOCX not found: {USER_DOCX}")
    shutil.copy2(USER_DOCX, FINAL_DOCX)
    doc = Document(FINAL_DOCX)
    insert_cover(doc)
    replace_toc(doc)
    apply_page_setup(doc)
    normalize_styles(doc)
    insert_stage9_content(doc)
    normalize_styles(doc)
    trim_trailing_empty_paragraphs(doc)
    doc.save(FINAL_DOCX)
    doc = Document(FINAL_DOCX)
    return {
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "inline_shapes": len(doc.inline_shapes),
        "code_blocks": sum(1 for p in doc.paragraphs if p.text.strip().startswith("Code ")),
    }


def trim_trailing_empty_paragraphs(doc: Document) -> None:
    for p in reversed(doc.paragraphs):
        if p.text.strip():
            break
        if p._p.xpath(".//w:drawing") or p._p.xpath(".//w:pict") or p._p.xpath(".//w:br"):
            break
        parent = p._p.getparent()
        if parent is not None:
            parent.remove(p._p)


def extract_docx_text(path: Path) -> str:
    doc = Document(path)
    chunks = []
    for p in doc.paragraphs:
        chunks.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    return "\n".join(chunks)


def audit_docx() -> dict[str, object]:
    text = extract_docx_text(FINAL_DOCX)
    doc = Document(FINAL_DOCX)
    chinese_paragraph_hits = [
        p.text.strip()
        for p in doc.paragraphs
        if re.search(r"[\u4e00-\u9fff]", p.text)
    ]
    ai_hits = {term: text.count(term) for term in BANNED_AI_TERMS if term in text}
    latex_hits = {term: text.count(term) for term in BANNED_LATEX_TERMS if term in text}
    p0_hits = len(re.findall(r"p0(?:\.\d+)?", text))
    missing_figure_hits = text.lower().count("missing figure")
    todo_hits = len(re.findall(r"\bTODO\b", text, flags=re.I))
    github_present = REPOSITORY_URL in text
    web_demo_present = (WEB_DEMO / "index.html").exists()
    readme_present = (WEB_DEMO / "README_WEB_DEMO.md").exists()
    screenshots = sorted(p.name for p in SCREENSHOTS.glob("*.png"))
    return {
        "ai_hits": ai_hits,
        "latex_hits": latex_hits,
        "p0_hits": p0_hits,
        "missing_figure_hits": missing_figure_hits,
        "todo_hits": todo_hits,
        "chinese_paragraph_hits": chinese_paragraph_hits,
        "github_present": github_present,
        "web_demo_present": web_demo_present,
        "readme_present": readme_present,
        "screenshots": screenshots,
    }


def try_word_pdf_export() -> dict[str, str]:
    pdf_path = FINAL_DOCX.with_suffix(".pdf")
    ps = f"""
$ErrorActionPreference = 'Stop'
$word = $null
try {{
  $word = New-Object -ComObject Word.Application
  $word.Visible = $false
  $doc = $word.Documents.Open('{str(FINAL_DOCX).replace("'", "''")}')
  $doc.ExportAsFixedFormat('{str(pdf_path).replace("'", "''")}', 17)
  $doc.Close($false)
  $word.Quit()
  Write-Output 'PDF_OK'
}} catch {{
  if ($word -ne $null) {{ $word.Quit() }}
  Write-Output ('PDF_FAIL: ' + $_.Exception.Message)
  exit 2
}}
"""
    proc = subprocess.run(["powershell", "-NoProfile", "-Command", ps], cwd=str(ROOT), capture_output=True, text=True, timeout=120)
    if proc.returncode == 0 and pdf_path.exists() and pdf_path.stat().st_size > 1000:
        return {"status": "generated", "path": str(pdf_path), "message": proc.stdout.strip()}
    return {"status": "not_generated", "path": str(pdf_path), "message": (proc.stdout + proc.stderr).strip()}


def write_summary(
    copied_figures: list[str],
    screenshot_failures: dict[str, str],
    doc_stats: dict[str, int | str],
    audit: dict[str, object],
    pdf_result: dict[str, str],
) -> None:
    lines = [
        "# Stage 9 Final Word Layout and Static Showcase Summary",
        "",
        "## Scope",
        "",
        "- This stage only adjusted the final Word layout, repository links, project screenshots, and the static showcase page.",
        "- No new experiments were added.",
        "- No notebooks were rerun.",
        "- No `results/*.csv` core metric files were modified by this stage script.",
        "- Classification, regression, productivity, and clustering conclusions were kept unchanged.",
        "",
        "## Word Layout Fixes",
        "",
        "- Started from the user-revised DOCX at `D:/resourcessssss/.../秦天examination report.docx`.",
        "- Rebuilt the cover as a clean standalone first page with the Big Data course-exam fields.",
        "- Added page breaks so Abstract, Table of Contents, and the body start on separate pages.",
        "- Replaced the plain TOC with a static dot-leader TOC with page numbers.",
        "- Cleared repeated page headers and set a dynamic footer as `Page X of Y`.",
        "- Normalized body paragraphs to left alignment to avoid stretched English word spacing.",
        "- Cleaned LaTeX table-width remnants such as `p0.44p0.16` and `p0.20p0.32`.",
        "- Added GitHub repository, final-report, LaTeX-package, and static-showcase links.",
        "- Inserted project-structure, notebook/source-code, web-homepage, and gallery screenshots.",
        "",
        "## Static Showcase",
        "",
        "- Generated `web_demo/index.html` as a pure static HTML/CSS/JS page.",
        "- Generated `web_demo/README_WEB_DEMO.md`.",
        f"- Copied {len(copied_figures)} existing figure artifacts into `web_demo/assets/figures/`.",
        "- Generated four screenshot PNGs under `web_demo/assets/screenshots/`.",
        "- The page does not require a server, GPU, checkpoint, backend, or online inference.",
        "",
        "## DOCX Structure Counts",
        "",
        f"- Paragraphs: {doc_stats['paragraphs']}",
        f"- Word tables: {doc_stats['tables']}",
        f"- Inline images/figures: {doc_stats['inline_shapes']}",
        f"- Code snippet labels: {doc_stats['code_blocks']}",
        f"- Inserted Stage 9 screenshots: {len(audit['screenshots'])}",
        "",
        "## Audit Results",
        "",
        f"- `p0.*` table residual count: {audit['p0_hits']}",
        f"- AI course residual hits: {json.dumps(audit['ai_hits'], ensure_ascii=False)}",
        f"- LaTeX residual hits: {json.dumps(audit['latex_hits'], ensure_ascii=False)}",
        f"- Missing figure text count: {audit['missing_figure_hits']}",
        f"- TODO text count: {audit['todo_hits']}",
        f"- Chinese paragraph hits: {json.dumps(audit['chinese_paragraph_hits'], ensure_ascii=False)}",
        f"- GitHub link present: {audit['github_present']}",
        f"- `web_demo/index.html` exists: {audit['web_demo_present']}",
        f"- `README_WEB_DEMO.md` exists: {audit['readme_present']}",
        f"- Screenshot files: {', '.join(audit['screenshots'])}",
        "",
        "## PDF Export",
        "",
        f"- PDF status: {pdf_result['status']}",
        f"- PDF path: `{pdf_result['path']}`",
        f"- Export message: {pdf_result['message'] or 'n/a'}",
        "",
        "## Screenshot Generation Notes",
        "",
    ]
    if screenshot_failures:
        for key, value in screenshot_failures.items():
            lines.append(f"- `{key}`: {value}")
    else:
        lines.append("- The homepage screenshot was generated with Microsoft Edge headless.")
        lines.append("- The gallery screenshot was composed from the static page figure assets for readability.")
    lines.extend(
        [
            "",
            "## Manual Follow-up",
            "",
            "- Fill Student ID / Name / College / Major and Grade / Institution / Grade and major in Word or WPS.",
            "- Open the DOCX in Word/WPS and manually inspect the page layout.",
            "- If the PDF generated here is not acceptable or if personal information is added later, export the final PDF again from Word/WPS.",
            "- Submit the DOCX/PDF and the complete project directory or GitHub repository link as required by the teacher.",
        ]
    )
    SUMMARY.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    ensure_dirs()
    copied = copy_figures()
    write_web_demo(copied)
    generate_tree_screenshots()
    screenshot_failures = capture_web_screenshots()
    doc_stats = save_docx_from_user()
    audit = audit_docx()
    pdf_result = try_word_pdf_export()
    write_summary(copied, screenshot_failures, doc_stats, audit, pdf_result)
    print(json.dumps({"doc_stats": doc_stats, "audit": audit, "pdf": pdf_result}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
