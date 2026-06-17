"""Rewrite the final DOCX into a standard data-analysis workflow report."""

from __future__ import annotations

import re
import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
REPO_ROOT = PROJECT_ROOT.parent
TEMPLATE_PATH = REPO_ROOT / "期末报告资料" / "课程模板" / "QinTian_experiment.docx"
OUTPUT_DOCX = PROJECT_ROOT / "final_submit" / "大数据分析与应用期末考查报告.docx"
FIGURE_DIR = PROJECT_ROOT / "figures" / "final_report"

REPORT_TITLE = (
    "High-Risk Identification, Digital Dependence Prediction, and User Profiling "
    "in Digital Lifestyles"
)


def clear_document_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def clear_headers_and_footers(doc: Document) -> None:
    for section in doc.sections:
        for part in [section.header, section.footer]:
            for paragraph in part.paragraphs:
                paragraph.text = ""
            for table in list(part.tables):
                part._element.remove(table._tbl)


def ensure_style(doc: Document, name: str, style_type=WD_STYLE_TYPE.PARAGRAPH):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, style_type)


def set_run_font(run, size: float = 11, bold: bool = False, color: str | None = None, font: str = "Calibri") -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    clear_headers_and_footers(doc)

    normal = ensure_style(doc, "Normal")
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.12
    normal.paragraph_format.space_after = Pt(5)

    for style_name, size, color in [
        ("Heading 1", 15, "1F4D78"),
        ("Heading 2", 12.5, "2E74B5"),
        ("Heading 3", 11.5, "1F4D78"),
    ]:
        style = ensure_style(doc, style_name)
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def add_centered(doc: Document, text: str, size: float = 12, bold: bool = False, after: float = 6) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(after)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)


def add_heading(doc: Document, text: str, level: int, page_break: bool = False) -> None:
    if page_break:
        doc.add_page_break()
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    run = paragraph.add_run(text)
    set_run_font(run, size=15 if level == 1 else 12.5 if level == 2 else 11.5, bold=True, color="1F4D78" if level == 1 else "2E74B5")


def add_para(doc: Document, text: str, bold_lead: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.12
    paragraph.paragraph_format.space_after = Pt(5)
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = paragraph.add_run(text[len(bold_lead):])
        set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.28)
    paragraph.paragraph_format.first_line_indent = Inches(-0.12)
    run = paragraph.add_run(f"- {text}")
    set_run_font(run)


def set_cell_text(cell, text: str, font_size: float = 10.5, bold: bool = False, align: str = "center") -> None:
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == "center" else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run_font(run, size=font_size, bold=bold)


def shade_cell(cell, fill: str = "F2F4F7") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_cover(doc: Document) -> None:
    add_centered(doc, "湖北大学 2025--2026 学年度第 2 学期课程考查试题纸", size=15, after=8)
    add_centered(doc, "The paper of course exam", size=12, after=18)

    info = doc.add_table(rows=4, cols=4)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ("Name:", "Big Data Analysis and Applications", "", ""),
        ("Content:", "Course Report", "", ""),
        ("Institution:", "________________", "Teacher:", "Li Jie"),
        ("Grade and major:", "________________", "", ""),
    ]
    for row, values in zip(info.rows, rows):
        for cell, value in zip(row.cells, values):
            set_cell_text(cell, value, bold=value.endswith(":") or value == "Teacher:", align="left")
    doc.add_paragraph()

    add_centered(doc, "Report Topic", size=15, bold=True, after=8)
    add_centered(doc, REPORT_TITLE, size=13, bold=True, after=18)

    personal = doc.add_table(rows=6, cols=2)
    personal.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row, (label, value) in zip(
        personal.rows,
        [
            ("Student ID:", "____________________________"),
            ("Name:", "____________________________"),
            ("College:", "____________________________"),
            ("Major and Grade:", "____________________________"),
            ("Institution:", "____________________________"),
            ("Grade and major:", "____________________________"),
        ],
    ):
        set_cell_text(row.cells[0], label, bold=True, align="right")
        set_cell_text(row.cells[1], value, align="left")
    doc.add_paragraph()

    score = doc.add_table(rows=2, cols=3)
    score.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text in zip(score.rows[0].cells, ["Teacher's comments", "Total score", "Grading teacher"]):
        set_cell_text(cell, text, bold=True)
        shade_cell(cell)
    for cell in score.rows[1].cells:
        set_cell_text(cell, "________________")
    doc.add_paragraph()
    add_centered(doc, "湖北大学", size=12)
    doc.add_page_break()


def add_placeholder(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, bold=True, color="9B1C1C", font="Consolas")


def add_code_placeholder(doc: Document, code_no: int, title: str) -> None:
    add_placeholder(doc, f"[PASTE CODE{code_no} HERE FROM report_code_snippets.md: Code{code_no} {title}]")


def add_code_explanation(doc: Document, purpose: str, necessity: str, output: str, support: str) -> None:
    add_para(doc, f"Explanation: This code is used to {purpose}.")
    add_para(doc, f"It is necessary because {necessity}.")
    add_para(doc, f"The output is {output}.")
    add_para(doc, f"This result supports {support}.")


def add_table_placeholder(doc: Document, table_no: int, title: str, path: str, purpose: str, interpretation: str, conclusion: str) -> None:
    add_placeholder(doc, f"[INSERT TABLE{table_no} SCREENSHOT HERE: {path}]")
    add_para(doc, f"Table{table_no} {title}", bold_lead=f"Table{table_no}")
    add_para(doc, f"Purpose: {purpose}", bold_lead="Purpose:")
    add_para(doc, f"Interpretation: {interpretation}", bold_lead="Interpretation:")
    add_para(doc, f"Conclusion: {conclusion}", bold_lead="Conclusion:")


def add_figure_block(doc: Document, fig_no: int, title: str, filename: str, purpose: str, interpretation: str, conclusion: str) -> None:
    path = FIGURE_DIR / filename
    if path.exists():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(path), width=Inches(5.85))
    else:
        add_placeholder(doc, f"[INSERT FIG{fig_no} HERE: figures/final_report/{filename}]")
    add_para(doc, f"Fig{fig_no} {title}", bold_lead=f"Fig{fig_no}")
    add_para(doc, f"Purpose: {purpose}", bold_lead="Purpose:")
    add_para(doc, f"Interpretation: {interpretation}", bold_lead="Interpretation:")
    add_para(doc, f"Conclusion: {conclusion}", bold_lead="Conclusion:")


def add_report_workflow(doc: Document) -> None:
    add_heading(doc, "Report Workflow", 1)
    add_para(
        doc,
        "This report strictly follows a complete standard data analysis workflow: defining the analysis objective, introducing the dataset, preprocessing the data, conducting statistical analysis and visualization, building and tuning machine learning models, evaluating model results, analyzing findings, and summarizing personal reflection. Each major step explains its experimental purpose, output, interpretation, and connection to later analysis. Core code snippets are embedded as clear paste positions with detailed explanations, while the complete runnable code is placed in Appendix A.",
    )


def add_objective(doc: Document) -> None:
    add_heading(doc, "1. Research Objective and Task Design", 1)
    add_heading(doc, "1.1 Background", 2)
    add_para(
        doc,
        "Digital lifestyle data records how people use devices, receive notifications, spend time on social media, study, sleep, and keep physical activity. These behaviors are not isolated numbers. They form a daily pattern that can be explored with standard data analysis methods.",
    )
    add_para(
        doc,
        "The purpose of this report is to build a complete course-level workflow. The report does not only compare model scores. It starts from dataset selection and data checking, then moves to visualization, modeling, result explanation, and reflection.",
    )
    add_heading(doc, "1.2 Research Questions", 2)
    for item in [
        "Can high_risk_flag be screened from non-leakage behavioral and lifestyle features?",
        "Can digital_dependence_score be predicted from digital behavior and lifestyle variables?",
        "Can productivity_score be explained by the same current feature set?",
        "Can clustering summarize users into readable digital lifestyle profiles?",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "1.3 Modeling Tasks", 2)
    add_para(
        doc,
        "The report implements classification, regression, and clustering. Classification is used for High Risk screening. Regression is used for digital dependence prediction and productivity weak-prediction checking. Clustering is used for exploratory user profiling. These three tasks are connected by the same preprocessing and feature-selection logic.",
    )


def add_dataset(doc: Document) -> None:
    add_heading(doc, "2. Dataset Description and Task Feasibility", 1, page_break=True)
    add_heading(doc, "2.1 Dataset Source", 2)
    add_para(
        doc,
        "The dataset is the 2025 Digital Lifestyle Benchmark Dataset. It is a CSV structured dataset with 3500 rows and 24 fields. It was selected because it naturally supports classification, regression, clustering, preprocessing evidence, and visualization in one coherent digital lifestyle scenario.",
    )
    add_para(
        doc,
        "Other candidate materials were reviewed, but the final dataset was selected because it better supports the complete workflow of classification, regression, clustering, preprocessing, and visualization.",
    )
    add_heading(doc, "2.2 Raw Data Evidence", 2)
    add_para(
        doc,
        "The raw data evidence is kept as an Excel screenshot file instead of a large Word table. This makes the report closer to a real analysis process: first inspect the original table, then clean it, transform it, visualize it, and finally model it.",
    )
    add_table_placeholder(
        doc,
        1,
        "Raw Dataset Preview",
        "screenshot_tables/table1_raw_dataset_preview.xlsx",
        "This table screenshot shows the original CSV structure before preprocessing.",
        "The raw dataset contains digital behavior variables, lifestyle variables, and target variables.",
        "The raw data structure supports classification, regression, and clustering, so the dataset meets the multi-task requirement.",
    )
    add_heading(doc, "2.3 Field Structure", 2)
    add_table_placeholder(
        doc,
        2,
        "Dataset Fields",
        "screenshot_tables/table2_dataset_fields.xlsx",
        "This table screenshot explains the meaning and type of each field before analysis.",
        "The dataset includes behavior variables, lifestyle variables, background variables, and target variables such as high_risk_flag, digital_dependence_score, and productivity_score.",
        "The field design is suitable for a full workflow because it provides both input features and several possible target variables.",
    )
    add_heading(doc, "2.4 Task Feasibility", 2)
    add_para(
        doc,
        "The dataset supports high_risk_flag classification, digital_dependence_score regression, productivity_score weak-prediction checking, and clustering based on behavior and lifestyle features. This is why it fits the final report better than a dataset that only supports one classification task.",
    )


def add_preprocessing(doc: Document) -> None:
    add_heading(doc, "3. Data Preprocessing and Feature Construction", 1, page_break=True)
    add_para(
        doc,
        "This section turns the raw CSV file into a reliable modeling table. The main purpose is to check data quality, construct useful behavior features, and control target leakage before modeling.",
    )
    add_heading(doc, "3.1 Data Loading", 2)
    add_para(
        doc,
        "This step reads the original CSV and confirms the basic shape of the dataset. The output is 3500 rows and 24 fields. The raw preview prepared in Table1 is the input evidence for this loading step.",
    )
    add_code_placeholder(doc, 1, "Data Loading and Basic Inspection")
    add_code_explanation(
        doc,
        "load the CSV file and inspect its shape, columns, dtypes, and first rows",
        "the analysis should start from the raw table instead of assuming the data is already correct",
        "the dataset shape, column list, data types, and preview rows",
        "the dataset introduction and later preprocessing steps",
    )
    add_heading(doc, "3.2 Missing and Duplicate Check", 2)
    add_para(doc, "This step checks missing values and duplicate records. The current result is missing values = 0, duplicate rows = 0, and duplicate id = 0.")
    add_table_placeholder(
        doc,
        3,
        "Missing Value Check",
        "screenshot_tables/table3_missing_value_check.xlsx",
        "This screenshot records missing counts and missing rates for all fields.",
        "Every field has zero missing values, so no row has to be removed for missing data.",
        "The clean missing-value result allows the modeling pipeline to focus on feature selection and evaluation.",
    )
    add_table_placeholder(
        doc,
        4,
        "Duplicate Check",
        "screenshot_tables/table4_duplicate_check.xlsx",
        "This screenshot checks whether repeated rows or repeated ids exist.",
        "The duplicate row count and duplicate id count are both 0.",
        "Because the duplicate check passes, the full 3500 records are kept for the following analysis.",
    )
    add_code_placeholder(doc, 2, "Missing, Duplicate, and Range Check")
    add_code_explanation(
        doc,
        "calculate missing values, duplicated rows, duplicated ids, and basic numeric ranges",
        "data quality problems can distort visualizations and model evaluation",
        "a missing-value table, duplicate counts, and numeric min/max summaries",
        "the decision to keep all records and continue to feature engineering",
    )
    add_heading(doc, "3.3 Range and Rationality Check", 2)
    add_para(
        doc,
        "The range check covers age, device_hours_per_day, phone_unlocks, notifications_per_day, social_media_mins, study_mins, sleep_hours, sleep_quality, productivity_score, and digital_dependence_score. The goal is not to make the data look perfect. The goal is to check whether the numbers are reasonable for this benchmark dataset.",
    )
    add_table_placeholder(
        doc,
        5,
        "Range and Rationality Check",
        "screenshot_tables/table5_range_and_rationality_check.xlsx",
        "This screenshot shows min, max, mean, reasonable range, and pass status for key numerical fields.",
        "The checked variables stay inside the expected ranges, so no serious abnormal record is removed.",
        "The range check supports later statistical analysis because the main numerical fields are usable.",
    )
    add_heading(doc, "3.4 Feature Engineering", 2)
    add_para(
        doc,
        "Feature engineering extracts behavior ratios and interaction information from the original columns. The new features are social_media_hours, study_hours, notifications_per_device_hour, unlocks_per_device_hour, device_to_sleep_ratio, activity_sleep_interaction, and social_to_study_ratio.",
    )
    add_table_placeholder(
        doc,
        6,
        "Engineered Features Preview",
        "screenshot_tables/table6_engineered_features_preview.xlsx",
        "This screenshot shows the engineered features beside key original behavior variables.",
        "The new features transform raw minutes and counts into more interpretable ratios and interaction terms.",
        "These features help classification, regression, and clustering use behavior intensity rather than only raw counts.",
    )
    add_code_placeholder(doc, 3, "Feature Engineering")
    add_code_explanation(
        doc,
        "create behavior-hour, ratio, and interaction features from the raw columns",
        "raw variables such as minutes and counts do not always describe intensity directly",
        "seven engineered features added to the processed dataset",
        "later modeling and clustering by adding behavior-density information",
    )
    add_heading(doc, "3.5 Feature Selection and Leakage Control", 2)
    add_para(
        doc,
        "Feature selection is also leakage control. For high_risk_flag classification, the input cannot use digital_dependence_score, productivity_score, anxiety_score, depression_score, stress_level, happiness_score, focus_score, id, or high_risk_flag itself. Otherwise the model would learn from outcome information rather than behavior evidence.",
    )
    add_table_placeholder(
        doc,
        7,
        "Feature Selection and Leakage Control",
        "screenshot_tables/table7_feature_selection_and_leakage_control.xlsx",
        "This screenshot records dropped columns and retained features for the main tasks.",
        "The classification task removes outcome columns. The clustering task uses only digital behavior and lifestyle numerical features.",
        "This step makes the modeling results more credible because the input features match the purpose of each task.",
    )
    add_code_placeholder(doc, 4, "Leakage Control and Task-Specific Feature Selection")
    add_code_explanation(
        doc,
        "define task-specific drop columns and clustering feature lists",
        "classification, regression, and clustering should not use the same input columns blindly",
        "separate feature sets for classification, regression, and clustering",
        "fairer model evaluation and more meaningful user-profile interpretation",
    )
    add_heading(doc, "3.6 PCA and Probabilistic Clustering Extension", 2)
    add_para(
        doc,
        "PCA is used only for dimensionality reduction visualization and structure understanding. The first two principal components explain about 42.41% variance. PCA is not used as input for classification or regression. Because the clustering inputs are mainly continuous numerical variables, a separate LCA model is not used. GaussianMixture is used as a probabilistic clustering extension to respond to the latent-group idea.",
    )
    add_table_placeholder(
        doc,
        8,
        "PCA Explained Variance",
        "screenshot_tables/table8_pca_explained_variance.xlsx",
        "This screenshot records the explained variance ratio and cumulative explained variance of PCA components.",
        "PC1 and PC2 explain about 42.41% of the total variance, so a two-dimensional PCA view only captures part of the high-dimensional behavior structure.",
        "The PCA evidence supports dimensionality-reduction discussion, but PCA is used only for auxiliary understanding rather than replacing the original modeling features.",
    )
    add_para(
        doc,
        "After preprocessing, the dataset keeps all 3500 records, adds behavior-intensity features, and separates task-specific feature sets. Therefore, the later models are built on a clean and leakage-controlled table.",
    )


def add_eda(doc: Document) -> None:
    add_heading(doc, "4. Statistical Exploration and Visualization", 1, page_break=True)
    add_para(
        doc,
        "This section uses statistical summaries and visualizations to understand the mathematical structure of the dataset before modeling. The goal is to connect raw variables with later modeling choices instead of using figures as decoration.",
    )
    add_table_placeholder(
        doc,
        9,
        "Descriptive Statistical Summary",
        "screenshot_tables/table9_descriptive_statistical_summary.xlsx",
        "This screenshot provides mathematical descriptive statistics before visual analysis.",
        "The variables have different means, ranges, and standard deviations. For example, notifications, phone unlocks, social media time, and device hours are measured on different scales.",
        "The descriptive statistics support later visualization and also explain why scaling is needed before PCA and clustering.",
    )
    add_heading(doc, "4.1 Target Distribution", 2)
    add_figure_block(
        doc,
        1,
        "High Risk vs No Risk Sample Structure",
        "fig1_high_risk_no_risk_distribution.png",
        "This figure checks whether the classification target is balanced before model training.",
        "The No Risk group is larger than the High Risk group. A model may get acceptable Accuracy even when it misses many High Risk samples.",
        "The classification task should focus on Recall, F1, PR-AUC, and the confusion matrix instead of relying only on Accuracy.",
    )
    add_heading(doc, "4.2 Numeric Feature Distributions", 2)
    add_figure_block(
        doc,
        2,
        "Core Digital Lifestyle Feature Distributions",
        "fig2_core_numeric_distributions.png",
        "This figure checks the distributions of the main numerical behavior and outcome variables.",
        "Device use, notifications, social media time, sleep, and digital dependence have different ranges and shapes.",
        "The distribution differences explain why scaling is needed before PCA and clustering, and why raw counts should be interpreted carefully.",
    )
    add_heading(doc, "4.3 Correlation Structure", 2)
    add_figure_block(
        doc,
        3,
        "Correlation Structure of Digital Lifestyle Variables",
        "fig3_correlation_heatmap.png",
        "This heatmap explores the mathematical relationships between numerical variables before modeling.",
        "Some digital behavior variables are related to digital_dependence_score and high_risk_flag, while productivity_score shows weaker relationships.",
        "This result explains why digital_dependence_score is easier to predict than productivity_score, because the correlation pattern around productivity_score is much weaker. The heatmap also warns that correlation is not causation.",
    )
    add_heading(doc, "4.4 High Risk and No Risk Group Difference", 2)
    add_figure_block(
        doc,
        4,
        "Behavioral Difference Between High Risk and No Risk Groups",
        "fig4_high_vs_no_risk_boxplots.png",
        "This figure compares behavior and lifestyle variables between High Risk and No Risk samples.",
        "The High Risk group tends to differ in device behavior and digital dependence, while some variables overlap between groups.",
        "The overlap between High Risk and No Risk groups explains why the classifier cannot be perfect. It also supports threshold tuning instead of relying on a default decision boundary.",
    )
    add_heading(doc, "4.5 EDA Summary", 2)
    add_para(
        doc,
        "The EDA results are not decoration. They explain why classification should use Recall and F1, why digital_dependence_score is more predictable than productivity_score, and why clustering needs scaling and cautious interpretation.",
    )
    add_para(
        doc,
        "The statistical summary and visualizations show three useful signals: class imbalance for classification, stronger behavior dependence for digital_dependence_score, and weaker visible patterns for productivity_score.",
    )


def add_modeling(doc: Document) -> None:
    add_heading(doc, "5. Modeling, Tuning, and Evaluation", 1, page_break=True)
    add_heading(doc, "5.1 Experimental Setup", 2)
    add_para(
        doc,
        "The supervised tasks use train/test split, random_state=42, cross-validation, and GridSearchCV-style hyperparameter tuning. Different tasks use different evaluation metrics. Classification uses Precision, Recall, F1, PR-AUC, ROC-AUC, Balanced Accuracy, and confusion matrix. Regression uses R2, MSE, RMSE, and MAE. Clustering uses elbow method, Silhouette, Calinski-Harabasz, and Davies-Bouldin.",
    )
    add_para(
        doc,
        "The model selection logic is different for each task. Classification focuses on recall-oriented screening and PR-AUC; regression focuses on R2, MSE, and MAE; clustering focuses on silhouette, elbow trend, and profile interpretability.",
    )
    add_heading(doc, "5.2 High Risk Classification", 2)
    add_para(
        doc,
        "The classification target is high_risk_flag. To avoid target leakage, digital_dependence_score, productivity_score, anxiety_score, depression_score, stress_level, happiness_score, focus_score, id, and high_risk_flag itself are excluded from the classification input. Logistic Regression, Random Forest, and Gradient Boosting are compared.",
    )
    add_para(
        doc,
        "Gradient Boosting is selected because it provides the most suitable screening trade-off after model comparison and threshold tuning, not because it is perfect on every metric. Accuracy alone is not enough because the High Risk group is smaller and missed High Risk samples matter more in a screening task.",
    )
    add_table_placeholder(
        doc,
        10,
        "Classification Model Comparison and Threshold Results",
        "screenshot_tables/table10_classification_model_comparison_and_threshold_results.xlsx",
        "This screenshot records candidate model comparison, threshold strategies, final test metrics, and confusion-matrix counts.",
        "The final Gradient Boosting threshold=0.14 gives Recall=0.6420, F1=0.5355, and PR-AUC=0.5084.",
        "The result is not perfect, but it makes sense as a recall-oriented High Risk screening strategy rather than a final individual-level judgment.",
    )
    add_figure_block(
        doc,
        5,
        "Threshold Tuning for Recall-Oriented Screening",
        "fig5_threshold_tuning.png",
        "This figure explains why the final classifier does not use the default threshold of 0.50.",
        "When the threshold decreases, Recall increases because more samples are predicted as High Risk. Precision may decrease because more No Risk samples are also marked as High Risk.",
        "The selected threshold is not chosen because it gives the highest Accuracy, but because it better matches the screening purpose. The threshold of 0.14 is selected to find more true High Risk samples, even though it increases false positives.",
    )
    add_figure_block(
        doc,
        6,
        "Confusion Matrix of the Final High Risk Classifier",
        "fig6_confusion_matrix.png",
        "This figure shows the final No Risk / High Risk prediction results.",
        "The matrix separates true No Risk, false High Risk, missed High Risk, and true High Risk cases. Recall=0.6420 means about 64.20% of true High Risk samples are identified.",
        "The confusion matrix makes the trade-off visible: lowering the threshold improves High Risk detection but creates more false alarms.",
    )
    add_figure_block(
        doc,
        7,
        "Precision-Recall Curve under Class Imbalance",
        "fig7_precision_recall_curve.png",
        "This figure evaluates High Risk ranking ability under different thresholds.",
        "The selected threshold lies on the curve as a recall-oriented point. PR-AUC=0.5084 means the model has some ability to rank High Risk samples above No Risk samples.",
        "The curve supports using PR-AUC together with Recall and F1 instead of relying only on Accuracy.",
    )
    add_code_placeholder(doc, 5, "Classification Model Comparison and Threshold Tuning")
    add_code_explanation(
        doc,
        "train the classifier with cross-validation and choose a recall-oriented threshold",
        "the default threshold may miss too many High Risk samples",
        "the tuned model and threshold-specific Precision, Recall, and F1 values",
        "the final threshold=0.14 screening strategy",
    )
    add_para(
        doc,
        "The classification result is not perfect, but it makes sense for a screening task. Lowering the threshold from 0.50 to 0.14 increases the chance of finding true High Risk samples. The cost is that more No Risk samples may be predicted as High Risk. Therefore, this model is more suitable as a screening reference rather than a final individual-level conclusion.",
    )
    add_heading(doc, "5.3 Digital Dependence and Productivity Regression", 2)
    add_para(
        doc,
        "The main regression target is digital_dependence_score. Gradient Boosting reaches R2=0.9839, MSE=3.1471, and MAE=0.9982. It is selected because it gives the strongest overall performance under the report's main criterion, especially R2 and MSE.",
    )
    add_para(
        doc,
        "Although some linear models may have a slightly lower MAE in digital dependence prediction, Gradient Boosting is kept as the final model because its overall R2 and MSE performance is stronger and it matches the non-linear feature setting. The auxiliary productivity_score task has R2=-0.0041, so it is kept as a weak-prediction result.",
    )
    add_table_placeholder(
        doc,
        11,
        "Regression Model Comparison",
        "screenshot_tables/table11_regression_model_comparison.xlsx",
        "This screenshot compares models for digital_dependence_score and productivity_score using R2, MSE, RMSE, MAE, and cross-validation scores.",
        "digital_dependence_score is strongly predicted, but productivity_score is not explained well by the current features.",
        "The contrast is important because it shows that the same feature set can strongly represent digital dependence but cannot explain productivity well.",
    )
    add_figure_block(
        doc,
        8,
        "Digital Dependence: Observed vs Predicted Scores",
        "fig8_digital_dependence_observed_predicted.png",
        "This figure checks whether predicted digital_dependence_score values are close to observed values.",
        "Most points are close to the diagonal line, which is consistent with the high R2=0.9839.",
        "The regression result supports using digital behavior features to predict digital dependence, but it should not be written as a causal conclusion.",
    )
    add_figure_block(
        doc,
        9,
        "Productivity Weak Prediction: Observed vs Predicted Scores",
        "fig9_productivity_observed_predicted.png",
        "This figure checks whether the current feature set can predict productivity_score as well as digital_dependence_score.",
        "The points are widely scattered instead of staying close to the diagonal reference line. The nearly horizontal prediction pattern suggests that the model tends to predict values near the average productivity level instead of capturing individual productivity differences. This pattern is consistent with the weak R2=-0.0041 result.",
        "The weak prediction is still meaningful because productivity may depend on motivation, task difficulty, learning environment, self-control, and work context, which are not fully captured by the current digital behavior features.",
    )
    add_code_placeholder(doc, 6, "Regression Model Evaluation")
    add_code_explanation(
        doc,
        "calculate R2, MSE, RMSE, and MAE for regression predictions",
        "regression quality cannot be judged by classification metrics",
        "a metrics table for each regression target",
        "the comparison between strong digital dependence prediction and weak productivity prediction",
    )
    add_para(
        doc,
        "The regression results show two different outcomes. Digital dependence can be strongly represented by the current behavioral features, but productivity cannot. This contrast is important because model performance depends on whether the selected features contain enough information for the target variable.",
    )
    add_heading(doc, "5.4 Digital Lifestyle Clustering", 2)
    add_para(
        doc,
        "Clustering is used to build exploratory digital lifestyle profiles. The input uses only digital behavior and lifestyle numerical features. It does not use high_risk_flag, digital_dependence_score, or productivity_score as clustering input because those variables are outcomes for interpretation, not grouping features.",
    )
    add_para(
        doc,
        "KMeans, AgglomerativeClustering, and GaussianMixture are compared. Candidate k values are evaluated with the elbow method, Silhouette, Calinski-Harabasz, and Davies-Bouldin. KMeans k=3 is selected because it gives the best tested silhouette result and produces readable profiles, but the low silhouette value means the profiles should be treated as exploratory summaries.",
    )
    add_table_placeholder(
        doc,
        12,
        "Clustering Model Comparison and Cluster Profiles",
        "screenshot_tables/table12_clustering_model_comparison_and_cluster_profiles.xlsx",
        "This screenshot records KMeans, AgglomerativeClustering, GaussianMixture, k-selection metrics, and the compact cluster profiles.",
        "Cluster 0 is a High social-media-use profile, Cluster 1 is a High device-dependence profile, and Cluster 2 is a Low-load balanced profile.",
        "The profiles are readable, but Silhouette=0.1860 is low, so the result is an exploratory summary rather than a strict grouping rule.",
    )
    add_figure_block(
        doc,
        10,
        "KMeans k Selection by Elbow and Silhouette",
        "fig10_kmeans_k_selection.png",
        "This figure combines elbow information and silhouette score to choose k.",
        "The inertia curve decreases as k increases, while the silhouette curve is highest around k=3 among the tested values.",
        "Although k=3 is selected, the silhouette value is still low. Therefore, k selection supports exploration, not strict segmentation.",
    )
    add_figure_block(
        doc,
        11,
        "Three Digital Lifestyle Cluster Profiles",
        "fig11_cluster_profile_heatmap.png",
        "This figure summarizes the average behavior patterns of the three clusters.",
        "Cluster 0 has high social media minutes. Cluster 1 has higher device hours, lower sleep quality, higher High Risk ratio, and higher digital dependence. Cluster 2 has lower device/social-media load and better sleep-related values.",
        "The heatmap makes the cluster names understandable even though the cluster boundaries are not strong.",
    )
    add_figure_block(
        doc,
        12,
        "PCA Explained Variance for Dimensionality Reduction",
        "fig12_pca_explained_variance.png",
        "This figure records how much variance is explained by the leading principal components.",
        "The first two principal components explain about 42.41% of the total variance.",
        "PCA is useful for auxiliary visualization and structure understanding, but it should not replace the original feature space for modeling.",
    )
    add_code_placeholder(doc, 7, "Clustering Model Comparison and k Selection")
    add_code_explanation(
        doc,
        "standardize clustering features and compare k values using inertia and silhouette",
        "clustering uses distance, so feature scale affects the result",
        "a k-selection table and cluster labels",
        "the final KMeans k=3 exploratory profile analysis",
    )
    add_para(
        doc,
        "The clustering result is useful not because the silhouette score is high, but because it provides a readable summary of lifestyle patterns. The weak silhouette score reminds us that behavioral data may form gradual transitions rather than clear-cut groups.",
    )
    add_heading(doc, "5.5 Integrated Result Analysis", 2)
    add_para(
        doc,
        "The three modeling tasks tell a consistent story. Classification can provide a High Risk screening reference, but more false positives appear when recall is improved. Regression shows that digital dependence is much easier to explain than productivity. Clustering provides readable profiles, but the profile boundaries are weak. The results are not all perfect, but they make sense when connected back to the data and features.",
    )


def add_conclusion(doc: Document) -> None:
    add_heading(doc, "6. Findings, Limitations, and Reflection", 1, page_break=True)
    add_heading(doc, "6.1 Main Findings", 2)
    for item in [
        "High Risk can be screened from behavior and lifestyle features, but the model is a reference rather than a final individual-level conclusion.",
        "digital_dependence_score can be predicted well by the current features, with R2=0.9839, MSE=3.1471, and MAE=0.9982.",
        "Clustering forms three readable lifestyle profiles, but Silhouette=0.1860 shows that the boundaries are weak.",
    ]:
        add_bullet(doc, item)
    add_para(
        doc,
        "The workflow is consistent from beginning to end. Data checks confirm that the original table is usable. Feature engineering turns raw behavior records into behavior-intensity variables. EDA explains why recall, PR-AUC, regression metrics, and clustering metrics are needed. The final models then show that High Risk can be screened as a reference, digital dependence can be strongly predicted, productivity is weakly explained, and lifestyle profiles are readable but not sharply separated.",
    )
    add_heading(doc, "6.2 Practical Value", 2)
    add_para(
        doc,
        "The value of this project is not that every model result is perfect. Its value is that it shows a complete data-analysis process: choosing a dataset, checking data quality, extracting features, visualizing patterns, modeling with suitable metrics, explaining results, and keeping the work reproducible.",
    )
    add_heading(doc, "6.3 Limitations", 2)
    for item in [
        "The dataset is a benchmark/synthetic educational dataset, so the results need external validation before real-world use.",
        "The classification Recall is not 100%, so some High Risk samples are still missed.",
        "productivity_score is weakly predicted, which means the current features are not enough for that target.",
        "The clustering silhouette score is low, so the cluster boundaries are not strong.",
        "No external dataset is used for validation.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "6.4 Personal Reflection", 2)
    add_para(
        doc,
        "Before this report, I tended to think that data analysis mainly meant running models and comparing scores. After completing this project, I realized that a complete big data analysis report should start much earlier: dataset selection, data cleaning, feature extraction, statistical analysis, visualization, model evaluation, and result interpretation are all necessary.",
    )
    add_para(
        doc,
        "The most important lesson for me is that a result does not have to be perfect to be meaningful. For example, the productivity regression result is weak, but it still tells me that the current features are not enough to explain productivity. The clustering silhouette score is also not high, but the cluster profiles still help summarize different digital lifestyle patterns. These results make the report more realistic.",
    )
    add_para(
        doc,
        "I also learned that visualizations should not be used as decoration. Each figure should answer a specific question and support a conclusion. This helped me connect the raw data, statistical patterns, model results, and final interpretation into one complete workflow. I also appreciate the teacher's repeated emphasis on data cleaning, feature extraction, visualization, evaluation, and reproducible reporting.",
    )


def add_appendix(doc: Document) -> None:
    add_heading(doc, "Appendix A Complete Runnable Code", 1, page_break=True)
    add_heading(doc, "A.1 Environment and Dependencies", 2)
    add_para(
        doc,
        "The experiment was completed with a CPU-based Python environment. The main dependencies are pandas, NumPy, scikit-learn, Matplotlib, openpyxl, python-docx, and joblib. The environment can be installed with:",
    )
    add_para(doc, "pip install -r requirements.txt")
    add_para(
        doc,
        "The complete runnable code below can regenerate preprocessing checks, screenshot Excel tables, figures, model metrics, and final result files.",
    )
    add_heading(doc, "A.2 Complete Runnable Code", 2)
    add_para(
        doc,
        "The complete runnable code should be pasted here before final submission. The code should be copied from:",
    )
    add_para(doc, "appendix_A_complete_code.py")
    add_placeholder(doc, "[PASTE COMPLETE RUNNABLE CODE HERE]")


def build_report() -> None:
    doc = Document(str(TEMPLATE_PATH))
    clear_document_body(doc)
    setup_document(doc)
    add_cover(doc)
    add_report_workflow(doc)
    add_objective(doc)
    add_dataset(doc)
    add_preprocessing(doc)
    add_eda(doc)
    add_modeling(doc)
    add_conclusion(doc)
    add_appendix(doc)
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)
    print(f"Saved workflow report to {OUTPUT_DOCX}")


if __name__ == "__main__":
    build_report()
