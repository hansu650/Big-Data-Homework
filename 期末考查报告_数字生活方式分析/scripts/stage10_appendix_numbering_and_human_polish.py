from __future__ import annotations

import json
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
PROJECT = ROOT / "\u671f\u672b\u8003\u67e5\u62a5\u544a_\u6570\u5b57\u751f\u6d3b\u65b9\u5f0f\u5206\u6790"
USER_DOCX = Path(
    "D:/resourcessssss/"
    "\u666e\u901a\u8bfe/"
    "\u5927\u6570\u636e\u5206\u6790\u4e0e\u5e94\u7528/"
    "\u7ec8/"
    "\u79e6\u5929examination report.docx"
)
FINAL_DOCX = PROJECT / "final_submit" / "\u5927\u6570\u636e\u5206\u6790\u4e0e\u5e94\u7528\u671f\u672b\u8003\u67e5\u62a5\u544a.docx"
SUMMARY = PROJECT / "results" / "stage10_appendix_numbering_and_human_polish_summary.md"

PROJECT_STRUCTURE_IMG = PROJECT / "web_demo" / "assets" / "screenshots" / "project_structure_screenshot.png"
NOTEBOOKS_SRC_IMG = PROJECT / "web_demo" / "assets" / "screenshots" / "notebooks_src_screenshot.png"
WEB_HOME_IMG = PROJECT / "web_demo" / "assets" / "screenshots" / "web_demo_homepage_screenshot.png"
WEB_GALLERY_IMG = PROJECT / "web_demo" / "assets" / "screenshots" / "web_demo_gallery_screenshot.png"

GITHUB_URL = "https://github.com/hansu650/Big-Data-Homework"
HTML_PREVIEW_URL = (
    "https://htmlpreview.github.io/?https://github.com/hansu650/Big-Data-Homework/blob/main/"
    "%E6%9C%9F%E6%9C%AB%E8%80%83%E6%9F%A5%E6%8A%A5%E5%91%8A_%E6%95%B0%E5%AD%97%E7%94%9F%E6%B4%BB%E6%96%B9%E5%BC%8F%E5%88%86%E6%9E%90/"
    "web_demo/index.html"
)

SECTION_24 = [
    "The main work of this report can be summarized as follows.",
    "First, I compared the available pfm_train / pfm_test files with the final Digital Lifestyle Benchmark Dataset and selected the latter because it better satisfies the course requirements and supports more task types.",
    "Second, I completed the basic data-quality checks, including missing values, duplicate records, rational value ranges, and task-specific feature selection. I also constructed several behavioral ratio and interaction features to describe digital lifestyle patterns more clearly.",
    "Third, I conducted exploratory analysis before modeling. The EDA part uses descriptive statistics, bar charts, histograms, boxplots, scatter plots, and heatmaps to understand distributions, correlations, group differences, and risk-rate patterns.",
    "Fourth, I built three types of machine learning tasks: classification for high-risk screening, regression for digital dependence prediction, and clustering for lifestyle profiling. The classification and regression tasks include hyperparameter tuning and professional evaluation metrics, while clustering uses k selection, the elbow method, and silhouette analysis.",
    "Finally, I organized the notebooks, source code, figures, result CSV files, Word report, and static showcase page into a reproducible project structure so that the report is supported by traceable experimental materials.",
]

SECTION_75 = [
    "Before working on this report, I tended to think of data analysis mainly as running models and comparing scores. During the course, the teacher repeatedly reminded us that a complete big data analysis report should start from dataset compliance, data cleaning, feature extraction, statistical analysis, visualization, model evaluation, and reproducibility. This changed the way I organized the project.",
    "In this final report, I learned that preprocessing and leakage control are as important as model selection. For example, when predicting high_risk_flag, I could not simply use outcome-related variables such as digital_dependence_score or productivity_score, because that would make the model look better but weaken the meaning of the experiment. I also learned that different metrics serve different purposes: Recall and F1 are more useful for high-risk screening, while R2, MSE, and MAE are more suitable for regression evaluation, and the silhouette coefficient helps judge whether clustering results are reliable.",
    "Another important lesson is that negative results should also be explained honestly. The productivity_score regression result was weak, but this made the report more realistic because it showed that not every target can be predicted well with the same features. I also realized that visualizations are not just decorations; they help connect data patterns with modeling results.",
    "I sincerely appreciate the teacher's clear explanations and repeated reminders about experimental workflow, report organization, and result interpretation. These requirements helped me turn the project from a set of separate notebooks into a complete and reproducible data analysis report.",
]

APPENDIX_B = [
    ("Appendix B Full Experimental Running Instructions and Repository Access", 0),
    ("B.1 Environment Setup and Notebook Execution", 1),
    ("The complete experiment can be reproduced from the project directory. First, install the required Python packages with:", None),
    ("pip install -r requirements.txt", None),
    ("Then run the notebooks in order from 00 to 05. The notebook sequence starts from dataset compliance checking, then proceeds to preprocessing and EDA, classification modeling, regression modeling, clustering analysis, and final result summarization. The script scripts/stage2_5_evidence_enhancement.py can be used to regenerate preprocessing-quality checks, PCA explained-variance evidence, and final report artifact lists.", None),
    ("All experiments use CPU-based scikit-learn workflows. The random_state is fixed at 42 to improve reproducibility. After the experiment files are checked, the final DOCX can be opened in Word or WPS, personal information can be filled in manually, and the report can be exported as PDF.", None),
    ("B.2 Repository Access and Local Opening Notes", 1),
    ("The complete repository is available at:", None),
    (GITHUB_URL, None),
    ("The final project directory is:", None),
    ("Big-Data-Homework / \u671f\u672b\u8003\u67e5\u62a5\u544a_\u6570\u5b57\u751f\u6d3b\u65b9\u5f0f\u5206\u6790", None),
    ("This directory contains the raw and processed data files, notebooks, source code, scripts, figures, result CSV files, the LaTeX package, the Word report, and the static visualization showcase. To inspect the static showcase, open:", None),
    ("web_demo / index.html", None),
    ("The static page is only used for displaying the project workflow and visual results. It does not perform online prediction, does not load a model checkpoint, and does not require GPU or a server.", None),
]

APPENDIX_D_TEXT = (
    "This report was completed independently for the course Big Data Analysis and Applications. "
    "My main work included selecting and verifying the dataset, designing the analytical topic, "
    "cleaning and checking the data, constructing engineered features, performing exploratory analysis, "
    "building classification, regression, and clustering models, organizing the result files, preparing "
    "visualizations, and formatting the final report."
)

APPENDIX_D_TEXT_2 = (
    "The previous experiments and other course reports were used only as references for workflow organization "
    "and formatting style. The final dataset, research topic, task design, feature scheme, figures, model "
    "results, interpretations, and conclusions in this report were developed for this course project. "
    "Personal information is intentionally left blank and should be filled in manually before final submission."
)

PATH_TABLE = [
    ("Artifact", "Location"),
    ("GitHub repository", GITHUB_URL),
    ("Project directory", "Big-Data-Homework / \u671f\u672b\u8003\u67e5\u62a5\u544a_\u6570\u5b57\u751f\u6d3b\u65b9\u5f0f\u5206\u6790"),
    ("Final Word report", "\u671f\u672b\u8003\u67e5\u62a5\u544a_\u6570\u5b57\u751f\u6d3b\u65b9\u5f0f\u5206\u6790 / final_submit / \u5927\u6570\u636e\u5206\u6790\u4e0e\u5e94\u7528\u671f\u672b\u8003\u67e5\u62a5\u544a.docx"),
    ("LaTeX package", "\u671f\u672b\u8003\u67e5\u62a5\u544a_\u6570\u5b57\u751f\u6d3b\u65b9\u5f0f\u5206\u6790 / overleaf_final"),
    ("Static visualization showcase", "\u671f\u672b\u8003\u67e5\u62a5\u544a_\u6570\u5b57\u751f\u6d3b\u65b9\u5f0f\u5206\u6790 / web_demo / index.html"),
    ("Optional HTML preview", HTML_PREVIEW_URL),
]

RESULT_TABLE = [
    ("Artifact group", "Examples"),
    ("EDA figures", "target distribution, histograms, correlation heatmap, boxplots, categorical risk-rate bars, scatter plots"),
    ("Classification figures", "tuned metrics comparison, confusion matrix, PR curve, ROC curve, permutation importance"),
    ("Regression figures", "target comparison, observed-vs-predicted plots, permutation importance"),
    ("Clustering/PCA figures", "elbow curve, silhouette curve, PCA explained variance, PCA cluster plot, profile heatmap"),
]


def set_run_font(run, size: float = 10.5, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def clear_paragraph(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def set_paragraph_text(paragraph, text: str, *, size: float = 10.5, bold: bool = False) -> None:
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)


def insert_paragraph_before(anchor, text: str, *, level: int | None = None):
    p = anchor.insert_paragraph_before()
    set_paragraph_text(p, text, size=heading_size(level), bold=level is not None)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(5)
    if level is not None:
        set_outline_level(p, level)
        for run in p.runs:
            run.font.color.rgb = RGBColor(31, 80, 132)
    return p


def heading_size(level: int | None) -> float:
    if level == 0:
        return 14
    if level == 1:
        return 12
    return 10.5


def set_outline_level(paragraph, level: int | None) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    for old in p_pr.findall(qn("w:outlineLvl")):
        p_pr.remove(old)
    if level is not None:
        outline = OxmlElement("w:outlineLvl")
        outline.set(qn("w:val"), str(level))
        p_pr.append(outline)


def clean_p0_text(text: str) -> str:
    return re.sub(r"p0(?:\.\d+)?(?:p0(?:\.\d+)?)*\s*", "", text)


def run_has_non_text(run) -> bool:
    return bool(run._r.xpath(".//w:drawing")) or bool(run._r.xpath(".//w:pict")) or bool(run._r.xpath(".//w:br"))


def find_paragraph(doc: Document, text: str):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    return None


def find_startswith(doc: Document, prefix: str):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    return None


def next_heading_after(doc: Document, start_text: str, patterns: list[str]):
    found = False
    for p in doc.paragraphs:
        stripped = p.text.strip()
        if not found:
            found = stripped == start_text
            continue
        for pat in patterns:
            if re.match(pat, stripped):
                return p
    return None


def remove_blocks_between(start_para, end_para, include_start: bool = True) -> None:
    body = start_para._p.getparent()
    children = list(body)
    start_idx = children.index(start_para._p)
    end_idx = children.index(end_para._p)
    for child in children[start_idx + (0 if include_start else 1):end_idx]:
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def remove_blocks_after(paragraph) -> None:
    body = paragraph._p.getparent()
    children = list(body)
    idx = children.index(paragraph._p)
    for child in children[idx + 1:]:
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def insert_table_before(anchor, rows: list[tuple[str, str]]):
    container = anchor._parent
    table = container.add_table(rows=len(rows), cols=2, width=Inches(6.2))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for ri, row in enumerate(rows):
        for ci, value in enumerate(row):
            cell = table.cell(ri, ci)
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    set_run_font(run, size=9.2, bold=(ri == 0))
            if ri == 0:
                set_cell_shading(cell, "EAF1F8")
    tbl = table._tbl
    tbl.getparent().remove(tbl)
    anchor._p.addprevious(tbl)
    return table


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for shd in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(shd)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def insert_image_block(anchor, image_path: Path, caption: str, notes: list[str]) -> None:
    p = anchor.insert_paragraph_before()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(5.9))
    cap = insert_paragraph_before(anchor, caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.bold = True
    for note in notes:
        insert_paragraph_before(anchor, note)


def clean_residuals_from_abstract(doc: Document) -> None:
    abstract = find_paragraph(doc, "Abstract")
    in_body = False
    for p in doc.paragraphs:
        if p is abstract:
            in_body = True
        if not in_body:
            continue
        for run in p.runs:
            if not run_has_non_text(run):
                run.text = clean_p0_text(run.text).replace("Table 10-1", "Table C-1")
    # Skip the first two cover tables so the exam-paper frame/template remains untouched.
    for table in doc.tables[2:]:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if not run_has_non_text(run):
                            run.text = clean_p0_text(run.text)


def replace_section_after_heading(doc: Document, heading_text: str, next_patterns: list[str], new_paragraphs: list[str]) -> None:
    heading = find_paragraph(doc, heading_text)
    if not heading:
        raise RuntimeError(f"Heading not found: {heading_text}")
    next_heading = next_heading_after(doc, heading_text, next_patterns)
    if not next_heading:
        raise RuntimeError(f"Next heading not found after: {heading_text}")
    set_paragraph_text(heading, heading_text, size=12, bold=True)
    set_outline_level(heading, 1)
    remove_blocks_between(heading, next_heading, include_start=False)
    for text in new_paragraphs:
        insert_paragraph_before(next_heading, text)


def replace_66_table(doc: Document) -> bool:
    anchor = find_paragraph(doc, "Project structure screenshot")
    if not anchor:
        raise RuntimeError("Project structure screenshot anchor not found.")
    removed = False
    for table in list(doc.tables):
        text = "\n".join(cell.text for row in table.rows for cell in row.cells)
        if "GitHub repository" in text and "Static visualization showcase" in text:
            table._tbl.getparent().remove(table._tbl)
            removed = True
            break
    insert_table_before(anchor, PATH_TABLE)
    return removed


def replace_appendices(doc: Document) -> None:
    appendix_b = find_paragraph(doc, "Appendix B Full Experimental Running Instructions")
    appendix_d = find_paragraph(doc, "Appendix D Personal Contribution and Academic Integrity Statement")
    if not appendix_b or not appendix_d:
        raise RuntimeError("Appendix B or Appendix D anchor not found.")

    remove_blocks_between(appendix_b, appendix_d, include_start=True)

    for text, level in APPENDIX_B:
        insert_paragraph_before(appendix_d, text, level=level)

    insert_paragraph_before(appendix_d, "Appendix C Result Files, Figure List, and Static Visualization Showcase", level=0)
    insert_paragraph_before(appendix_d, "C.1 Result Files and Figure Categories", level=1)
    insert_paragraph_before(
        appendix_d,
        "Important result files include dataset_compliance_summary.csv, preprocessing_quality_check.csv, classification_tuned_metrics.csv, classification_threshold_tuning.csv, regression_target_comparison.csv, regression_digital_dependence_metrics.csv, regression_productivity_metrics.csv, clustering_model_comparison.csv, clustering_lifestyle_profiles_compact.csv, and pca_explained_variance.csv.",
    )
    insert_paragraph_before(appendix_d, "Table C-1 Result file and figure categories")
    insert_table_before(appendix_d, RESULT_TABLE)
    insert_paragraph_before(appendix_d, "C.2 Static Visualization Showcase", level=1)
    for text in [
        "GitHub repository:",
        GITHUB_URL,
        "Static visualization showcase:",
        "web_demo / index.html",
        "Optional HTML preview:",
        HTML_PREVIEW_URL,
        "The static page can be opened locally in a browser. It summarizes the project theme, dataset scale, task types, modeling results, and visualization gallery. It is designed only as a lightweight showcase page, not as an online prediction system.",
    ]:
        insert_paragraph_before(appendix_d, text)
    insert_image_block(
        appendix_d,
        WEB_HOME_IMG,
        "Figure C-1 Static visualization showcase homepage",
        [
            "The homepage gives a quick overview of the dataset, workflow, model tasks, and key results. It helps the reader understand the project before reading all technical details.",
        ],
    )
    insert_image_block(
        appendix_d,
        WEB_GALLERY_IMG,
        "Figure C-2 Visualization gallery in the static showcase page",
        [
            "The gallery collects representative EDA, classification, regression, and clustering figures. It provides a visual summary of the analytical evidence used in the report.",
        ],
    )

    set_paragraph_text(appendix_d, "Appendix D Personal Contribution and Academic Integrity Statement", size=14, bold=True)
    set_outline_level(appendix_d, 0)
    for run in appendix_d.runs:
        run.font.color.rgb = RGBColor(31, 80, 132)
    remove_blocks_after(appendix_d)
    p = doc.add_paragraph()
    set_paragraph_text(p, APPENDIX_D_TEXT)
    p = doc.add_paragraph()
    set_paragraph_text(p, APPENDIX_D_TEXT_2)


def normalize_body_headings(doc: Document) -> None:
    abstract = find_paragraph(doc, "Abstract")
    in_body = False
    appendix_titles = {
        "Appendix A Core Code",
        "Appendix B Full Experimental Running Instructions and Repository Access",
        "Appendix C Result Files, Figure List, and Static Visualization Showcase",
        "Appendix D Personal Contribution and Academic Integrity Statement",
    }
    top_level = re.compile(r"^(Abstract|References|Chapter \d+ .+)$")
    second_level = re.compile(r"^(\d+\.\d+ .+|[BC]\.\d+ .+)$")
    for p in doc.paragraphs:
        if p is abstract:
            in_body = True
        if not in_body:
            continue
        text = p.text.strip()
        if top_level.match(text) or text in appendix_titles:
            set_outline_level(p, 0)
            for run in p.runs:
                set_run_font(run, size=14 if text.startswith(("Chapter", "Appendix")) else 12, bold=True)
                run.font.color.rgb = RGBColor(31, 80, 132)
        elif second_level.match(text):
            set_outline_level(p, 1)
            for run in p.runs:
                set_run_font(run, size=12, bold=True)
                run.font.color.rgb = RGBColor(31, 80, 132)
        else:
            set_outline_level(p, None)


def check_images(doc: Document) -> dict[str, bool]:
    text = "\n".join(p.text for p in doc.paragraphs)
    return {
        "figure_6_1_caption": "Figure 6-1 Project directory and engineering structure" in text,
        "figure_6_2_caption": "Figure 6-2 Notebook and source-code organization" in text,
        "figure_c_1_caption": "Figure C-1 Static visualization showcase homepage" in text,
        "figure_c_2_caption": "Figure C-2 Visualization gallery in the static showcase page" in text,
        "inline_shape_count_at_least_24": len(doc.inline_shapes) >= 24,
    }


def audit_docx(path: Path) -> dict[str, object]:
    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += "\n" + cell.text
    expected_appendices = {
        "Appendix A Core Code",
        "Appendix B Full Experimental Running Instructions and Repository Access",
        "Appendix C Result Files, Figure List, and Static Visualization Showcase",
        "Appendix D Personal Contribution and Academic Integrity Statement",
    }
    appendix_top = [p.text.strip() for p in doc.paragraphs if p.text.strip() in expected_appendices]
    footer_texts = []
    header_texts = []
    for section in doc.sections:
        header_texts.extend(p.text.strip() for p in section.header.paragraphs if p.text.strip())
        footer_texts.extend(p.text.strip() for p in section.footer.paragraphs if p.text.strip())
    return {
        "p0_residual_count": len(re.findall(r"p0(?:\.\d+)?", text)),
        "table_10_1_count": text.count("Table 10-1"),
        "table_c_1_count": text.count("Table C-1"),
        "appendix_top_level": appendix_top,
        "appendix_b_count": sum(1 for item in appendix_top if item.startswith("Appendix B ")),
        "appendix_c_count": sum(1 for item in appendix_top if item.startswith("Appendix C ")),
        "b_subheadings_present": all(item in text for item in ["B.1 Environment Setup and Notebook Execution", "B.2 Repository Access and Local Opening Notes"]),
        "c_subheadings_present": all(item in text for item in ["C.1 Result Files and Figure Categories", "C.2 Static Visualization Showcase"]),
        "github_present": GITHUB_URL in text,
        "html_preview_present": HTML_PREVIEW_URL in text,
        "image_checks": check_images(doc),
        "header_texts": header_texts,
        "footer_texts": footer_texts,
        "fixed_footer_detected": any("\u7b2c1 \u9875" in item or "\u51714\u9875" in item or "\u51714\u9875" in item.replace(" ", "") for item in footer_texts),
        "footer_preserved_by_request": True,
        "inline_shapes": len(doc.inline_shapes),
        "tables": len(doc.tables),
        "paragraphs": len(doc.paragraphs),
    }


def write_summary(audit: dict[str, object], path_table_replaced: bool) -> None:
    lines = [
        "# Stage 10 Appendix Numbering and Human Polish Summary",
        "",
        "## Scope",
        "",
        "- Started from the user's latest DOCX: `D:/resourcessssss/普通课/大数据分析与应用/终/秦天examination report.docx`.",
        "- Preserved the cover page, outside frame, exam-paper template, and page border/header/footer structure.",
        "- Edited only content from `Abstract` onward.",
        "- No new experiments were added.",
        "- No notebooks were rerun.",
        "- No `results/*.csv` core metric files were modified.",
        "",
        "## Fixes",
        "",
        f"- Duplicate Appendix B/C fixed: {audit['appendix_b_count'] == 1 and audit['appendix_c_count'] == 1}",
        f"- Top-level appendices: {json.dumps(audit['appendix_top_level'], ensure_ascii=False)}",
        f"- B.1/B.2 present: {audit['b_subheadings_present']}",
        f"- C.1/C.2 present: {audit['c_subheadings_present']}",
        f"- `p0.*` residual count: {audit['p0_residual_count']}",
        f"- `Table 10-1` count: {audit['table_10_1_count']}",
        f"- `Table C-1` count: {audit['table_c_1_count']}",
        f"- 6.6 path table replaced: {path_table_replaced}",
        f"- Screenshot captions/images checked: {json.dumps(audit['image_checks'], ensure_ascii=False)}",
        f"- Footer texts checked: {json.dumps(audit['footer_texts'], ensure_ascii=False)}",
        f"- Fixed template footer detected: {audit['fixed_footer_detected']}",
        f"- Footer/header/frame preserved by latest user instruction: {audit['footer_preserved_by_request']}",
        f"- Inline image count: {audit['inline_shapes']}",
        f"- Word table count: {audit['tables']}",
        "",
        "## Text Polish",
        "",
        "- Replaced Section 2.4 with a more natural contribution/workflow summary.",
        "- Replaced Section 7.5 with a more natural course-learning reflection.",
        "- Replaced Appendix D with a concise personal contribution and academic integrity statement.",
        "- Replaced Appendix B and Appendix C with merged, numbered appendix structures.",
        "",
        "## Remaining Manual Step",
        "",
        "- Fill personal information in the preserved cover/template fields before final submission.",
    ]
    SUMMARY.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    if not USER_DOCX.exists():
        raise FileNotFoundError(USER_DOCX)
    FINAL_DOCX.parent.mkdir(parents=True, exist_ok=True)
    SUMMARY.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(USER_DOCX, FINAL_DOCX)
    doc = Document(FINAL_DOCX)
    clean_residuals_from_abstract(doc)
    replace_section_after_heading(doc, "2.4 Main Contributions and Technical Workflow", [r"^Chapter 3 "], SECTION_24)
    replace_section_after_heading(doc, "7.5 Course Learning Reflection and Appreciation", [r"^References$"], SECTION_75)
    path_table_replaced = replace_66_table(doc)
    replace_appendices(doc)
    normalize_body_headings(doc)
    clean_residuals_from_abstract(doc)
    doc.save(FINAL_DOCX)
    audit = audit_docx(FINAL_DOCX)
    write_summary(audit, path_table_replaced)
    print(json.dumps(audit, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
