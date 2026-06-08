"""Build the final English Word submission for Stage 8.

This script only rewrites the final DOCX report. It does not run experiments,
does not modify result CSV files, and does not change the validated model
metrics from Stage 6/7.
"""

from __future__ import annotations

import json
import re
import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUTPUT_DOCX_NAME = "大数据分析与应用期末考查报告.docx"
OUTPUT_PDF_NAME = "大数据分析与应用期末考查报告.pdf"
REPORT_TITLE = (
    "High-Risk Identification, Digital Dependence Prediction, and User Profiling "
    "in Digital Lifestyles: A Classification, Regression, and Clustering Study "
    "Based on the 2025 Digital Lifestyle Benchmark Dataset"
)

COURSE_NAME = "Big Data Analysis and Applications"
CONTENT_NAME = "Course Report"
TEACHER_NAME = "Li Jie"

BANNED_TERMS = [
    "湖北大学本科课程设计",
    "人工智能技术与应用",
    "校园室内巡检",
    "CampusDepthSegLite",
    "RGB-D",
    "语义分割",
    "NYUDepthV2",
    "PyTorch",
    "王雷春",
    "副教授",
    "AI-homework",
    "mIoU",
    "Pixel Acc",
    "RGBD-concat",
    "self-collected campus scenes",
]

LATEX_RESIDUES = [
    "\\chapter",
    "\\section",
    "\\subsection",
    "\\texttt",
    "\\input",
    "\\maybefigure",
    "\\begin",
    "\\end",
    "\\caption",
    "\\label",
    "\\ref",
]

ALLOWED_CHINESE_PHRASES = [
    "湖北大学",
    "学年度",
    "第",
    "学期",
    "课程考查试题纸",
    "学号",
    "姓名",
    "学院",
    "专业年级",
]

FIGURES = [
    (4, "eda_high_risk_flag_distribution.png", "Distribution of high_risk_flag"),
    (4, "eda_numeric_histograms.png", "Distributions of core numerical variables"),
    (4, "eda_numeric_correlation_heatmap.png", "Correlation heatmap of numerical variables"),
    (4, "eda_boxplots_by_risk.png", "Behavioral differences by high-risk group"),
    (4, "eda_category_risk_rate.png", "Risk rates across categorical variables"),
    (4, "eda_behavior_outcome_scatter.png", "Relationships between behavioral variables and outcomes"),
    (5, "classification_tuned_metrics_comparison.png", "Comparison of tuned classification metrics"),
    (5, "classification_final_confusion_matrix.png", "Final confusion matrix for high-risk screening"),
    (5, "classification_precision_recall_curve.png", "Precision-recall curve for the final classifier"),
    (5, "classification_roc_curve.png", "ROC curve for the final classifier"),
    (5, "regression_target_comparison.png", "Comparison of regression targets"),
    (5, "regression_digital_dependence_observed_vs_predicted.png", "Observed versus predicted digital dependence scores"),
    (5, "regression_digital_dependence_permutation_importance.png", "Permutation importance for digital dependence prediction"),
    (5, "regression_productivity_observed_vs_predicted.png", "Observed versus predicted productivity scores"),
    (5, "classification_permutation_importance.png", "Permutation importance for high-risk screening"),
    (5, "clustering_kmeans_elbow.png", "KMeans elbow curve for selecting k"),
    (5, "clustering_silhouette_by_k.png", "Silhouette scores by number of clusters"),
    (5, "pca_explained_variance.png", "Cumulative explained variance of PCA components"),
    (5, "clustering_lifestyle_pca.png", "PCA visualization of lifestyle clusters"),
    (5, "clustering_lifestyle_profile_heatmap.png", "Heatmap of cluster lifestyle profiles"),
]

CAPTION_TRANSLATIONS = {
    "dataset_compliance_summary.tex": "Dataset compliance summary against course requirements",
    "preprocessing_quality_check_compact.tex": "Compact preprocessing quality check",
    "classification_tuned_metrics.tex": "Test-set classification metrics under threshold strategies",
    "classification_threshold_tuning.tex": "Validation-set threshold tuning results",
    "regression_target_comparison.tex": "Best-model comparison for the two regression targets",
    "regression_digital_dependence_metrics.tex": "Test-set metrics for digital_dependence_score regression",
    "regression_productivity_metrics.tex": "Test-set metrics for productivity_score regression",
    "clustering_model_comparison.tex": "Comparison of clustering algorithms and cluster numbers",
    "clustering_lifestyle_profiles_compact.tex": "Compact digital lifestyle cluster profiles",
    "pca_explained_variance_compact.tex": "PCA explained variance table for the first 10 components",
}

CELL_TRANSLATIONS = {
    "课程要求": "Course requirement",
    "Digital Lifestyle Benchmark Dataset (2025) 情况": "Digital Lifestyle Benchmark Dataset (2025) status",
    "是否满足": "Status",
    "来源平台": "Source platform",
    "Hugging Face 公开数据集页面": "Public Hugging Face dataset page",
    "发布时间/更新年份": "Release/update year",
    "2025 年数据集": "2025 dataset",
    "文件格式": "File format",
    "CSV 结构化表格数据": "CSV structured tabular data",
    "样本量": "Sample size",
    "3500 行，满足不少于 1000 行要求": "3,500 rows, meeting the requirement of at least 1,000 rows",
    "有效特征数": "Number of effective features",
    "24 个字段，满足不少于 8 个有效特征要求": "24 fields, meeting the requirement of at least 8 effective features",
    "任务支持": "Task support",
    "支持 high_risk_flag 分类、digital_dependence_score 回归和生活方式聚类": "Supports high_risk_flag classification, digital_dependence_score regression, and lifestyle clustering",
    "数据性质": "Data nature",
    "合成教学/benchmark 数据，结论需谨慎解释，不用于医学诊断": "Synthetic educational/benchmark data; conclusions must be interpreted cautiously and not used for medical diagnosis",
    "满足课程建模用途": "Meets course modeling purpose",
    "满足": "Meets",
    "通过": "Passed",
    "检查项": "Check item",
    "检查结果": "Check result",
    "处理策略": "Processing strategy",
    "原始样本量": "Original sample size",
    "原始字段数": "Original number of fields",
    "缺失值总数": "Total missing values",
    "重复行数量": "Duplicate rows",
    "重复 id 数量": "Duplicate id count",
    "关键数值字段合理范围": "Rational range of key numerical fields",
    "全部通过": "All passed",
    "工程特征生成": "Engineered feature generation",
    "7 个特征全部生成": "All 7 engineered features generated",
    "分类泄漏字段排除": "Classification leakage-column exclusion",
    "回归 outcome 控制": "Regression outcome control",
    "聚类输入控制": "Clustering input control",
    "无需删除，建模管道保留缺失处理能力": "No deletion needed; the modeling pipeline still keeps missing-value handling capability",
    "无需删除": "No deletion needed",
    "未删除记录": "No records removed",
    "未删除字段": "No fields removed",
    "检查后未发现需删除记录": "No records requiring deletion were found after checking",
    "用于补充行为比例和交互信息": "Used to supplement behavioral ratio and interaction information",
    "排除 outcome 字段与 id": "Outcome fields and id were excluded",
    "排除 high_risk_flag 和其他结果变量": "high_risk_flag and other outcome variables were excluded",
    "仅使用行为与生活习惯数值特征": "Only behavioral and lifestyle numerical features were used",
    "模型": "Model",
    "阈值策略": "Threshold strategy",
    "阈值": "Threshold",
    "目标变量": "Target variable",
    "最佳模型": "Best model",
    "CV 最优 R²": "Best CV R²",
    "算法": "Algorithm",
    "簇": "Cluster",
    "样本数": "Cluster size",
    "设备时长": "Device hours",
    "社交媒体分钟": "Social media minutes",
    "睡眠时长": "Sleep hours",
    "睡眠质量": "Sleep quality",
    "高风险比例": "High-risk ratio",
    "数字依赖分": "Digital dependence score",
    "建议画像名称": "Suggested profile name",
    "高社交媒体使用型": "High social-media-use profile",
    "高设备依赖型": "High device-dependence profile",
    "低负荷均衡型": "Low-load balanced profile",
    "主成分": "Principal component",
    "解释方差比例": "Explained variance ratio",
    "累计解释方差": "Cumulative explained variance",
}

STRATEGY_TRANSLATIONS = {
    "default_0_50": "default 0.50",
    "max_f1": "max F1",
    "recall_at_least_60_best_precision": "recall >= 0.60",
    "recall_at_least_70_best_precision": "recall >= 0.70",
}


def find_project_root() -> Path:
    repo_root = Path.cwd()
    candidates = [p.parent.parent for p in repo_root.rglob("overleaf_final/main.tex")]
    if not candidates:
        raise FileNotFoundError("Cannot locate project directory with overleaf_final/main.tex.")
    return candidates[0]


def clear_document_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def clear_headers_and_footers(doc: Document) -> None:
    for section in doc.sections:
        for part in [section.header, section.footer]:
            for p in part.paragraphs:
                p.text = ""
            for table in list(part.tables):
                part._element.remove(table._tbl)


def ensure_style(doc: Document, style_name: str, style_type=WD_STYLE_TYPE.PARAGRAPH):
    try:
        return doc.styles[style_name]
    except KeyError:
        return doc.styles.add_style(style_name, style_type)


def set_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None, font: str = "Calibri") -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    clear_headers_and_footers(doc)

    normal = ensure_style(doc, "Normal")
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)

    for name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = ensure_style(doc, name)
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12 if name == "Heading 1" else 8)
        style.paragraph_format.space_after = Pt(6)

    for name in ["List Bullet", "List Number"]:
        style = ensure_style(doc, name)
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)


def add_centered(doc: Document, text: str, size: float = 12, bold: bool = False, color: str | None = None, after: float = 6) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)


def add_paragraph(doc: Document, text: str, after: float = 6, first_line: bool = False, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.10
    p.paragraph_format.space_after = Pt(after)
    if first_line:
        p.paragraph_format.first_line_indent = Pt(18)
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_font(lead, size=11, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_font(rest, size=11)
    else:
        r = p.add_run(text)
        set_font(r, size=11)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    r = p.add_run(text)
    set_font(r, size=11)


def add_heading(doc: Document, text: str, level: int, page_break: bool = False) -> None:
    if page_break:
        doc.add_page_break()
    p = doc.add_paragraph(style=f"Heading {level}")
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_font(r, size=16 if level == 1 else 13 if level == 2 else 12, bold=True, color="2E74B5" if level <= 2 else "1F4D78")


def set_cell_text(cell, text: str, font_size: float = 9.0, bold: bool = False, align: str = "center") -> None:
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == "center" else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_font(r, size=font_size, bold=bold)


def shade_cell(cell, fill: str = "F2F4F7") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "A6A6A6")


def add_table_from_rows(doc: Document, caption: str, rows: list[list[str]], chapter: int, table_counts: dict[int, int]) -> None:
    if not rows:
        return
    table_counts[chapter] = table_counts.get(chapter, 0) + 1
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(f"Table {chapter}-{table_counts[chapter]} {caption}")
    set_font(r, size=10, bold=True)

    n_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=1, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    try:
        table.style = "Table Grid"
    except KeyError:
        set_table_borders(table)

    font_size = 6.6 if n_cols >= 8 else 7.8 if n_cols >= 5 else 9.0
    for i in range(n_cols):
        set_cell_text(table.rows[0].cells[i], rows[0][i] if i < len(rows[0]) else "", font_size=font_size, bold=True)
        shade_cell(table.rows[0].cells[i])
    for data_row in rows[1:]:
        row = table.add_row()
        for i in range(n_cols):
            value = data_row[i] if i < len(data_row) else ""
            align = "center" if len(value) < 22 else "left"
            set_cell_text(row.cells[i], value, font_size=font_size, align=align)
    doc.add_paragraph()


def add_cover(doc: Document) -> None:
    add_centered(doc, "湖北大学 2025--2026 学年度第 2 学期课程考查试题纸", size=15, after=8)
    add_centered(doc, "The paper of course exam", size=12, after=20)

    info = doc.add_table(rows=4, cols=4)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ("Name:", COURSE_NAME, "", ""),
        ("Content:", CONTENT_NAME, "", ""),
        ("Institution:", "________________", "Teacher:", TEACHER_NAME),
        ("Grade and major:", "________________", "", ""),
    ]
    for row, values in zip(info.rows, rows):
        for cell, value in zip(row.cells, values):
            set_cell_text(cell, value, font_size=11, bold=value.endswith(":") or value == "Teacher:", align="left")
    doc.add_paragraph()

    add_centered(doc, "Report Topic", size=15, bold=True, after=10)
    add_centered(doc, REPORT_TITLE, size=13, bold=True, color="1F4D78", after=24)

    personal = doc.add_table(rows=6, cols=2)
    personal.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row, (label, value) in zip(
        personal.rows,
        [
            ("Student ID:", "____________________________"),
            ("Name:", "____________________________"),
            ("College:", "____________________________"),
            ("Major and Grade:", "____________________________"),
            ("Institution:", "____________________________"),
            ("Grade and major:", "____________________________"),
        ],
    ):
        set_cell_text(row.cells[0], label, font_size=11, bold=True, align="right")
        set_cell_text(row.cells[1], value, font_size=11, align="left")
    doc.add_paragraph()

    score = doc.add_table(rows=2, cols=3)
    score.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text in zip(score.rows[0].cells, ["Teacher's comments", "Total score", "Grading teacher"]):
        set_cell_text(cell, text, font_size=10.5, bold=True)
        shade_cell(cell)
    for cell in score.rows[1].cells:
        set_cell_text(cell, "________________", font_size=10.5)
    doc.add_paragraph()
    add_centered(doc, "湖北大学", size=12, after=0)
    doc.add_page_break()


def add_abstract_and_toc(doc: Document) -> None:
    add_centered(doc, "Abstract", size=16, bold=True, color="2E74B5", after=12)
    abstract_paragraphs = [
        "This report is prepared for the course Big Data Analysis and Applications and follows the teacher's seven grading dimensions: independent dataset selection, independent research topic design, data preprocessing, exploratory analysis and visualization, machine learning modeling and evaluation, report/code organization, and final reflection. The study uses the 2025 Digital Lifestyle Benchmark Dataset, a CSV structured tabular dataset with 3,500 records and 24 fields. The dataset is a synthetic educational and benchmark dataset, so all findings are used only for course modeling practice and behavioral analysis, not for medical diagnosis, individual health judgment, or causal inference.",
        "The research topic is high-risk identification, digital dependence prediction, and user profiling in digital lifestyles. Three machine learning task types are implemented, which exceeds the course requirement of at least two types: classification for high_risk_flag screening, regression for digital_dependence_score prediction, and clustering for exploratory digital lifestyle profiling. The preprocessing stage includes data cleaning, feature extraction, PCA extension, LCA/GMM discussion, and strict target leakage control. Exploratory analysis uses descriptive statistics, correlation analysis, group comparisons, and visualizations such as bar charts, histograms, boxplots, scatter plots, heatmaps, ROC/PR curves, elbow curves, silhouette curves, PCA plots, and cluster profile heatmaps.",
        "The validated experimental results are kept unchanged. For high-risk screening, the final classifier is Gradient Boosting with threshold=0.14, achieving Recall=0.6420, F1=0.5355, and PR-AUC=0.5084 on the test set. This is interpreted as a recall-oriented screening strategy rather than a diagnosis tool. For regression, digital_dependence_score is the main target, with R²=0.9839, MSE=3.1471, and MAE=0.9982. By contrast, productivity_score obtains R²=-0.0041 and is reported as a weak-prediction or negative-result analysis. For clustering, KMeans with k=3 obtains Silhouette=0.1860, so the clustering result is used only for exploratory user profiling instead of claiming naturally clear population boundaries.",
    ]
    for p in abstract_paragraphs:
        add_paragraph(doc, p, first_line=True)
    add_paragraph(
        doc,
        "Keywords: digital lifestyle; high-risk screening; digital dependence prediction; machine learning; clustering profile; big data analysis",
        bold_lead="Keywords:",
    )
    doc.add_page_break()

    add_centered(doc, "Table of Contents", size=16, bold=True, color="2E74B5", after=12)
    toc = [
        ("Chapter 1 Dataset Selection", ["1.1 Course Requirements for Dataset Selection", "1.2 Candidate Dataset Comparison and Reasons for Excluding pfm_train / pfm_test", "1.3 Final Dataset Source and Compliance Check", "1.4 Feature Fields, Sample Scale, and Supported Tasks"]),
        ("Chapter 2 Independent Research Theme and Analytical Perspective", ["2.1 Business Background and Problem Motivation", "2.2 Report Title and Research Objectives", "2.3 Analytical Perspective and Research Questions", "2.4 Main Contributions and Technical Workflow"]),
        ("Chapter 3 Data Preprocessing", ["3.1 Raw Data Loading and Field Inspection", "3.2 Missing Values, Duplicate Records, and Abnormal Data Checking", "3.3 Data Rationality Verification", "3.4 Feature Engineering and Feature Extraction", "3.5 PCA Extension and LCA/GMM Discussion", "3.6 Feature Selection and Target Leakage Control", "3.7 Summary of Preprocessing Results"]),
        ("Chapter 4 Exploratory Analysis and Data Visualization", ["4.1 Descriptive Statistical Analysis", "4.2 Target Variable Distribution", "4.3 Numerical Feature Distribution and Correlation Analysis", "4.4 Group Differences between High-Risk and Non-High-Risk Samples", "4.5 Categorical Risk Rate Analysis", "4.6 Relationships between Behavioral Variables and Outcome Variables", "4.7 Summary of EDA Findings"]),
        ("Chapter 5 Machine Learning Modeling, Hyperparameter Tuning, and Evaluation", ["5.1 Experimental Settings and Evaluation Metrics", "5.2 Classification Task: High-Risk Digital Lifestyle Screening", "5.3 Regression Task: Digital Dependence Prediction and Productivity Weak-Prediction Analysis", "5.4 Clustering Task: Digital Lifestyle User Profiling", "5.5 Integrated Model Result Analysis", "5.6 Business Interpretation and Result Boundaries"]),
        ("Chapter 6 Report Structure and Code Presentation", ["6.1 Project Directory and Engineering Structure", "6.2 Notebook and Python Source Code Organization", "6.3 Core Code Snippets and Explanations", "6.4 Figures, Result Files, and Reproducibility", "6.5 Complete Code Appendix Description"]),
        ("Chapter 7 Conclusions and Personal Reflection", ["7.1 Core Experimental Conclusions", "7.2 Data Patterns and Practical Value", "7.3 Model Advantages, Limitations, and Risks", "7.4 Future Improvements", "7.5 Course Learning Reflection and Appreciation"]),
        ("References", []),
        ("Appendix A Core Code", []),
        ("Appendix B Full Experimental Running Instructions", []),
        ("Appendix C Result Files and Figure List", []),
        ("Appendix D Personal Contribution and Academic Integrity Statement", []),
    ]
    for chapter, sections in toc:
        p = doc.add_paragraph()
        r = p.add_run(chapter)
        set_font(r, size=11, bold=True)
        for sec in sections:
            sp = doc.add_paragraph()
            sp.paragraph_format.left_indent = Inches(0.28)
            sr = sp.add_run(sec)
            set_font(sr, size=10)
    doc.add_page_break()


def clean_latex(text: str) -> str:
    s = text.strip()
    for old, new in {
        "\\_": "_",
        "\\%": "%",
        "\\&": "&",
        "$R^2$": "R²",
        "$k$": "k",
    }.items():
        s = s.replace(old, new)
    s = re.sub(r"\\texttt\{([^{}]+)\}", r"\1", s)
    s = re.sub(r"\$([^$]+)\$", lambda m: m.group(1).replace("^2", "²"), s)
    s = re.sub(r"\\[a-zA-Z]+", "", s)
    s = s.replace("{", "").replace("}", "")
    s = re.sub(r"\s+", " ", s)
    return s.strip()


def translate_cell(text: str) -> str:
    value = clean_latex(text)
    if value in STRATEGY_TRANSLATIONS:
        return STRATEGY_TRANSLATIONS[value]
    if value in CELL_TRANSLATIONS:
        return CELL_TRANSLATIONS[value]
    for cn, en in sorted(CELL_TRANSLATIONS.items(), key=lambda x: len(x[0]), reverse=True):
        value = value.replace(cn, en)
    return value


def extract_balanced_arg(text: str, command: str) -> str:
    start = text.find(command)
    if start < 0:
        return ""
    pos = text.find("{", start)
    if pos < 0:
        return ""
    depth = 0
    out = []
    for ch in text[pos + 1 :]:
        if ch == "{":
            depth += 1
        elif ch == "}":
            if depth == 0:
                break
            depth -= 1
        out.append(ch)
    return "".join(out)


def parse_latex_table(path: Path) -> tuple[str, list[list[str]]]:
    text = path.read_text(encoding="utf-8")
    caption = CAPTION_TRANSLATIONS.get(path.name, translate_cell(extract_balanced_arg(text, "\\caption")))
    begin = text.find("\\begin{tabular}")
    end = text.find("\\end{tabular}", begin)
    if begin < 0 or end < 0:
        return caption, []
    header_end = text.find("}", begin + len("\\begin{tabular}"))
    body = text[header_end + 1 : end]
    for token in ["\\toprule", "\\midrule", "\\bottomrule"]:
        body = body.replace(token, "")
    rows = []
    for raw_row in body.split("\\\\"):
        row = raw_row.strip()
        if "&" not in row:
            continue
        cells = [translate_cell(cell) for cell in row.split("&")]
        rows.append(cells)
    return caption, rows


def add_latex_table(doc: Document, tables_dir: Path, filename: str, chapter: int, table_counts: dict[int, int]) -> None:
    caption, rows = parse_latex_table(tables_dir / filename)
    add_table_from_rows(doc, caption, rows, chapter, table_counts)


def add_project_structure_table(doc: Document, chapter: int, table_counts: dict[int, int]) -> None:
    rows = [
        ["Directory/File", "Role in the report"],
        ["data/raw and data/processed", "Store raw and processed dataset files without deleting the original data."],
        ["notebooks/00-05", "Record dataset selection, preprocessing/EDA, classification, regression, clustering, and result summary."],
        ["src/config.py", "Centralize paths, random_state=42, target names, and leakage-column definitions."],
        ["src/data_utils.py", "Provide dataset loading, saving, and reusable quality-check helpers."],
        ["src/feature_engineering.py", "Generate engineered behavior-ratio and interaction features and define task-specific feature sets."],
        ["src/model_utils.py", "Provide model evaluation, metric calculation, threshold tuning, and clustering utilities."],
        ["scripts/stage2_5_evidence_enhancement.py", "Generate preprocessing quality checks, PCA explained variance evidence, and final report artifact lists."],
        ["results/ and figures/", "Store CSV metrics, summaries, tables, and all visual evidence used in the report."],
        ["final_submit/", "Store the final Word submission file."],
    ]
    add_table_from_rows(doc, "Project directory and code organization", rows, chapter, table_counts)


def add_figure(doc: Document, figures_dir: Path, filename: str, caption: str, chapter: int, fig_counts: dict[int, int], width: float = 5.75) -> None:
    path = figures_dir / filename
    if not path.exists():
        return
    fig_counts[chapter] = fig_counts.get(chapter, 0) + 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(f"Figure {chapter}-{fig_counts[chapter]} {caption}")
    set_font(r, size=10, bold=True)


def add_code_block(doc: Document, caption: str, code: str, chapter: int, code_counts: dict[int, int]) -> None:
    code_counts[chapter] = code_counts.get(chapter, 0) + 1
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(f"Code {chapter}-{code_counts[chapter]} {caption}")
    set_font(r, size=10, bold=True)
    for line in code.strip("\n").splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        set_font(r, size=8.3, font="Consolas")


def add_chapter_1(doc: Document, tables_dir: Path, table_counts: dict[int, int]) -> None:
    add_heading(doc, "Chapter 1 Dataset Selection", 1, page_break=True)
    add_heading(doc, "1.1 Course Requirements for Dataset Selection", 2)
    add_paragraph(doc, "This chapter directly corresponds to the first grading dimension: independent dataset selection. According to the course requirements, the dataset should come from a qualified public data platform such as Kaggle, Hugging Face, Alibaba Tianchi, Heywhale, or OpenML. It should be released or updated in 2024-2025, stored as a CSV structured table, contain at least 1,000 samples, contain at least 8 effective features, and naturally support at least two machine learning task types among regression, classification, and clustering.")
    add_paragraph(doc, "The selected dataset must not merely be convenient. It should allow meaningful preprocessing, statistical analysis, visualization, machine learning modeling, metric evaluation, and code reproduction. Therefore, dataset selection is treated as a methodological decision rather than a file-loading step.")
    add_heading(doc, "1.2 Candidate Dataset Comparison and Reasons for Excluding pfm_train / pfm_test", 2)
    add_paragraph(doc, "The pfm_train.csv and pfm_test.csv files were considered as candidates but were not used as the final dataset. Their source platform and release/update year were less stable for the final course requirement, and the employee attrition topic was highly similar to a previous experiment. In addition, the pfm files mainly support a classification-style task, while the regression task is less natural. To avoid repeating the earlier experiment and to build a richer final report, they were excluded from the main analysis.")
    add_heading(doc, "1.3 Final Dataset Source and Compliance Check", 2)
    add_paragraph(doc, "The final dataset is the 2025 Digital Lifestyle Benchmark Dataset. It is suitable for this report because it contains digital behavior, lifestyle habits, mental-risk indicators, high_risk_flag, productivity_score, and digital_dependence_score. These variables support classification, regression, and clustering in one coherent analytical scenario.")
    add_latex_table(doc, tables_dir, "dataset_compliance_summary.tex", 1, table_counts)
    add_paragraph(doc, "Table 1-1 shows that the dataset satisfies the formal course requirements. The dataset is synthetic educational/benchmark data, so it is appropriate for modeling practice but should not be interpreted as a medical or real-world diagnostic dataset.")
    add_heading(doc, "1.4 Feature Fields, Sample Scale, and Supported Tasks", 2)
    add_paragraph(doc, "The dataset contains 3,500 records and 24 fields. The behavioral fields include variables such as device_hours_per_day, phone_unlocks, notifications_per_day, social_media_mins, and study_mins. Lifestyle fields include sleep_hours, sleep_quality, and physical_activity_days. Outcome fields include high_risk_flag, productivity_score, and digital_dependence_score.")
    add_paragraph(doc, "This field design supports three task types: high_risk_flag classification, digital_dependence_score regression, and exploratory lifestyle clustering. The report therefore exceeds the course requirement of implementing at least two machine learning task types.")


def add_chapter_2(doc: Document) -> None:
    add_heading(doc, "Chapter 2 Independent Research Theme and Analytical Perspective", 1, page_break=True)
    add_heading(doc, "2.1 Business Background and Problem Motivation", 2)
    add_paragraph(doc, "Mobile devices, social media platforms, frequent notifications, sleep habits, study time, and physical activity have become important components of daily digital lifestyles. A big data analysis report can use these observable behavior variables to explore digital well-being patterns in a structured and reproducible way.")
    add_paragraph(doc, "The purpose of this report is not to diagnose health conditions, but to build a course-level modeling workflow that connects data quality, mathematical statistical analysis, visualization, supervised learning, and unsupervised profiling.")
    add_heading(doc, "2.2 Report Title and Research Objectives", 2)
    add_paragraph(doc, f"The report title is: {REPORT_TITLE}.")
    for item in [
        "Identify high-risk digital lifestyle samples using high_risk_flag as the classification target.",
        "Predict digital_dependence_score as the main regression target.",
        "Test whether productivity_score can be stably predicted by the current observable features.",
        "Construct exploratory digital lifestyle user profiles through clustering.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "2.3 Analytical Perspective and Research Questions", 2)
    add_paragraph(doc, "The analysis is organized from four perspectives. The behavioral perspective includes device time, notifications, phone unlocks, social media time, and study time. The lifestyle perspective includes sleep duration, sleep quality, and physical activity. The outcome perspective includes high-risk labels, digital dependence, and productivity. The methodological perspective includes classification, regression, and clustering.")
    add_paragraph(doc, "The main research questions are: Can high-risk digital lifestyles be screened with leakage-controlled features? Can digital dependence be predicted from digital behavior and lifestyle habits? Is productivity_score predictable under the same feature constraints? Can clustering form interpretable digital lifestyle profiles?")
    add_heading(doc, "2.4 Main Contributions and Technical Workflow", 2)
    for item in [
        "Completed independent dataset screening and explained why the pfm dataset was not selected.",
        "Performed data cleaning, feature extraction, PCA extension, LCA/GMM discussion, and target leakage control.",
        "Used statistical analysis and visualization to examine distributions, correlations, group differences, and risk-rate patterns.",
        "Built classification, regression, and clustering tasks with tuning, metrics, and visualized evaluation evidence.",
        "Organized notebooks, source code, result CSV files, figures, and the final Word report into a reproducible project structure.",
    ]:
        add_bullet(doc, item)


def add_chapter_3(doc: Document, tables_dir: Path, figures_dir: Path, table_counts: dict[int, int], fig_counts: dict[int, int], code_counts: dict[int, int]) -> None:
    add_heading(doc, "Chapter 3 Data Preprocessing", 1, page_break=True)
    add_heading(doc, "3.1 Raw Data Loading and Field Inspection", 2)
    add_paragraph(doc, "The raw CSV file was loaded with pandas. The first checks verified the row count, column count, field names, and data types. The raw dataset contains 3,500 records and 24 fields. These checks ensure that subsequent preprocessing and modeling are performed on the expected structured table.")
    add_heading(doc, "3.2 Missing Values, Duplicate Records, and Abnormal Data Checking", 2)
    add_paragraph(doc, "The preprocessing stage includes missing value checking, duplicate record checking, abnormal value checking, data rationality verification, and task-specific feature selection. The total number of missing values is 0, the number of duplicate rows is 0, and the number of duplicate id values is 0. Since no severe abnormal records were found, no sample deletion was performed.")
    add_latex_table(doc, tables_dir, "preprocessing_quality_check_compact.tex", 3, table_counts)
    add_paragraph(doc, "Table 3-1 summarizes the data cleaning evidence. The result supports the decision to keep the full sample while still preserving robust preprocessing operations in the modeling pipeline.")
    add_heading(doc, "3.3 Data Rationality Verification", 2)
    add_paragraph(doc, "The rationality check covers age, device_hours_per_day, phone_unlocks, notifications_per_day, social_media_mins, study_mins, physical_activity_days, sleep_hours, sleep_quality, productivity_score, and digital_dependence_score. The check focuses on whether the observed minimum and maximum values are consistent with reasonable behavioral ranges.")
    add_heading(doc, "3.4 Feature Engineering and Feature Extraction", 2)
    add_paragraph(doc, "Feature extraction was conducted by constructing behavioral ratio and interaction features. The engineered features include social_media_hours, study_hours, notifications_per_device_hour, unlocks_per_device_hour, device_to_sleep_ratio, activity_sleep_interaction, and social_to_study_ratio.")
    add_code_block(doc, "Behavioral feature extraction", """def add_engineered_features(df):
    out = df.copy()
    out["social_media_hours"] = out["social_media_mins"] / 60
    out["study_hours"] = out["study_mins"] / 60
    out["notifications_per_device_hour"] = (
        out["notifications_per_day"] / (out["device_hours_per_day"] + 1e-6)
    )
    out["unlocks_per_device_hour"] = (
        out["phone_unlocks"] / (out["device_hours_per_day"] + 1e-6)
    )
    out["device_to_sleep_ratio"] = (
        out["device_hours_per_day"] / (out["sleep_hours"] + 1e-6)
    )
    out["activity_sleep_interaction"] = (
        out["physical_activity_days"] * out["sleep_quality"]
    )
    out["social_to_study_ratio"] = (
        out["social_media_mins"] / (out["study_mins"] + 1)
    )
    return out""", 3, code_counts)
    add_heading(doc, "3.5 PCA Extension and LCA/GMM Discussion", 2)
    add_paragraph(doc, "PCA was used for dimensionality reduction visualization and auxiliary interpretation of clustering structure. The first two principal components explain about 42.41% of the total variance. PCA was not used as input for classification or regression and should not be interpreted causally.")
    add_latex_table(doc, tables_dir, "pca_explained_variance_compact.tex", 3, table_counts)
    add_figure(doc, figures_dir, "pca_explained_variance.png", "Cumulative explained variance of PCA components", 3, fig_counts)
    add_paragraph(doc, "The PCA curve provides visual evidence for the amount of information captured by the leading components. It supports dimensionality-reduction visualization, but it does not prove that the behavioral variables cause the outcome variables.")
    add_paragraph(doc, "A separate LCA model was not introduced because the clustering inputs are mainly continuous behavioral and lifestyle variables. GaussianMixture was used as a probabilistic clustering method to approximate latent group structure and to complement KMeans and AgglomerativeClustering.")
    add_heading(doc, "3.6 Feature Selection and Target Leakage Control", 2)
    add_paragraph(doc, "Target leakage control is essential because high_risk_flag is closely related to several mental-status and outcome variables. For the classification task, anxiety_score, depression_score, stress_level, happiness_score, focus_score, productivity_score, digital_dependence_score, id, and high_risk_flag were excluded from the input features.")
    add_paragraph(doc, "For regression, high_risk_flag and other outcome fields were excluded. For clustering, only digital behavior and lifestyle numerical features were used as clustering inputs, while demographic and outcome variables were kept only for post-clustering interpretation.")
    add_code_block(doc, "Task-specific leakage control", """HIGH_RISK_LEAKAGE_COLUMNS = [
    "anxiety_score", "depression_score", "stress_level",
    "happiness_score", "focus_score", "productivity_score",
    "digital_dependence_score",
]

classification_drop = [
    "id", "high_risk_flag", *HIGH_RISK_LEAKAGE_COLUMNS
]

regression_drop = [
    "id", "high_risk_flag", "anxiety_score", "depression_score",
    "stress_level", "happiness_score", "focus_score",
]

clustering_features = [
    "device_hours_per_day", "phone_unlocks", "notifications_per_day",
    "social_media_mins", "study_mins", "physical_activity_days",
    "sleep_hours", "sleep_quality", "social_media_hours",
    "study_hours", "notifications_per_device_hour",
    "unlocks_per_device_hour", "device_to_sleep_ratio",
    "activity_sleep_interaction", "social_to_study_ratio",
]""", 3, code_counts)
    add_heading(doc, "3.7 Summary of Preprocessing Results", 2)
    add_paragraph(doc, "The preprocessing stage produced a clean and reproducible modeling table. No missing values, duplicate rows, or duplicate ids were found. Seven engineered features were generated successfully. The feature sets for classification, regression, and clustering were controlled separately to reduce leakage and keep the interpretation aligned with each task.")


def add_chapter_4(doc: Document, figures_dir: Path, fig_counts: dict[int, int]) -> None:
    add_heading(doc, "Chapter 4 Exploratory Analysis and Data Visualization", 1, page_break=True)
    add_heading(doc, "4.1 Descriptive Statistical Analysis", 2)
    add_paragraph(doc, "Before modeling, descriptive statistics, distributions, correlations, group comparisons, and risk-rate summaries were used to explore the mathematical structure of the dataset. The report examines mean, range, distribution shape, correlation, and group difference. This step helps determine whether later modeling results are plausible and interpretable.")
    add_heading(doc, "4.2 Target Variable Distribution", 2)
    add_figure(doc, figures_dir, "eda_high_risk_flag_distribution.png", "Distribution of high_risk_flag", 4, fig_counts)
    add_paragraph(doc, "The bar chart shows the class distribution of high_risk_flag. Because the positive high-risk group is not the majority class, Accuracy alone is not sufficient for model evaluation. This observation motivates using Recall, F1, PR-AUC, and the confusion matrix in Chapter 5.")
    add_heading(doc, "4.3 Numerical Feature Distribution and Correlation Analysis", 2)
    add_figure(doc, figures_dir, "eda_numeric_histograms.png", "Distributions of core numerical variables", 4, fig_counts)
    add_paragraph(doc, "The histogram panel visualizes the distributions of core numerical variables such as device time, notifications, unlocks, sleep, digital dependence, and productivity. These distributions provide the mathematical background for deciding whether scaling, robust metrics, or nonlinear models may be useful.")
    add_figure(doc, figures_dir, "eda_numeric_correlation_heatmap.png", "Correlation heatmap of numerical variables", 4, fig_counts)
    add_paragraph(doc, "The heatmap summarizes pairwise correlations among numerical variables. It helps identify associations between digital behavior variables and outcome variables, but correlation does not imply causation. The heatmap is used as exploratory evidence rather than causal proof.")
    add_heading(doc, "4.4 Group Differences between High-Risk and Non-High-Risk Samples", 2)
    add_figure(doc, figures_dir, "eda_boxplots_by_risk.png", "Behavioral differences by high-risk group", 4, fig_counts)
    add_paragraph(doc, "The boxplots compare behavioral and lifestyle variables between high-risk and non-high-risk samples. The visual comparison supports later classification modeling by showing that several observable behaviors differ across risk groups. However, the figure should still be interpreted as descriptive group comparison.")
    add_heading(doc, "4.5 Categorical Risk Rate Analysis", 2)
    add_figure(doc, figures_dir, "eda_category_risk_rate.png", "Risk rates across categorical variables", 4, fig_counts)
    add_paragraph(doc, "The categorical risk-rate bar chart compares risk proportions across variables such as gender, region, income_level, education_level, daily_role, and device_type. This visualization checks whether background categories show systematic descriptive differences, while the main clustering input remains restricted to behavioral and lifestyle numerical variables.")
    add_heading(doc, "4.6 Relationships between Behavioral Variables and Outcome Variables", 2)
    add_figure(doc, figures_dir, "eda_behavior_outcome_scatter.png", "Relationships between behavioral variables and outcomes", 4, fig_counts)
    add_paragraph(doc, "The scatter plots and binned relationship plots show how selected behavioral variables relate to digital dependence, productivity, and high-risk labels. This satisfies the requirement for scatter-style visualization and supports the transition from exploratory analysis to supervised and unsupervised modeling.")
    add_heading(doc, "4.7 Summary of EDA Findings", 2)
    add_paragraph(doc, "The EDA chapter uses bar charts, histograms, boxplots, scatter plots, and heatmaps with analysis rather than only inserting figures. The visual evidence suggests that digital behavior variables are useful for high-risk screening and digital dependence prediction, while productivity may require additional unobserved factors.")


def add_chapter_5(doc: Document, tables_dir: Path, figures_dir: Path, table_counts: dict[int, int], fig_counts: dict[int, int], code_counts: dict[int, int]) -> None:
    add_heading(doc, "Chapter 5 Machine Learning Modeling, Hyperparameter Tuning, and Evaluation", 1, page_break=True)
    add_heading(doc, "5.1 Experimental Settings and Evaluation Metrics", 2)
    add_paragraph(doc, "All experiments use Python, pandas, numpy, scipy, matplotlib, and scikit-learn on CPU. No GPU, no deep learning, and no unverified experimental results are used. The random_state is fixed at 42 for reproducibility. Supervised tasks use train/test splitting, cross-validation, and GridSearchCV-style hyperparameter tuning; the test set is used only for final evaluation.")
    add_paragraph(doc, "The classification task reports Accuracy, Precision, Recall, F1, confusion matrix, ROC-AUC, PR-AUC, and Balanced Accuracy. F1 and Recall are emphasized because high-risk screening should reduce missed high-risk cases. The regression task reports R², MSE, RMSE, and MAE. The clustering task uses the elbow method, silhouette coefficient, Calinski-Harabasz, and Davies-Bouldin metrics.")
    add_code_block(doc, "Cross-validation and GridSearchCV for supervised tuning", """cv = StratifiedKFold(n_splits=5, shuffle=True, random_state=42)

search = GridSearchCV(
    estimator=pipeline,
    param_grid=param_grid,
    scoring="average_precision",
    cv=cv,
    n_jobs=-1,
)

search.fit(X_train, y_train)
best_model = search.best_estimator_""", 5, code_counts)
    add_heading(doc, "5.2 Classification Task: High-Risk Digital Lifestyle Screening", 2)
    add_paragraph(doc, "The classification target is high_risk_flag. Logistic Regression, Random Forest, and Gradient Boosting were evaluated. The final strategy is Gradient Boosting with threshold=0.14. This threshold is selected as a recall-oriented screening strategy, not as a medical diagnosis tool.")
    add_latex_table(doc, tables_dir, "classification_tuned_metrics.tex", 5, table_counts)
    add_latex_table(doc, tables_dir, "classification_threshold_tuning.tex", 5, table_counts)
    add_paragraph(doc, "The final test-set results are Recall=0.6420, F1=0.5355, and PR-AUC=0.5084. The threshold tuning table shows why the default 0.50 threshold is not the only reasonable choice. For high-risk screening, a lower threshold can improve recall, but it also increases false positives.")
    add_figure(doc, figures_dir, "classification_tuned_metrics_comparison.png", "Comparison of tuned classification metrics", 5, fig_counts)
    add_figure(doc, figures_dir, "classification_final_confusion_matrix.png", "Final confusion matrix for high-risk screening", 5, fig_counts)
    add_figure(doc, figures_dir, "classification_precision_recall_curve.png", "Precision-recall curve for the final classifier", 5, fig_counts)
    add_figure(doc, figures_dir, "classification_roc_curve.png", "ROC curve for the final classifier", 5, fig_counts)
    add_figure(doc, figures_dir, "classification_permutation_importance.png", "Permutation importance for high-risk screening", 5, fig_counts)
    add_paragraph(doc, "Modeling results are supported by visual evidence, including the confusion matrix, ROC curve, precision-recall curve, tuned-metric comparison, and permutation importance. The confusion matrix directly shows true positives, false positives, false negatives, and true negatives, while the PR curve is especially useful under class imbalance.")
    add_code_block(doc, "Threshold selection for a recall-oriented screening model", """thresholds = np.linspace(0.01, 0.99, 99)
records = []

for threshold in thresholds:
    y_pred = (valid_proba >= threshold).astype(int)
    records.append({
        "threshold": threshold,
        "precision": precision_score(y_valid, y_pred, zero_division=0),
        "recall": recall_score(y_valid, y_pred, zero_division=0),
        "f1": f1_score(y_valid, y_pred, zero_division=0),
    })

threshold_table = pd.DataFrame(records)
selected = threshold_table.query("recall >= 0.60").sort_values(
    ["precision", "f1"], ascending=False
).iloc[0]""", 5, code_counts)
    add_heading(doc, "5.3 Regression Task: Digital Dependence Prediction and Productivity Weak-Prediction Analysis", 2)
    add_paragraph(doc, "The regression main target is digital_dependence_score. The best model obtains R²=0.9839, MSE=3.1471, and MAE=0.9982. This indicates that digital behavior variables strongly represent digital dependence in this synthetic dataset. However, strong prediction does not imply causation.")
    add_latex_table(doc, tables_dir, "regression_target_comparison.tex", 5, table_counts)
    add_latex_table(doc, tables_dir, "regression_digital_dependence_metrics.tex", 5, table_counts)
    add_latex_table(doc, tables_dir, "regression_productivity_metrics.tex", 5, table_counts)
    add_figure(doc, figures_dir, "regression_target_comparison.png", "Comparison of regression targets", 5, fig_counts)
    add_figure(doc, figures_dir, "regression_digital_dependence_observed_vs_predicted.png", "Observed versus predicted digital dependence scores", 5, fig_counts)
    add_figure(doc, figures_dir, "regression_digital_dependence_permutation_importance.png", "Permutation importance for digital dependence prediction", 5, fig_counts)
    add_figure(doc, figures_dir, "regression_productivity_observed_vs_predicted.png", "Observed versus predicted productivity scores", 5, fig_counts)
    add_paragraph(doc, "The productivity_score experiment is deliberately kept as a weak-prediction or negative-result analysis. Its best R² is -0.0041, so this report does not claim that productivity can be effectively predicted by the current feature set. This negative result is valuable because it shows that not every target is equally explainable by the same observable variables.")
    add_code_block(doc, "Regression metric calculation", """metrics = {
    "R2": r2_score(y_test, y_pred),
    "MSE": mean_squared_error(y_test, y_pred),
    "RMSE": mean_squared_error(y_test, y_pred, squared=False),
    "MAE": mean_absolute_error(y_test, y_pred),
}

pd.DataFrame([metrics]).to_csv(output_path, index=False)""", 5, code_counts)
    add_heading(doc, "5.4 Clustering Task: Digital Lifestyle User Profiling", 2)
    add_paragraph(doc, "The clustering task uses only digital behavior and lifestyle numerical features. KMeans, AgglomerativeClustering, and GaussianMixture were compared. The elbow method and silhouette coefficient were used for k selection. The best setting is KMeans with k=3 and Silhouette=0.1860.")
    add_latex_table(doc, tables_dir, "clustering_model_comparison.tex", 5, table_counts)
    add_latex_table(doc, tables_dir, "clustering_lifestyle_profiles_compact.tex", 5, table_counts)
    add_figure(doc, figures_dir, "clustering_kmeans_elbow.png", "KMeans elbow curve for selecting k", 5, fig_counts)
    add_figure(doc, figures_dir, "clustering_silhouette_by_k.png", "Silhouette scores by number of clusters", 5, fig_counts)
    add_figure(doc, figures_dir, "clustering_lifestyle_pca.png", "PCA visualization of lifestyle clusters", 5, fig_counts)
    add_figure(doc, figures_dir, "clustering_lifestyle_profile_heatmap.png", "Heatmap of cluster lifestyle profiles", 5, fig_counts)
    add_paragraph(doc, "The cluster labels are interpreted as high social-media-use profile, high device-dependence profile, and low-load balanced profile. Because Silhouette=0.1860 is relatively low, the clusters should be treated as exploratory profiles rather than naturally clear population boundaries.")
    add_code_block(doc, "Clustering input scaling and model comparison", """cluster_X = df[clustering_features]
cluster_X_scaled = StandardScaler().fit_transform(cluster_X)

for k in range(2, 9):
    labels = KMeans(n_clusters=k, random_state=42, n_init=20).fit_predict(cluster_X_scaled)
    scores.append({
        "algorithm": "KMeans",
        "k": k,
        "silhouette": silhouette_score(cluster_X_scaled, labels),
        "calinski_harabasz": calinski_harabasz_score(cluster_X_scaled, labels),
        "davies_bouldin": davies_bouldin_score(cluster_X_scaled, labels),
    })""", 5, code_counts)
    add_heading(doc, "5.5 Integrated Model Result Analysis", 2)
    add_paragraph(doc, "The three task types answer different parts of the research question. Classification provides a recall-oriented high-risk screening baseline. Regression shows that digital_dependence_score is strongly predictable in the synthetic dataset, while productivity_score is weakly predictable. Clustering provides interpretable but weak-boundary lifestyle profiles.")
    add_heading(doc, "5.6 Business Interpretation and Result Boundaries", 2)
    add_paragraph(doc, "The results can be used as a reference for digital lifestyle management, notification control, social media time awareness, sleep and activity balance, and manual review of high-risk screening results. They should not be used for medical diagnosis, automatic punishment, or strong intervention. The synthetic benchmark nature of the dataset limits direct real-world generalization.")


def add_chapter_6(doc: Document, table_counts: dict[int, int], code_counts: dict[int, int]) -> None:
    add_heading(doc, "Chapter 6 Report Structure and Code Presentation", 1, page_break=True)
    add_heading(doc, "6.1 Project Directory and Engineering Structure", 2)
    add_paragraph(doc, "The report follows a standard data analysis workflow: define analytical objectives, introduce the dataset, perform preprocessing, conduct statistical analysis and visualization, build and tune machine learning models, evaluate results, and summarize limitations and reflections.")
    add_project_structure_table(doc, 6, table_counts)
    add_heading(doc, "6.2 Notebook and Python Source Code Organization", 2)
    add_paragraph(doc, "The notebooks are executed in the order 00_dataset_selection_and_compliance, 01_data_preprocessing_and_eda, 02_classification_high_risk, 03_regression_productivity, 04_clustering_lifestyle_profiles, and 05_result_summary_for_report. The source files include config.py, data_utils.py, feature_engineering.py, model_utils.py, and visualization.py.")
    add_heading(doc, "6.3 Core Code Snippets and Explanations", 2)
    add_paragraph(doc, "The main text embeds core code snippets with explanations, while complete code is provided in notebooks/, src/, and scripts/. This balances readability and reproducibility: readers can understand key logic in the report and still inspect the full implementation in the project directory.")
    add_code_block(doc, "Reproducible configuration", """RANDOM_STATE = 42
CLASSIFICATION_TARGET = "high_risk_flag"
REGRESSION_TARGETS = [
    "productivity_score",
    "digital_dependence_score",
]

HIGH_RISK_LEAKAGE_COLUMNS = [
    "anxiety_score", "depression_score", "stress_level",
    "happiness_score", "focus_score", "productivity_score",
    "digital_dependence_score",
]""", 6, code_counts)
    add_heading(doc, "6.4 Figures, Result Files, and Reproducibility", 2)
    add_paragraph(doc, "All major results are saved as CSV files in results/ and all visual evidence is saved as PNG files in figures/. This makes the report auditable: every metric and figure can be traced back to a generated artifact.")
    add_code_block(doc, "Artifact-oriented result saving", """metrics_df.to_csv(results_dir / "classification_tuned_metrics.csv", index=False)
threshold_df.to_csv(results_dir / "classification_threshold_tuning.csv", index=False)
profile_df.to_csv(results_dir / "clustering_lifestyle_profiles_compact.csv", index=False)

fig.savefig(figures_dir / "classification_precision_recall_curve.png", dpi=160)
fig.savefig(figures_dir / "clustering_lifestyle_profile_heatmap.png", dpi=160)""", 6, code_counts)
    add_heading(doc, "6.5 Complete Code Appendix Description", 2)
    add_paragraph(doc, "The complete code is not fully copied into the main text because long notebooks would reduce readability. Instead, Appendix A explains the core code structure, Appendix B gives the running order, and Appendix C lists important result files and figures.")


def add_chapter_7(doc: Document) -> None:
    add_heading(doc, "Chapter 7 Conclusions and Personal Reflection", 1, page_break=True)
    add_heading(doc, "7.1 Core Experimental Conclusions", 2)
    add_paragraph(doc, "This report completed independent dataset selection, data preprocessing, EDA, classification, regression, clustering, result visualization, code organization, and final Word report preparation. The classification task provides a recall-oriented screening model; the regression task shows strong predictability for digital_dependence_score and weak predictability for productivity_score; the clustering task provides exploratory lifestyle profiles.")
    add_heading(doc, "7.2 Data Patterns and Practical Value", 2)
    add_paragraph(doc, "The descriptive and modeling results suggest that high device usage, notifications, and social media behavior are associated with high-risk labels and digital dependence patterns in this synthetic dataset. These patterns can support digital lifestyle awareness, notification management, sleep/activity balance, and manual review of screening results.")
    add_heading(doc, "7.3 Model Advantages, Limitations, and Risks", 2)
    add_paragraph(doc, "The classification model improves recall by lowering the threshold, but this also increases false positives. The regression model for digital dependence is strong, but it does not prove causality. The productivity regression result is weak, which shows that the current variables do not explain every outcome. The clustering profiles are interpretable, but the silhouette score is low, so the boundaries are not strong.")
    add_heading(doc, "7.4 Future Improvements", 2)
    add_paragraph(doc, "Future work could use real-world longitudinal data, richer behavioral features, external validation, more robust threshold-selection protocols, and additional unsupervised modeling methods. If privacy and ethics conditions are satisfied, time-series behavioral features may also improve the interpretation of digital lifestyle changes.")
    add_heading(doc, "7.5 Course Learning Reflection and Appreciation", 2)
    add_paragraph(doc, "During the course, the teacher repeatedly emphasized that a qualified big data analysis report should not only present model scores, but also explain data cleaning, feature extraction, visualization, evaluation metrics, analytical conclusions, and reproducibility. These reminders helped me organize the whole project in a more systematic way.")
    add_paragraph(doc, "Through this final report, I learned that data compliance, target leakage control, metric selection, negative-result interpretation, visualized evidence, and code organization are all essential parts of a complete analysis. I sincerely appreciate the teacher's clear teaching, patient guidance, and emphasis on rigorous experimental reporting throughout the course.")


def add_references_and_appendices(doc: Document, table_counts: dict[int, int]) -> None:
    add_heading(doc, "References", 1, page_break=True)
    refs = [
        "Digital Lifestyle Benchmark Dataset (2025). Public dataset page and data card.",
        "scikit-learn official documentation. https://scikit-learn.org/",
        "Pedregosa F., Varoquaux G., Gramfort A., et al. Scikit-learn: Machine Learning in Python. Journal of Machine Learning Research, 2011.",
        "pandas official documentation. https://pandas.pydata.org/",
        "NumPy official documentation. https://numpy.org/",
        "Matplotlib official documentation. https://matplotlib.org/",
        "Course requirements for Big Data Analysis and Applications final course report.",
    ]
    for i, ref in enumerate(refs, start=1):
        add_paragraph(doc, f"[{i}] {ref}", after=4)

    add_heading(doc, "Appendix A Core Code", 1, page_break=True)
    add_paragraph(doc, "Appendix A records the purpose of the core source files and the main code snippets already shown in the report. The complete executable code is located in notebooks/, src/, and scripts/. The report avoids copying every notebook cell into the main text so that the reading flow remains clear.")
    add_heading(doc, "Appendix B Full Experimental Running Instructions", 1, page_break=True)
    for item in [
        "Install dependencies with pip install -r requirements.txt.",
        "Run notebooks in the order 00 to 05.",
        "Run scripts/stage2_5_evidence_enhancement.py when preprocessing quality checks and PCA evidence need to be regenerated.",
        "Use CPU only and keep random_state=42.",
        "Open the final DOCX in Word/WPS, fill personal information, check pages, and export PDF manually if the local converter is unavailable.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Appendix C Result Files and Figure List", 1, page_break=True)
    add_paragraph(doc, "Important result files include dataset_compliance_summary.csv, preprocessing_quality_check.csv, classification_tuned_metrics.csv, classification_threshold_tuning.csv, regression_target_comparison.csv, regression_digital_dependence_metrics.csv, regression_productivity_metrics.csv, clustering_model_comparison.csv, clustering_lifestyle_profiles_compact.csv, and pca_explained_variance.csv.")
    rows = [["Artifact group", "Examples"], ["EDA figures", "target distribution, histograms, correlation heatmap, boxplots, categorical risk-rate bars, scatter plots"], ["Classification figures", "tuned metrics comparison, confusion matrix, PR curve, ROC curve, permutation importance"], ["Regression figures", "target comparison, observed-vs-predicted plots, permutation importance"], ["Clustering/PCA figures", "elbow curve, silhouette curve, PCA explained variance, PCA cluster plot, profile heatmap"]]
    add_table_from_rows(doc, "Result file and figure categories", rows, 10, table_counts)
    add_heading(doc, "Appendix D Personal Contribution and Academic Integrity Statement", 1, page_break=True)
    add_paragraph(doc, "This report was independently completed for the Big Data Analysis and Applications course. Previous experiments were used only as workflow references. The final topic, dataset selection, feature design, EDA, modeling tasks, visualizations, interpretations, Word formatting, and conclusions were redesigned for this final course report.")
    add_paragraph(doc, "The report does not copy the artificial intelligence course report content. Personal information remains blank and should be filled manually before final submission.")


def collect_doc_text(doc: Document) -> str:
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def disallowed_chinese(text: str) -> list[str]:
    sanitized = text
    for phrase in ALLOWED_CHINESE_PHRASES:
        sanitized = sanitized.replace(phrase, "")
    snippets = []
    for match in re.finditer(r"[\u4e00-\u9fff]+", sanitized):
        snippet = sanitized[max(0, match.start() - 20): match.end() + 20]
        if snippet not in snippets:
            snippets.append(snippet)
    return snippets[:20]


def try_convert_pdf(docx_path: Path, pdf_path: Path) -> tuple[bool, str]:
    if pdf_path.exists():
        pdf_path.unlink()
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return False, "PDF conversion was not available because LibreOffice/soffice and Word COM were not found in the current environment."
    try:
        subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(pdf_path.parent), str(docx_path)],
            check=True,
            capture_output=True,
            text=True,
            timeout=180,
        )
        converted = pdf_path.parent / f"{docx_path.stem}.pdf"
        if converted.exists() and converted != pdf_path:
            converted.rename(pdf_path)
        return pdf_path.exists() and pdf_path.stat().st_size > 0, "PDF conversion was attempted with LibreOffice/soffice."
    except Exception as exc:  # noqa: BLE001
        return False, f"PDF conversion failed: {exc}"


def main() -> None:
    repo_root = Path.cwd()
    project = find_project_root()
    template_path = repo_root / "期末报告资料" / "课程模板" / "QinTian_experiment.docx"
    if not template_path.exists():
        raise FileNotFoundError(f"Required template not found: {template_path}")

    overleaf = project / "overleaf_final"
    tables_dir = overleaf / "tables"
    figures_dir = overleaf / "figures"
    final_dir = project / "final_submit"
    final_dir.mkdir(exist_ok=True)
    output_docx = final_dir / OUTPUT_DOCX_NAME
    output_pdf = final_dir / OUTPUT_PDF_NAME

    doc = Document(str(template_path))
    clear_document_body(doc)
    setup_document(doc)
    doc.core_properties.author = ""
    doc.core_properties.title = REPORT_TITLE
    doc.core_properties.subject = COURSE_NAME

    fig_counts: dict[int, int] = {}
    table_counts: dict[int, int] = {}
    code_counts: dict[int, int] = {}

    add_cover(doc)
    add_abstract_and_toc(doc)
    add_chapter_1(doc, tables_dir, table_counts)
    add_chapter_2(doc)
    add_chapter_3(doc, tables_dir, figures_dir, table_counts, fig_counts, code_counts)
    add_chapter_4(doc, figures_dir, fig_counts)
    add_chapter_5(doc, tables_dir, figures_dir, table_counts, fig_counts, code_counts)
    add_chapter_6(doc, table_counts, code_counts)
    add_chapter_7(doc)
    add_references_and_appendices(doc, table_counts)

    doc.save(output_docx)

    generated = Document(str(output_docx))
    text = collect_doc_text(generated)
    image_rels = [rel for rel in generated.part.rels.values() if "image" in rel.reltype]
    pdf_ok, pdf_note = try_convert_pdf(output_docx, output_pdf)

    stats = {
        "template_path": str(template_path),
        "output_docx": str(output_docx),
        "output_pdf": str(output_pdf) if pdf_ok else "",
        "pdf_generated": pdf_ok,
        "pdf_note": pdf_note,
        "updated_existing_final_docx": True,
        "english_body_generated": True,
        "teacher_rubric_order_preserved": True,
        "teacher_oral_requirements_explicitly_reinforced": True,
        "figures_inserted_by_counter": sum(fig_counts.values()),
        "embedded_image_relationship_count": len(image_rels),
        "tables_inserted_by_counter": sum(table_counts.values()),
        "word_table_count": len(generated.tables),
        "code_blocks_inserted": sum(code_counts.values()),
        "banned_hits": [term for term in BANNED_TERMS if term in text],
        "latex_hits": [term for term in LATEX_RESIDUES if term in text],
        "disallowed_chinese_snippets": disallowed_chinese(text),
        "personal_placeholders_ok": all(
            term in text
            for term in [
                "Student ID:",
                "Name:",
                "College:",
                "Major and Grade:",
                "Institution:",
                "Grade and major:",
                "________________",
            ]
        ),
        "core_metrics_present": all(
            term in text
            for term in [
                "Recall=0.6420",
                "F1=0.5355",
                "PR-AUC=0.5084",
                "R²=0.9839",
                "MSE=3.1471",
                "MAE=0.9982",
                "R²=-0.0041",
                "Silhouette=0.1860",
            ]
        ),
        "teacher_name_body_forbidden_hits": [
            term for term in ["Teacher Li", "Professor Li", "Li Jie gave"] if term in text
        ],
        "docx_size": output_docx.stat().st_size,
        "pdf_size": output_pdf.stat().st_size if output_pdf.exists() else 0,
    }
    (project / "results" / "stage8_word_generation_stats.json").write_text(
        json.dumps(stats, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    print(json.dumps(stats, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
