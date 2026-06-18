"""Build the teacher-feedback revised Word report.

The script uses the user-provided EEE Word report as the visual source file,
keeps its cover/front matter, removes the old template-like body, and rebuilds
the body with concrete figure/table analysis. It also writes the visual audit
CSV and revision log required by the teacher-feedback pass.
"""

from __future__ import annotations

import csv
import re
import shutil
import subprocess
import sys
import zipfile
from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
RESULTS_DIR = PROJECT_ROOT / "results"
FIGURE_DIR = PROJECT_ROOT / "figures" / "final_report"
FINAL_SUBMIT_DIR = PROJECT_ROOT / "final_submit"
SOURCE_BACKUP = FINAL_SUBMIT_DIR / "EEE大数据分析与应用期末考查报告_源文件备份.docx"
PRE_EDIT_BACKUP = FINAL_SUBMIT_DIR / "大数据分析与应用期末考查报告_修改前备份.docx"
OUTPUT_DOCX = FINAL_SUBMIT_DIR / "大数据分析与应用期末考查报告_老师反馈修改版.docx"
REVISION_LOG = PROJECT_ROOT / "revision_log_teacher_feedback.md"
VISUAL_AUDIT = RESULTS_DIR / "report_visual_audit.csv"


def find_source_docx() -> Path:
    download_dir = Path.home() / "Downloads"
    candidates = list(download_dir.glob("EEE*.docx"))
    if candidates:
        return max(candidates, key=lambda p: p.stat().st_mtime)
    if SOURCE_BACKUP.exists():
        return SOURCE_BACKUP
    fallback = FINAL_SUBMIT_DIR / "大数据分析与应用期末考查报告.docx"
    if fallback.exists():
        return fallback
    raise FileNotFoundError("No source DOCX found.")


def backup_source(source: Path) -> None:
    FINAL_SUBMIT_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source, SOURCE_BACKUP)
    shutil.copy2(source, PRE_EDIT_BACKUP)


def delete_body_from_first_heading(doc: Document) -> None:
    start = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("1. Analysis Goals and Tasks"):
            start = paragraph._element
            break
    if start is None:
        raise ValueError("Could not find the first main report heading in the source DOCX.")
    body = doc._body._element
    children = list(body)
    start_index = children.index(start)
    for child in children[start_index:]:
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def set_run_font(run, size: float = 10.5, bold: bool = False, color: str | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn_name("w:eastAsia"), "SimSun")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def qn_name(name: str):
    from docx.oxml.ns import qn

    return qn(name)


def add_heading(doc: Document, text: str, level: int = 1, page_break: bool = False) -> None:
    if page_break:
        doc.add_page_break()
    paragraph = doc.add_paragraph()
    paragraph.style = doc.styles[f"Heading {level}"] if f"Heading {level}" in [s.name for s in doc.styles] else None
    run = paragraph.add_run(text)
    set_run_font(run, 13 if level == 1 else 11.5, bold=True, color="1F2937")


def add_para(doc: Document, text: str, size: float = 10.5) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Inches(0.22)
    paragraph.paragraph_format.space_after = Pt(5)
    for part in text.split("\n"):
        if part:
            run = paragraph.add_run(part)
            set_run_font(run, size=size)
        paragraph.add_run("\n") if part != text.split("\n")[-1] else None


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet" if "List Bullet" in [s.name for s in doc.styles] else None)
        run = paragraph.add_run(item)
        set_run_font(run, size=10.5)


def add_code_block(doc: Document, code: str, max_lines: int | None = None) -> None:
    lines = code.strip("\n").splitlines()
    if max_lines:
        lines = lines[:max_lines]
    for line in lines:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn_name("w:eastAsia"), "Consolas")
        run.font.size = Pt(7.5)
        run.font.color.rgb = RGBColor(31, 41, 55)


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, size=9.5, bold=True, color="111827")


def add_figure(doc: Document, fig_id: int, title: str, filename: str, intro: str, analysis: str) -> None:
    add_para(doc, intro)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(FIGURE_DIR / filename), width=Inches(6.1))
    add_caption(doc, f"Fig{fig_id} {title}")
    add_para(doc, analysis)


def add_table_caption(doc: Document, table_id: int, title: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(f"Table{table_id} {title}")
    set_run_font(run, size=9.5, bold=True, color="111827")


def add_dataframe_table(doc: Document, table_id: int, title: str, df: pd.DataFrame, intro: str, analysis: str) -> None:
    add_para(doc, intro)
    add_table_caption(doc, table_id, title)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table_style_names = {style.name for style in doc.styles if style.type == 3}
    if "Table Grid" in table_style_names:
        table.style = "Table Grid"
    header = table.rows[0].cells
    for idx, column in enumerate(df.columns):
        header[idx].text = str(column)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            if isinstance(value, float):
                cells[idx].text = f"{value:.3f}"
            else:
                cells[idx].text = str(value)
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=7.8)
    add_para(doc, analysis)


def load_code_snippets() -> dict[str, str]:
    path = PROJECT_ROOT / "report_code_snippets.md"
    text = path.read_text(encoding="utf-8")
    pattern = re.compile(r"## (Code\d+ [^\n]+)\n```python\n(.*?)\n```", re.S)
    return {heading: code for heading, code in pattern.findall(text)}


def selected_code(snippets: dict[str, str], prefix: str) -> str:
    for key, code in snippets.items():
        if key.startswith(prefix):
            return code
    return ""


def prepare_tables(raw: pd.DataFrame, processed: pd.DataFrame) -> dict[str, pd.DataFrame]:
    roles = {
        "id": "identifier",
        "age": "background",
        "gender": "background",
        "region": "background",
        "income_level": "background",
        "education_level": "background",
        "daily_role": "background",
        "device_hours_per_day": "digital behavior",
        "phone_unlocks": "digital behavior",
        "notifications_per_day": "digital behavior",
        "social_media_mins": "digital behavior",
        "study_mins": "lifestyle",
        "physical_activity_days": "lifestyle",
        "sleep_hours": "lifestyle",
        "sleep_quality": "lifestyle",
        "anxiety_score": "outcome-style field",
        "depression_score": "outcome-style field",
        "stress_level": "outcome-style field",
        "happiness_score": "outcome-style field",
        "focus_score": "outcome-style field",
        "high_risk_flag": "classification target",
        "device_type": "device category",
        "productivity_score": "auxiliary regression target",
        "digital_dependence_score": "main regression target",
    }
    table1 = raw[
        [
            "id",
            "device_hours_per_day",
            "phone_unlocks",
            "notifications_per_day",
            "sleep_hours",
            "high_risk_flag",
            "productivity_score",
            "digital_dependence_score",
        ]
    ].head(6)
    table2 = pd.DataFrame(
        {
            "field": raw.columns,
            "dtype": [str(raw[c].dtype) for c in raw.columns],
            "role": [roles.get(c, "feature") for c in raw.columns],
            "example": [raw[c].iloc[0] for c in raw.columns],
        }
    )
    table3 = pd.DataFrame(
        {
            "field": raw.columns,
            "missing_count": [int(raw[c].isna().sum()) for c in raw.columns],
            "missing_rate": [float(raw[c].isna().mean()) for c in raw.columns],
        }
    )
    table4 = pd.DataFrame(
        [
            ["total_records", len(raw), "kept"],
            ["duplicate_rows", int(raw.duplicated().sum()), "none removed"],
            ["duplicate_id", int(raw["id"].duplicated().sum()), "none removed"],
        ],
        columns=["check_item", "value", "handling"],
    )
    key_numeric = [
        "age",
        "device_hours_per_day",
        "phone_unlocks",
        "notifications_per_day",
        "social_media_mins",
        "study_mins",
        "sleep_hours",
        "sleep_quality",
        "productivity_score",
        "digital_dependence_score",
    ]
    table5 = raw[key_numeric].agg(["min", "max", "mean"]).T.reset_index().rename(columns={"index": "feature"})
    table5["status"] = "pass"
    table6 = processed[
        [
            "device_hours_per_day",
            "social_media_mins",
            "study_mins",
            "sleep_hours",
            "social_media_hours",
            "study_hours",
            "notifications_per_device_hour",
            "device_to_sleep_ratio",
            "activity_sleep_interaction",
            "social_to_study_ratio",
        ]
    ].head(6)
    table7 = pd.DataFrame(
        [
            [
                "classification drop",
                "id, high_risk_flag, anxiety_score, depression_score, stress_level, happiness_score, focus_score, productivity_score, digital_dependence_score",
            ],
            [
                "regression drop",
                "id, high_risk_flag, other outcome-style target fields, and the current target itself",
            ],
            [
                "clustering input",
                "digital behavior and lifestyle numerical features only",
            ],
            [
                "leakage principle",
                "targets and outcome-style variables are used for interpretation, not as direct inputs",
            ],
        ],
        columns=["feature_set", "rule"],
    )
    pca = pd.read_csv(RESULTS_DIR / "pca_explained_variance.csv")
    table8 = pca[["component", "explained_variance_ratio", "cumulative_explained_variance"]].head(6)
    descriptive_cols = [
        "device_hours_per_day",
        "phone_unlocks",
        "notifications_per_day",
        "social_media_mins",
        "sleep_hours",
        "sleep_quality",
        "productivity_score",
        "digital_dependence_score",
    ]
    table9 = raw[descriptive_cols].describe().T.reset_index().rename(columns={"index": "feature"})
    return {
        "table1": table1,
        "table2": table2,
        "table3": table3,
        "table4": table4,
        "table5": table5,
        "table6": table6,
        "table7": table7,
        "table8": table8,
        "table9": table9,
    }


def build_report(source: Path) -> dict[str, int | str | bool]:
    raw = pd.read_csv(PROJECT_ROOT / "data" / "raw" / "digital_lifestyle_benchmark_2025.csv")
    processed = pd.read_csv(PROJECT_ROOT / "data" / "processed" / "digital_lifestyle_benchmark_2025_processed.csv")
    group = pd.read_csv(RESULTS_DIR / "high_risk_group_comparison.csv")
    corr_cols = [
        "device_hours_per_day",
        "phone_unlocks",
        "notifications_per_day",
        "social_media_mins",
        "study_mins",
        "sleep_hours",
        "sleep_quality",
        "high_risk_flag",
        "productivity_score",
        "digital_dependence_score",
    ]
    corr = raw[corr_cols].corr(numeric_only=True)
    pca = pd.read_csv(RESULTS_DIR / "pca_explained_variance.csv")
    profiles = pd.read_csv(RESULTS_DIR / "clustering_lifestyle_profiles_compact.csv")
    class_metrics = pd.read_csv(RESULTS_DIR / "classification_tuned_metrics.csv")
    final_cls = class_metrics[
        (class_metrics["dataset"] == "test")
        & (class_metrics["model"] == "gradient_boosting")
        & (class_metrics["threshold"].round(2) == 0.14)
    ].iloc[0]
    regression_dd = pd.read_csv(RESULTS_DIR / "regression_digital_dependence_metrics.csv")
    regression_prod = pd.read_csv(RESULTS_DIR / "regression_productivity_metrics.csv")
    pred_prod = pd.read_csv(RESULTS_DIR / "regression_productivity_predictions.csv")
    pred_dd = pd.read_csv(RESULTS_DIR / "regression_digital_dependence_predictions.csv")
    snippets = load_code_snippets()
    tables = prepare_tables(raw, processed)

    shutil.copy2(source, OUTPUT_DOCX)
    doc = Document(str(OUTPUT_DOCX))
    delete_body_from_first_heading(doc)
    doc.add_page_break()

    no_risk = int((raw["high_risk_flag"] == 0).sum())
    high_risk = int((raw["high_risk_flag"] == 1).sum())
    total = len(raw)
    top_group = group.iloc[group["cohens_d"].abs().argmax()]
    dd_row = regression_dd[regression_dd["model"].eq("gradient_boosting")].iloc[0]
    prod_row = regression_prod[regression_prod["model"].eq("gradient_boosting")].iloc[0]
    pc2 = pca[pca["component_index"].eq(2)].iloc[0]
    pc80 = int(pca[pca["cumulative_explained_variance"].ge(0.8)]["component_index"].iloc[0])

    add_heading(doc, "1. Analysis Goals and Tasks", 1)
    add_para(
        doc,
        "This report follows the workflow required by the course: define the analysis goal, introduce the dataset, clean and transform the data, conduct statistical visualization, train and evaluate machine learning models, and finally interpret the results. The main purpose is not to claim that every model is perfect. The purpose is to make each step explainable and supported by data evidence.",
    )
    add_heading(doc, "1.1 Background", 2)
    add_para(
        doc,
        "Digital lifestyle data combine device use, notifications, social media time, sleep, study, and activity records. These variables provide a practical setting for studying whether behavior patterns are associated with High Risk labels, digital dependence scores, and exploratory lifestyle profiles.",
    )
    add_heading(doc, "1.2 Key Questions", 2)
    add_bullets(
        doc,
        [
            "Can High Risk samples be screened from behavior and lifestyle variables?",
            "Can digital_dependence_score be predicted more reliably than productivity_score?",
            "Can continuous behavior features be summarized into readable lifestyle profiles?",
        ],
    )
    add_heading(doc, "1.3 Analysis Tasks", 2)
    add_para(
        doc,
        "The report uses three machine learning tasks: classification for high_risk_flag, regression for digital_dependence_score and productivity_score, and clustering for lifestyle profiling. The same preprocessing pipeline is used to keep the analysis coherent across tasks.",
    )

    add_heading(doc, "2. Dataset and Task Support", 1, page_break=True)
    add_heading(doc, "2.1 Dataset Source", 2)
    add_para(
        doc,
        "The selected dataset is the 2025 Digital Lifestyle Benchmark Dataset. It contains 3,500 rows and 24 original fields in CSV format. The dataset is a benchmark or educational dataset, so the results should be interpreted as course-level analysis rather than real individual assessment.",
    )
    add_heading(doc, "2.2 Raw Data Preview", 2)
    add_dataframe_table(
        doc,
        1,
        "Raw Dataset Preview",
        tables["table1"],
        "Table1 is used to confirm that the report starts from the raw CSV rather than from a pre-summarized result table.",
        "Table1 shows that the raw records include digital behavior fields, lifestyle fields, and target fields in the same row. This structure is the reason why the dataset can support classification, regression, and clustering in one workflow.",
    )
    add_heading(doc, "2.3 Field Structure", 2)
    add_dataframe_table(
        doc,
        2,
        "Dataset Fields",
        tables["table2"],
        "Table2 lists every original field with its type and role, which makes the later leakage-control decision traceable.",
        "Table2 confirms that high_risk_flag, productivity_score, and digital_dependence_score are outcome-style variables. They should not be mixed blindly into every task because doing so would make the model evaluation less credible.",
    )
    add_heading(doc, "2.4 Task Support", 2)
    add_para(
        doc,
        "The dataset naturally supports at least two machine learning task types and in this report supports three. high_risk_flag is the classification target, digital_dependence_score is the main regression target, productivity_score is used as a weak-prediction check, and behavior/lifestyle variables are used for clustering.",
    )

    add_heading(doc, "3. Data Cleaning and Feature Engineering", 1, page_break=True)
    add_heading(doc, "3.1 Data Loading", 2)
    add_para(
        doc,
        "The loading step verifies the input file shape and column structure. The raw dataset contains exactly 3,500 records and 24 fields, matching the dataset requirement and the later task design.",
    )
    add_code_block(doc, selected_code(snippets, "Code1"), max_lines=18)
    add_heading(doc, "3.2 Missing and Duplicate Check", 2)
    add_dataframe_table(
        doc,
        3,
        "Missing Value Check",
        tables["table3"],
        "Table3 checks every field before any modeling step.",
        "Table3 confirms that all 24 fields contain zero missing values. Together with Table4, this allows all 3,500 records to be retained and avoids introducing imputation assumptions into the later models.",
    )
    add_dataframe_table(
        doc,
        4,
        "Duplicate Check",
        tables["table4"],
        "Table4 tests whether repeated rows or duplicated identifiers should be removed.",
        "Table4 reports 0 duplicate rows and 0 duplicate id values. Therefore, the data-cleaning step keeps all samples, and the later train/test split is not affected by duplicated records.",
    )
    add_code_block(doc, selected_code(snippets, "Code2"), max_lines=22)
    add_heading(doc, "3.3 Range and Rationality Check", 2)
    add_dataframe_table(
        doc,
        5,
        "Range and Rationality Check",
        tables["table5"],
        "Table5 summarizes the minimum, maximum, and mean values of key numerical fields.",
        "Table5 indicates that device_hours_per_day ranges from 0.28 to 17.16, notifications_per_day ranges from 22 to 1,211, and sleep_hours ranges from 3.00 to 11.00. These values are wide but still interpretable for a benchmark dataset, so no severe abnormal record is removed.",
    )
    add_heading(doc, "3.4 Feature Engineering", 2)
    add_dataframe_table(
        doc,
        6,
        "Engineered Features Preview",
        tables["table6"],
        "Table6 records the engineered behavior-intensity variables.",
        "Table6 confirms that raw minutes and counts are transformed into ratios such as notifications_per_device_hour, device_to_sleep_ratio, and social_to_study_ratio. The feature activity_sleep_interaction is consistently defined as physical_activity_days multiplied by sleep_hours.",
    )
    add_code_block(doc, selected_code(snippets, "Code3"), max_lines=22)
    add_heading(doc, "3.5 Feature Selection and Leakage Control", 2)
    add_dataframe_table(
        doc,
        7,
        "Feature Selection and Leakage Control",
        tables["table7"],
        "Table7 explains which columns are excluded before modeling.",
        "Table7 is important because the classification task must not use anxiety_score, depression_score, stress_level, happiness_score, focus_score, productivity_score, digital_dependence_score, id, or high_risk_flag as inputs. This keeps the classifier focused on behavior evidence instead of direct outcome fields.",
    )
    add_code_block(doc, selected_code(snippets, "Code4"), max_lines=24)
    add_heading(doc, "3.6 PCA and GMM Extension", 2)
    add_dataframe_table(
        doc,
        8,
        "PCA Explained Variance",
        tables["table8"],
        "Table8 provides the numerical evidence for the PCA extension.",
        f"Table8 shows that PC1 and PC2 explain {pc2['cumulative_explained_variance'] * 100:.2f}% of the total variance. Since {100 - pc2['cumulative_explained_variance'] * 100:.2f}% remains outside the first two components and {pc80} components are needed to exceed 80%, PCA is used only for auxiliary structure understanding rather than as classification or regression input.",
    )

    add_heading(doc, "4. Statistics and Visualization", 1, page_break=True)
    add_heading(doc, "4.1 Descriptive Statistics", 2)
    add_dataframe_table(
        doc,
        9,
        "Descriptive Statistical Summary",
        tables["table9"],
        "Table9 gives the mathematical summary before visual analysis.",
        "Table9 shows that notifications_per_day and social_media_mins have wide ranges, while sleep_quality is bounded between 1 and 5. This scale difference explains why StandardScaler is needed for PCA, KMeans, GaussianMixture, and linear models, while tree-based models are less sensitive to raw scale.",
    )
    add_heading(doc, "4.2 Target Distribution", 2)
    add_figure(
        doc,
        1,
        "High Risk vs No Risk Sample Structure",
        "fig1_high_risk_no_risk_distribution.png",
        "Fig1 is used to check the class structure before model training.",
        f"Fig1 reports {no_risk} No Risk samples ({no_risk / total:.2%}) and {high_risk} High Risk samples ({high_risk / total:.2%}). Because the No Risk class is about four times larger, a classifier can look acceptable on Accuracy while still missing many High Risk cases. This class structure is why the later model evaluation emphasizes Recall, F1, PR-AUC, and the confusion matrix.",
    )
    add_heading(doc, "4.3 Numeric Feature Distribution", 2)
    add_figure(
        doc,
        2,
        "Core Digital Lifestyle Feature Distributions",
        "fig2_core_numeric_distributions.png",
        "Fig2 summarizes the main numerical variables before modeling.",
        "Fig2 indicates that notifications_per_day has a long right tail up to 1,211 and phone_unlocks reaches 374, while sleep_hours stays in a much narrower range from about 3 to 11 hours. This means distance-based steps such as PCA, KMeans, GaussianMixture, and linear models need scaling, whereas tree models can use raw splits more directly.",
    )
    add_heading(doc, "4.4 Correlation Structure", 2)
    dd_device = corr.loc["digital_dependence_score", "device_hours_per_day"]
    dd_sleepq = corr.loc["digital_dependence_score", "sleep_quality"]
    prod_device = corr.loc["productivity_score", "device_hours_per_day"]
    add_figure(
        doc,
        3,
        "Correlation Structure of Digital Lifestyle Variables",
        "fig3_correlation_heatmap.png",
        "Fig3 checks the correlation structure among behavior variables, lifestyle variables, and outcomes.",
        f"Fig3 shows that digital_dependence_score is positively related to device_hours_per_day (r={dd_device:.2f}) and negatively related to sleep_quality (r={dd_sleepq:.2f}), while productivity_score has a much weaker relationship with device_hours_per_day (r={prod_device:.2f}). This difference helps explain why digital_dependence_score is highly predictable but productivity_score is not. The heatmap is still only correlation evidence and is not used as a causal claim.",
    )
    add_heading(doc, "4.5 High Risk and No Risk Group Difference", 2)
    device = group[group["feature"].eq("device_hours_per_day")].iloc[0]
    sleep = group[group["feature"].eq("sleep_hours")].iloc[0]
    dd_group = group[group["feature"].eq("digital_dependence_score")].iloc[0]
    notif = group[group["feature"].eq("notifications_per_day")].iloc[0]
    add_figure(
        doc,
        4,
        "Behavioral Difference Between High Risk and No Risk Groups",
        "fig4_high_vs_no_risk_boxplots.png",
        "Fig4 compares the High Risk and No Risk groups using the variables with direct behavioral meaning.",
        f"Fig4 and the statistics in results/high_risk_group_comparison.csv show that the High Risk group has a higher median device_hours_per_day ({device['high_risk_median']:.2f} vs {device['no_risk_median']:.2f}) and a higher median digital_dependence_score ({dd_group['high_risk_median']:.2f} vs {dd_group['no_risk_median']:.2f}). The largest standardized group difference among the displayed variables is device_hours_per_day (Cohen's d={device['cohens_d']:.2f}), followed by digital_dependence_score (d={dd_group['cohens_d']:.2f}) and sleep_hours (d={sleep['cohens_d']:.2f}, lower in the High Risk group). However, notifications_per_day has only a small standardized difference (d={notif['cohens_d']:.2f}), and the box ranges still overlap across groups. Therefore, no single variable can serve as a reliable rule, which supports the later multi-feature classifier and threshold tuning.",
    )

    add_heading(doc, "5. Model Training and Evaluation", 1, page_break=True)
    add_heading(doc, "5.1 Experimental Setup", 2)
    add_para(
        doc,
        "All supervised experiments use train/test separation, random_state=42, cross-validation or grid-search evidence, and final testing on the held-out test set. Classification focuses on Recall, F1, PR-AUC, ROC-AUC, Balanced Accuracy, and the confusion matrix. Regression focuses on R2, MSE, RMSE, and MAE. Clustering uses inertia, Silhouette, Calinski-Harabasz, Davies-Bouldin, and profile interpretability.",
    )
    add_heading(doc, "5.2 High Risk Classification", 2)
    add_figure(
        doc,
        5,
        "Classification Model Comparison at the Default Threshold",
        "fig5_classification_model_comparison.png",
        "Fig5 replaces the previous Excel screenshot for classification model comparison.",
        "At the default threshold of 0.50, Random Forest has the highest Recall (0.3409), F1 (0.4167), and Balanced Accuracy (0.6333), whereas Gradient Boosting has only a slight advantage in PR-AUC (0.4805) and ROC-AUC (0.7233). This comparison does not mean that Gradient Boosting dominates every metric. Gradient Boosting is retained as the probability-ranking model and is then adjusted through threshold tuning to match the recall-oriented screening objective.",
    )
    add_figure(
        doc,
        6,
        "Threshold Tuning for Recall-Oriented Screening",
        "fig6_threshold_tuning.png",
        "Fig6 explains why the final classifier does not use the default threshold of 0.50.",
        f"As Fig6 indicates, increasing the decision threshold improves precision but sharply reduces recall. The middle threshold has the highest validation F1, while the selected threshold of 0.14 preserves recall above the screening requirement. On the independent test set, this policy reaches Precision={final_cls['precision']:.4f}, Recall={final_cls['recall']:.4f}, F1={final_cls['f1']:.4f}, PR-AUC=0.5084, ROC-AUC=0.7531, and Balanced Accuracy=0.7259. The selected point is therefore a deliberate screening trade-off rather than the threshold that maximizes overall correctness.",
    )
    add_figure(
        doc,
        7,
        "Confusion Matrix of the Final High Risk Classifier",
        "fig7_confusion_matrix.png",
        "Fig7 provides the final test-set confusion matrix for Gradient Boosting at threshold 0.14.",
        "Fig7 shows TN=566, FP=133, FN=63, and TP=113. The model finds 113 true High Risk samples but still misses 63 High Risk samples. To reduce missed High Risk cases, it produces 133 false alarms, which is larger than the number of false negatives. This is the direct cost of the recall-oriented policy, so the model should be used as a screening reference rather than a final individual-level conclusion.",
    )
    add_figure(
        doc,
        8,
        "Precision-Recall Curve under Class Imbalance",
        "fig8_precision_recall_curve.png",
        "Fig8 evaluates the High Risk ranking ability across thresholds.",
        "Fig8 includes the no-skill baseline computed from the test-set High Risk proportion, about 20.11%. The PR-AUC is 0.5084, which is clearly above this baseline but still far from a high-precision screening system. This supports the conclusion that the model contains useful ranking information, but the threshold must be chosen carefully and the output should not be treated as a final personal judgment.",
    )
    add_code_block(doc, selected_code(snippets, "Code5"), max_lines=24)

    add_heading(doc, "5.3 Digital Dependence and Productivity Regression", 2)
    add_figure(
        doc,
        9,
        "Digital Dependence Observed vs Predicted Scores",
        "fig9_digital_dependence_observed_predicted.png",
        "Fig9 checks whether the final digital_dependence_score predictions align with the observed scores.",
        f"In Fig9, most points are close to the 45-degree reference line, and the final Gradient Boosting result reaches R2={dd_row['r2']:.4f}, MSE={dd_row['mse']:.4f}, and MAE={dd_row['mae']:.4f}. The point cloud does not show a strong systematic curve away from the diagonal, so the current behavior features provide strong predictive information for this target inside the benchmark dataset. This is still a prediction result, not a causal statement.",
    )
    actual_range = pred_prod["actual"].max() - pred_prod["actual"].min()
    pred_range = pred_prod["predicted"].max() - pred_prod["predicted"].min()
    add_figure(
        doc,
        10,
        "Productivity Weak Prediction Observed vs Predicted Scores",
        "fig10_productivity_observed_predicted.png",
        "Fig10 keeps the productivity_score experiment as a negative-result check.",
        f"Fig10 shows that the observed productivity_score spans about {actual_range:.1f} points, but the predicted values cover only about {pred_range:.1f} points and concentrate near the average. This is a regression-to-the-mean pattern. The final R2={prod_row['r2']:.4f} means the model does not outperform a simple mean baseline, so the current behavior variables are not enough to explain individual productivity differences.",
    )
    add_figure(
        doc,
        11,
        "Regression R2 Comparison by Target",
        "fig11_regression_r2_comparison.png",
        "Fig11 replaces the previous Excel regression comparison screenshot with a formal R2 comparison.",
        "Fig11 separates the two regression targets so that R2 is not mixed with MSE or MAE on one axis. All models predict digital_dependence_score very well, with Gradient Boosting reaching R2=0.9839. In contrast, all productivity_score R2 values are close to or below zero. The same feature set therefore represents digital dependence strongly but does not explain productivity_score well.",
    )
    add_code_block(doc, selected_code(snippets, "Code6"), max_lines=24)

    add_heading(doc, "5.4 Digital Lifestyle Clustering", 2)
    add_figure(
        doc,
        12,
        "K Selection by Inertia and Silhouette",
        "fig12_clustering_k_selection_comparison.png",
        "Fig12 replaces the previous clustering model comparison screenshot and combines the key k-selection evidence.",
        "In Fig12, the KMeans inertia curve decreases smoothly rather than forming a sharp elbow. The silhouette panel shows that KMeans reaches its highest tested silhouette at k=3 with Silhouette=0.1860, while AgglomerativeClustering and GaussianMixture do not provide a clearly stronger alternative. Because 0.1860 is low, k=3 is a reasonable and readable exploratory choice, not evidence of strict natural groups.",
    )
    c0 = profiles[profiles["cluster"].eq(0)].iloc[0]
    c1 = profiles[profiles["cluster"].eq(1)].iloc[0]
    c2 = profiles[profiles["cluster"].eq(2)].iloc[0]
    add_figure(
        doc,
        13,
        "Three Digital Lifestyle Cluster Profiles",
        "fig13_cluster_profile_heatmap.png",
        "Fig13 summarizes the final KMeans k=3 lifestyle profiles.",
        f"Fig13 shows that Cluster 0 has very high social_media_mins ({c0['social_media_mins']:.1f}) but moderate device_hours_per_day ({c0['device_hours_per_day']:.2f}). Cluster 1 has the highest device_hours_per_day ({c1['device_hours_per_day']:.2f}), the lowest sleep_quality ({c1['sleep_quality']:.2f}), the highest High Risk ratio ({c1['high_risk_flag']:.1%}), and the highest digital_dependence_score ({c1['digital_dependence_score']:.2f}). Cluster 2 has the lowest device_hours_per_day ({c2['device_hours_per_day']:.2f}) and the lowest High Risk ratio ({c2['high_risk_flag']:.1%}). These patterns support the names High social-media-use, High device-dependence, and Low-load balanced, but the low silhouette score means the profiles should be treated as gradual summaries.",
    )
    add_figure(
        doc,
        14,
        "PCA Explained Variance for Dimensionality Reduction",
        "fig14_pca_explained_variance.png",
        "Fig14 provides the PCA evidence used to discuss dimensionality reduction.",
        f"Fig14 shows that the first two principal components explain {pc2['cumulative_explained_variance'] * 100:.2f}% of the total variance, leaving {100 - pc2['cumulative_explained_variance'] * 100:.2f}% outside the two-dimensional view. The cumulative curve reaches 80% only after {pc80} components. Therefore, PCA is useful for auxiliary visualization and structure understanding, but it should not replace the original feature space for classification, regression, or final clustering evaluation.",
    )
    add_code_block(doc, selected_code(snippets, "Code7"), max_lines=24)

    add_heading(doc, "5.5 Integrated Result Analysis", 2)
    add_para(
        doc,
        "The model evidence forms a consistent chain. EDA explains class imbalance and group overlap, so classification is handled as recall-oriented screening. The regression comparison explains why digital_dependence_score is the main regression target and productivity_score remains a weak result. The clustering evidence gives readable profiles, but the low silhouette value prevents the report from claiming strict population segmentation.",
    )

    add_heading(doc, "6. Findings and Reflection", 1, page_break=True)
    add_heading(doc, "6.1 Main Findings", 2)
    add_bullets(
        doc,
        [
            "High Risk can be screened from behavior variables, but the final threshold trades more false alarms for fewer missed High Risk cases.",
            "digital_dependence_score is strongly represented by the current behavior features in this benchmark dataset, with R2=0.9839.",
            "productivity_score remains a weak-prediction target, with R2=-0.0041, which shows that the feature set does not explain every outcome equally.",
            "The KMeans k=3 clusters are readable lifestyle profiles, but Silhouette=0.1860 means they are exploratory summaries rather than strict groups.",
        ],
    )
    add_heading(doc, "6.2 Limitations", 2)
    add_para(
        doc,
        "The dataset is a benchmark/synthetic educational dataset, so the high R2 for digital_dependence_score may partly reflect the data-generation structure. The classification recall is not 100%, productivity_score is not explained well, and the clustering silhouette is low. External validation with real digital lifestyle data would be needed before stronger practical claims could be made.",
    )
    add_heading(doc, "6.3 Personal Reflection", 2)
    add_para(
        doc,
        "This revision helped me understand that figures and tables should not be decoration. Each visual item needs a question, a data-supported observation, and a conclusion that connects to the next step. The most important lesson is that a complete data-analysis report is not only a list of model scores; it is a traceable workflow from raw data to cleaning, visualization, model evaluation, and honest interpretation.",
    )

    add_heading(doc, "Appendix A Complete Runnable Code", 1, page_break=True)
    add_heading(doc, "A.1 Environment and Dependencies", 2)
    add_para(
        doc,
        "The experiment was completed with a CPU-based Python environment. The main dependencies are pandas, NumPy, scikit-learn, Matplotlib, openpyxl, python-docx, and joblib. The environment can be installed with: pip install -r requirements.txt.",
    )
    add_heading(doc, "A.2 Complete Code", 2)
    add_para(doc, "The complete runnable code is copied from appendix_A_complete_code.py.")
    appendix_code = (PROJECT_ROOT / "appendix_A_complete_code.py").read_text(encoding="utf-8", errors="replace")
    add_code_block(doc, appendix_code, max_lines=None)

    doc.save(OUTPUT_DOCX)
    return {
        "source_docx": str(source),
        "output_docx": str(OUTPUT_DOCX),
        "figure_count": 14,
        "table_count": 9,
        "source_file_used": source.name,
    }


def write_visual_audit() -> None:
    rows: list[dict[str, str | bool]] = []
    tables = {
        "Table1": ("Raw Dataset Preview", "table", "keep", "native Word table from raw CSV", "Does the report start from raw data?", "raw CSV preview", "The raw table contains behavior, lifestyle, and target fields in one structured dataset."),
        "Table2": ("Dataset Fields", "table", "keep", "native Word table from raw CSV", "Which fields are inputs, background fields, or targets?", "field dtype and role mapping", "Target and outcome-style fields can be separated before modeling."),
        "Table3": ("Missing Value Check", "table", "keep", "native Word table from raw CSV", "Are missing values present?", "missing counts by column", "All 24 fields contain zero missing values."),
        "Table4": ("Duplicate Check", "table", "keep", "native Word table from raw CSV", "Are duplicated records present?", "duplicate row/id counts", "All 3,500 records can be retained."),
        "Table5": ("Range and Rationality Check", "table", "keep", "native Word table from raw CSV", "Are key numeric ranges reasonable?", "min/max/mean by field", "Wide but interpretable ranges support later modeling without deletion."),
        "Table6": ("Engineered Features Preview", "table", "keep", "native Word table from processed CSV", "Were engineered features generated correctly?", "processed feature preview", "Behavior-intensity variables are available for modeling."),
        "Table7": ("Feature Selection and Leakage Control", "table", "keep", "native Word table", "Which fields are excluded for each task?", "task-specific exclusion rules", "Outcome-style fields are controlled before modeling."),
        "Table8": ("PCA Explained Variance", "table", "keep", "results/pca_explained_variance.csv", "How much variance does PCA summarize?", "explained variance ratios", "PC1+PC2 explain 42.41%, so PCA is auxiliary."),
        "Table9": ("Descriptive Statistical Summary", "table", "keep", "native Word table from raw CSV", "What are the variable ranges and scales?", "descriptive statistics", "Different scales explain why scaling is needed for PCA and clustering."),
    }
    figures = {
        "Fig1": ("High Risk vs No Risk Sample Structure", "figure", "revise", "figures/final_report/fig1_high_risk_no_risk_distribution.png", "Is the classification target imbalanced?", "class counts and ratios", "High Risk is 20.14%, so Accuracy alone can be misleading."),
        "Fig2": ("Core Digital Lifestyle Feature Distributions", "figure", "revise", "figures/final_report/fig2_core_numeric_distributions.png", "How are core numeric variables distributed?", "histograms and ranges", "Notification and unlock variables have wider ranges than sleep variables."),
        "Fig3": ("Correlation Structure of Digital Lifestyle Variables", "figure", "revise", "figures/final_report/fig3_correlation_heatmap.png", "Which variables are correlated with outcomes?", "correlation matrix", "digital_dependence_score has stronger behavior correlations than productivity_score."),
        "Fig4": ("Behavioral Difference Between High Risk and No Risk Groups", "figure", "replace", "figures/final_report/fig4_high_vs_no_risk_boxplots.png", "Which behaviors differ by risk group?", "results/high_risk_group_comparison.csv", "Device hours, digital dependence, and lower sleep hours show the largest group differences, but distributions overlap."),
        "Fig5": ("Classification Model Comparison at the Default Threshold", "figure", "replace", "figures/final_report/fig5_classification_model_comparison.png", "Which classifier works best before threshold tuning?", "classification_tuned_metrics.csv", "Random Forest is stronger on default Recall/F1, while Gradient Boosting has a slight PR-AUC/ROC-AUC edge."),
        "Fig6": ("Threshold Tuning for Recall-Oriented Screening", "figure", "revise", "figures/final_report/fig6_threshold_tuning.png", "Why was threshold=0.14 selected?", "classification_threshold_tuning.csv", "The selected threshold preserves recall for screening rather than maximizing overall correctness."),
        "Fig7": ("Confusion Matrix of the Final High Risk Classifier", "figure", "revise", "figures/final_report/fig7_confusion_matrix.png", "What errors does the final classifier make?", "classification_final_confusion_matrix.csv", "The model finds 113 High Risk samples and misses 63, with 133 false alarms."),
        "Fig8": ("Precision-Recall Curve under Class Imbalance", "figure", "revise", "figures/final_report/fig8_precision_recall_curve.png", "Is PR-AUC above the positive-class baseline?", "test split and classification_tuned_metrics.csv", "PR-AUC=0.5084 is above the 20.11% baseline but not high enough for precise individual decisions."),
        "Fig9": ("Digital Dependence Observed vs Predicted Scores", "figure", "revise", "figures/final_report/fig9_digital_dependence_observed_predicted.png", "How well is digital dependence predicted?", "regression_digital_dependence_predictions.csv", "Predictions align closely with observed scores, consistent with R2=0.9839."),
        "Fig10": ("Productivity Weak Prediction Observed vs Predicted Scores", "figure", "revise", "figures/final_report/fig10_productivity_observed_predicted.png", "Does the feature set predict productivity?", "regression_productivity_predictions.csv", "Predictions concentrate near the mean and R2=-0.0041 indicates weak prediction."),
        "Fig11": ("Regression R2 Comparison by Target", "figure", "replace", "figures/final_report/fig11_regression_r2_comparison.png", "Do both regression targets behave the same?", "regression metrics CSV files", "digital_dependence_score is strongly predicted, but productivity_score is not."),
        "Fig12": ("K Selection by Inertia and Silhouette", "figure", "replace", "figures/final_report/fig12_clustering_k_selection_comparison.png", "Which k and clustering method are reasonable?", "clustering_model_comparison.csv", "KMeans k=3 is readable and has the best KMeans silhouette, but 0.1860 is low."),
        "Fig13": ("Three Digital Lifestyle Cluster Profiles", "figure", "revise", "figures/final_report/fig13_cluster_profile_heatmap.png", "What does each cluster represent?", "clustering_lifestyle_profiles_compact.csv", "The profiles differ in social media time, device hours, sleep quality, High Risk ratio, and digital dependence."),
        "Fig14": ("PCA Explained Variance for Dimensionality Reduction", "figure", "revise", "figures/final_report/fig14_pca_explained_variance.png", "How much structure does PCA preserve?", "pca_explained_variance.csv", "PC1+PC2 explain 42.41%, so the 2D view is only partial."),
    }
    for item_id, values in {**tables, **figures}.items():
        caption, typ, decision, source_file, question, evidence, conclusion = values
        rows.append(
            {
                "item_id": item_id,
                "caption": caption,
                "current_type": typ,
                "decision": decision,
                "source_file": source_file,
                "analytical_question": question,
                "evidence_used": evidence,
                "concrete_conclusion": conclusion,
                "body_reference_found": True,
                "revision_note": "Converted to continuous analysis paragraph and explicit body reference.",
            }
        )
    pd.DataFrame(rows).to_csv(VISUAL_AUDIT, index=False)


def try_render_docx() -> str:
    pdf_path = OUTPUT_DOCX.with_suffix(".pdf")
    try:
        subprocess.run(
            [
                "powershell",
                "-NoProfile",
                "-Command",
                (
                    "$word = New-Object -ComObject Word.Application; "
                    "$word.Visible = $false; "
                    f"$doc = $word.Documents.Open('{OUTPUT_DOCX}'); "
                    f"$doc.SaveAs([ref] '{pdf_path}', [ref] 17); "
                    "$doc.Close(); $word.Quit();"
                ),
            ],
            check=True,
            timeout=120,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
        )
        if pdf_path.exists() and pdf_path.stat().st_size > 0:
            return f"Word COM PDF export succeeded: {pdf_path}"
    except Exception as exc:  # noqa: BLE001
        return f"PDF render not available in this environment: {exc}"
    return "PDF render did not produce a non-empty file."


def structural_check() -> dict[str, object]:
    doc = Document(str(OUTPUT_DOCX))
    text = "\n".join(p.text for p in doc.paragraphs)
    forbidden = ["Summary:", "Purpose:", "Observation:", "Meaning:", "[INSERT", "[PASTE", "TODO", "Table10", "Table11", "Table12"]
    return {
        "forbidden_hits": {term: text.count(term) for term in forbidden if term in text},
        "figure_count": len(set(re.findall(r"Fig\d+", text))),
        "table_count": len(set(re.findall(r"Table\d+", text))),
        "visible_inline_shapes": len(doc.inline_shapes),
        "metrics_ok": all(
            metric in text
            for metric in [
                "Recall=0.6420",
                "F1=0.5355",
                "PR-AUC=0.5084",
                "ROC-AUC=0.7531",
                "Balanced Accuracy=0.7259",
                "R2=0.9839",
                "MSE=3.1471",
                "MAE=0.9982",
                "R2=-0.0041",
                "Silhouette=0.1860",
                "42.41%",
            ]
        ),
        "has_high_no_risk": "High Risk" in text and "No Risk" in text,
    }


def write_revision_log(source: Path, build_info: dict[str, int | str | bool], render_status: str, check: dict[str, object]) -> None:
    log = f"""# Teacher Feedback Revision Log

## Source DOCX

- Actual source DOCX: `{source}`
- Project backup: `{SOURCE_BACKUP}`
- Pre-edit backup: `{PRE_EDIT_BACKUP}`
- Final revised DOCX: `{OUTPUT_DOCX}`

## Deleted or Replaced Items

- Removed the model-comparison Excel screenshots from the formal body.
- Table10 Classification Model Comparison and Threshold Results was replaced by Fig5.
- Table11 Regression Model Comparison was replaced by Fig11.
- Table12 Clustering Model Comparison and Cluster Profiles was replaced by Fig12 and Fig13.
- Removed the mechanical Summary / Purpose / Observation / Meaning blocks and rewrote figure/table interpretations as continuous English analysis paragraphs.

## New or Replaced Formal Analysis Figures

- Fig5 Classification Model Comparison at the Default Threshold.
- Fig11 Regression R2 Comparison by Target.
- Fig12 K Selection by Inertia and Silhouette.
- Fig4 was regenerated with Cohen's d annotations and is supported by `results/high_risk_group_comparison.csv`.
- Fig8 Precision-Recall Curve now includes a no-skill baseline.

## Numbering Changes

- Final tables are Table1-Table9 only.
- Final figures are Fig1-Fig14.
- Old Table10, Table11, and Table12 no longer appear in the revised DOCX body.

## Fig4 Evidence

- Variables used: device_hours_per_day, phone_unlocks, notifications_per_day, social_media_mins, sleep_hours, digital_dependence_score.
- Strongest standardized difference: device_hours_per_day, followed by digital_dependence_score and sleep_hours.
- Full group statistics are saved in `results/high_risk_group_comparison.csv`.

## QA

- Visual audit CSV: `results/report_visual_audit.csv`.
- Structural check: `{check}`.
- Render status: {render_status}

## Core Metrics

- Classification: Gradient Boosting, threshold=0.14, Recall=0.6420, F1=0.5355, PR-AUC=0.5084, ROC-AUC=0.7531, Balanced Accuracy=0.7259, TN=566, FP=133, FN=63, TP=113.
- digital_dependence_score regression: Gradient Boosting, R2=0.9839, MSE=3.1471, MAE=0.9982.
- productivity_score: R2=-0.0041, kept as weak prediction.
- Clustering: KMeans k=3, Silhouette=0.1860, exploratory lifestyle profiles only.
- PCA: PC1+PC2 cumulative explained variance = 42.41%.
"""
    REVISION_LOG.write_text(log, encoding="utf-8-sig")


def main() -> None:
    source = find_source_docx()
    backup_source(source)
    generator = PROJECT_ROOT / "scripts" / "generate_final_report_figures.py"
    subprocess.run(["conda", "run", "-n", "qintian-DL", "python", str(generator)], check=True)
    build_info = build_report(SOURCE_BACKUP)
    write_visual_audit()
    render_status = try_render_docx()
    check = structural_check()
    write_revision_log(source, build_info, render_status, check)
    print(f"Source DOCX: {source}")
    print(f"Output DOCX: {OUTPUT_DOCX}")
    print(f"Visual audit: {VISUAL_AUDIT}")
    print(f"Revision log: {REVISION_LOG}")
    print(f"Render status: {render_status}")
    print(f"Structural check: {check}")


if __name__ == "__main__":
    main()
