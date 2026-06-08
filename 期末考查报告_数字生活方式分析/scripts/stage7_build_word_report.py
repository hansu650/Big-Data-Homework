"""Build the final Word submission from the Stage 6 LaTeX report.

This script does not run experiments or touch result CSV files. It converts the
validated Stage 6 LaTeX structure into a DOCX based on the teacher's Word
template and writes a structural QA summary for Stage 7.
"""

from __future__ import annotations

import json
import re
import shutil
import subprocess
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


REPORT_TITLE = (
    "数字生活方式下的高风险识别、数字依赖预测与用户画像分析"
    "——基于 2025 Digital Lifestyle Benchmark 数据集的分类、回归与聚类研究"
)

COURSE_NAME_EN = "Big Data Analysis and Applications"
CONTENT_NAME = "Course Report"
TEACHER_NAME = "Li Jie"
OUTPUT_DOCX_NAME = "大数据分析与应用期末考查报告.docx"
OUTPUT_PDF_NAME = "大数据分析与应用期末考查报告.pdf"
TEMPLATE_PATH = Path(r"C:\Users\qintian\Downloads\examination report.docx")

SECTION_ORDER = [
    "01_dataset_selection.tex",
    "02_research_theme.tex",
    "03_preprocessing.tex",
    "04_eda_visualization.tex",
    "05_modeling_evaluation.tex",
    "06_report_code_structure.tex",
    "07_conclusion_reflection.tex",
    "references.tex",
    "appendix_a_code.tex",
    "appendix_b_reproducibility.tex",
    "appendix_c_artifacts.tex",
    "appendix_d_independence.tex",
]

FINAL_TOC = [
    ("第1章 数据集自主选取", ["1.1 课程数据集要求说明", "1.2 候选数据集比较与 pfm 数据集弃用原因", "1.3 最终数据集来源与合规性检查", "1.4 数据字段、样本规模与任务支持情况"]),
    ("第2章 自主选题与分析视角", ["2.1 业务背景与问题来源", "2.2 报告题目与研究目标", "2.3 分析视角与研究问题", "2.4 本文主要工作与技术路线"]),
    ("第3章 数据预处理", ["3.1 原始数据读取与字段检查", "3.2 缺失值、重复值与异常值检查", "3.3 数据合理性校验", "3.4 特征工程与特征提取", "3.5 PCA 降维扩展与 LCA/GMM 说明", "3.6 特征筛选与目标泄漏控制", "3.7 预处理结果小结"]),
    ("第4章 探索性分析与数据可视化", ["4.1 描述性统计分析", "4.2 目标变量分布分析", "4.3 数值变量分布与相关性分析", "4.4 高风险组与非高风险组差异分析", "4.5 类别变量风险率分析", "4.6 行为变量与结果变量关系分析", "4.7 EDA 结论小结"]),
    ("第5章 机器学习建模、调参与模型评估", ["5.1 实验设置与评价指标", "5.2 分类任务：高风险数字生活方式筛查", "5.3 回归任务：数字依赖预测与生产力弱预测", "5.4 聚类任务：数字生活方式用户画像", "5.5 模型结果综合分析", "5.6 业务场景解释与结果边界"]),
    ("第6章 报告架构与代码规范呈现", ["6.1 项目目录与工程结构", "6.2 Notebook 与 Python 源码组织", "6.3 核心代码片段与说明", "6.4 图表、结果文件与复现说明", "6.5 完整代码附录说明"]),
    ("第7章 实验结论与个人反思", ["7.1 核心实验结论", "7.2 数据规律与应用价值", "7.3 模型优点、不足与局限性", "7.4 后续改进方向", "7.5 课程学习反思"]),
    ("参考文献", []),
    ("附录A 核心代码", []),
    ("附录B 完整实验运行说明", []),
    ("附录C 结果文件与图表清单", []),
    ("附录D 个人完成情况与独立性说明", []),
]

BANNED_TERMS = [
    "湖北大学本科课程设计",
    "人工智能技术与应用",
    "校园室内巡检",
    "CampusDepthSegLite",
    "RGB-D",
    "语义分割",
    "NYUDepthV2",
    "PyTorch",
    "王雷春",
    "副教授",
]

LATEX_RESIDUES = [
    "\\chapter",
    "\\section",
    "\\texttt",
    "\\input",
    "\\maybefigure",
    "\\begin",
    "\\end",
    "Missing figure",
    "placeholder",
    "TODO",
]


def find_project_root() -> Path:
    repo_root = Path.cwd()
    for child in repo_root.iterdir():
        if child.is_dir() and child.name.startswith("期末考查报告_数字生活方式分析"):
            return child
    raise FileNotFoundError("Cannot locate final report project directory.")


def clear_document_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None, font: str = "宋体") -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_style_font(doc: Document, style_name: str, font: str, size: float, bold: bool = False, color: str | None = None) -> None:
    style = doc.styles[style_name]
    style.font.name = font
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    style.font.size = Pt(size)
    style.font.bold = bold
    if color:
        style.font.color.rgb = RGBColor.from_string(color)


def setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.6)
    section.bottom_margin = Cm(2.4)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.6)
    ensure_paragraph_style(doc, "Normal")
    set_style_font(doc, "Normal", "宋体", 10.5)
    for style_name, font, size in [
        ("Heading 1", "黑体", 16),
        ("Heading 2", "黑体", 14),
        ("Heading 3", "黑体", 12),
        ("List Bullet", "宋体", 10.5),
        ("List Number", "宋体", 10.5),
    ]:
        ensure_paragraph_style(doc, style_name)
        set_style_font(doc, style_name, font, size, bold=style_name.startswith("Heading"), color="000000")


def ensure_paragraph_style(doc: Document, style_name: str):
    try:
        return doc.styles[style_name]
    except KeyError:
        return doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)


def add_centered_text(doc: Document, text: str, size: float = 12, bold: bool = False, spacing_after: float = 6) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(spacing_after)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)


def add_plain_paragraph(doc: Document, text: str, first_line_indent: bool = True) -> None:
    if not text.strip():
        return
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(4)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Pt(21)
    r = p.add_run(text.strip())
    set_run_font(r, size=10.5)


def set_cell_text(cell, text: str, font_size: float = 10.5, bold: bool = False, align: str = "center") -> None:
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == "center" else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=font_size, bold=bold)


def shade_cell(cell, fill: str = "EAF2F8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "666666")


def add_cover(doc: Document) -> None:
    add_centered_text(doc, "湖北大学 2025--2026 学年度第 2 学期课程考查试题纸", size=15, bold=False, spacing_after=10)
    add_centered_text(doc, "The paper of course exam", size=12, spacing_after=24)

    t = doc.add_table(rows=4, cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    info = [
        ("Name:", COURSE_NAME_EN, "", ""),
        ("Content:", CONTENT_NAME, "", ""),
        ("Institution:", "________________", "Teacher:", TEACHER_NAME),
        ("Grade and major:", "________________", "", ""),
    ]
    for row, values in zip(t.rows, info):
        for cell, value in zip(row.cells, values):
            set_cell_text(cell, value, font_size=11, bold=value.endswith(":") or value == "Teacher:", align="left")
    doc.add_paragraph()

    add_centered_text(doc, "课程报告题目", size=16, bold=True, spacing_after=12)
    add_centered_text(doc, f"《{REPORT_TITLE}》", size=13.5, bold=True, spacing_after=30)

    info_table = doc.add_table(rows=6, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row, (label, value) in zip(
        info_table.rows,
        [
            ("学号：", "____________________________"),
            ("姓名：", "____________________________"),
            ("学院：", "____________________________"),
            ("专业年级：", "____________________________"),
            ("Institution:", "____________________________"),
            ("Grade and major:", "____________________________"),
        ],
    ):
        set_cell_text(row.cells[0], label, font_size=12, bold=True, align="right")
        set_cell_text(row.cells[1], value, font_size=12, align="left")
    doc.add_paragraph()

    score_table = doc.add_table(rows=2, cols=3)
    score_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, value in zip(score_table.rows[0].cells, ["Teacher's comments", "Total score", "Grading teacher"]):
        set_cell_text(cell, value, font_size=11, bold=True)
        shade_cell(cell, "F2F2F2")
    for cell in score_table.rows[1].cells:
        set_cell_text(cell, "________________", font_size=11)
    doc.add_paragraph()
    add_centered_text(doc, "湖北大学", size=12, spacing_after=0)
    doc.add_page_break()


def extract_balanced_args(text: str, command: str) -> list[str]:
    start = text.find(command)
    if start < 0:
        return []
    pos = start + len(command)
    args: list[str] = []
    while pos < len(text):
        while pos < len(text) and text[pos].isspace():
            pos += 1
        if pos >= len(text) or text[pos] != "{":
            break
        depth = 0
        arg_start = pos + 1
        pos += 1
        while pos < len(text):
            if text[pos] == "{":
                depth += 1
            elif text[pos] == "}":
                if depth == 0:
                    args.append(text[arg_start:pos])
                    pos += 1
                    break
                depth -= 1
            pos += 1
        else:
            break
    return args


def clean_latex(text: str) -> str:
    s = text.strip()
    if not s:
        return ""
    replacements = {
        "\\%": "%",
        "\\_": "_",
        "\\&": "&",
        "\\$": "$",
        "--": "-",
        "``": "“",
        "''": "”",
    }
    for old, new in replacements.items():
        s = s.replace(old, new)
    for cmd in ["texttt", "textbf", "emph", "url"]:
        pattern = re.compile(rf"\\{cmd}\{{([^{{}}]*(?:\{{[^{{}}]*\}}[^{{}}]*)*)\}}")
        prev = None
        while prev != s:
            prev = s
            s = pattern.sub(lambda m: clean_latex(m.group(1)), s)
    s = s.replace("$R^2$", "R²").replace("$R^2", "R²").replace("R^2", "R²")
    s = re.sub(r"\$([^$]+)\$", lambda m: m.group(1).replace("^2", "²"), s)
    s = s.replace("\\quad", " ")
    s = re.sub(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?", "", s)
    s = s.replace("{", "").replace("}", "")
    s = re.sub(r"\s+", " ", s)
    return s.strip()


def extract_main_abstract(main_tex: str) -> tuple[list[str], str]:
    start = main_tex.index("\\chapter*{摘 要}")
    key = "\\noindent\\textbf{关键词：}"
    end = main_tex.index(key, start)
    body = main_tex[start:end]
    body = body.split("\\addcontentsline{toc}{chapter}{摘 要}", 1)[-1]
    paragraphs = [clean_latex(p) for p in re.split(r"\n\s*\n", body) if clean_latex(p)]
    kw_part = main_tex[end + len(key) :]
    keywords = clean_latex(kw_part.split("\\clearpage", 1)[0])
    return paragraphs, keywords


def add_abstract_and_toc(doc: Document, main_tex: str) -> None:
    paragraphs, keywords = extract_main_abstract(main_tex)
    add_centered_text(doc, "摘 要", size=16, bold=True, spacing_after=12)
    for p in paragraphs:
        if p != "摘 要":
            add_plain_paragraph(doc, p)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    r1 = p.add_run("关键词：")
    set_run_font(r1, size=10.5, bold=True)
    r2 = p.add_run(keywords)
    set_run_font(r2, size=10.5)
    doc.add_page_break()

    add_centered_text(doc, "目 录", size=16, bold=True, spacing_after=12)
    for chapter, sections in FINAL_TOC:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(chapter)
        set_run_font(r, size=10.5, bold=True)
        for sec in sections:
            sp = doc.add_paragraph()
            sp.paragraph_format.left_indent = Pt(21)
            sp.paragraph_format.space_after = Pt(1)
            sr = sp.add_run(sec)
            set_run_font(sr, size=10)
    doc.add_page_break()


def parse_table(table_path: Path) -> tuple[str, list[list[str]]]:
    text = table_path.read_text(encoding="utf-8")
    caption_args = extract_balanced_args(text, "\\caption")
    caption = clean_latex(caption_args[0]) if caption_args else table_path.stem
    begin = text.find("\\begin{tabular}")
    end = text.find("\\end{tabular}", begin)
    if begin < 0 or end < 0:
        return caption, []
    after = text.find("}", begin + len("\\begin{tabular}"))
    content = text[after + 1 : end]
    content = content.replace("\\toprule", "").replace("\\midrule", "").replace("\\bottomrule", "")
    raw_rows = [r.strip() for r in content.split("\\\\") if r.strip()]
    rows: list[list[str]] = []
    for row in raw_rows:
        if "&" not in row:
            continue
        cells = [clean_latex(c) for c in row.split("&")]
        if any(cells):
            rows.append(cells)
    return caption, rows


def add_docx_table(doc: Document, caption: str, rows: list[list[str]], chapter_no: int, table_counter: dict[int, int]) -> None:
    if not rows:
        return
    table_counter[chapter_no] = table_counter.get(chapter_no, 0) + 1
    caption_text = f"表 {chapter_no}-{table_counter[chapter_no]} {caption}"
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption_text)
    set_run_font(cr, size=10.5, bold=True)
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=1, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except KeyError:
        set_table_borders(table)
    table.autofit = True
    font_size = 7.2 if n_cols >= 8 else 8.5 if n_cols >= 5 else 9.5
    for i in range(n_cols):
        set_cell_text(table.rows[0].cells[i], rows[0][i] if i < len(rows[0]) else "", font_size=font_size, bold=True)
        shade_cell(table.rows[0].cells[i], "EAF2F8")
    for source_row in rows[1:]:
        row = table.add_row()
        for i in range(n_cols):
            value = source_row[i] if i < len(source_row) else ""
            align = "center" if len(value) <= 18 else "left"
            set_cell_text(row.cells[i], value, font_size=font_size, align=align)
    doc.add_paragraph()


def add_figure(doc: Document, figure_path: Path, caption: str, chapter_no: int, fig_counter: dict[int, int]) -> None:
    if not figure_path.exists():
        return
    fig_counter[chapter_no] = fig_counter.get(chapter_no, 0) + 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        p.add_run().add_picture(str(figure_path), width=Inches(5.65))
    except Exception:
        p.add_run().add_picture(str(figure_path), width=Inches(5.2))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(f"图 {chapter_no}-{fig_counter[chapter_no]} {clean_latex(caption)}")
    set_run_font(r, size=10.5, bold=True)


def add_code_block(doc: Document, code: str, caption: str, chapter_no: int, code_counter: dict[int, int]) -> None:
    code_counter[chapter_no] = code_counter.get(chapter_no, 0) + 1
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(f"代码 {chapter_no}-{code_counter[chapter_no]} {clean_latex(caption)}")
    set_run_font(cr, size=10.5, bold=True)
    for line in code.rstrip("\n").splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(14)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        set_run_font(r, size=8.5, font="Consolas")


def add_heading(doc: Document, title: str, level: int) -> None:
    if level == 1:
        doc.add_page_break()
        p = doc.add_paragraph(style="Heading 1")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(title)
    set_run_font(r, size=16 if level == 1 else 13 if level == 2 else 11, bold=True, font="黑体")


def parse_and_add_section(
    doc: Document,
    tex_path: Path,
    figures_dir: Path,
    tables_dir: Path,
    fig_counter: dict[int, int],
    table_counter: dict[int, int],
    code_counter: dict[int, int],
) -> int:
    chapter_no = max(fig_counter.keys() or [0, 0])
    paragraph_lines: list[str] = []
    list_mode: str | None = None
    in_code = False
    code_lines: list[str] = []
    code_caption = ""

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        if paragraph_lines:
            text = clean_latex(" ".join(paragraph_lines))
            if text:
                add_plain_paragraph(doc, text)
            paragraph_lines = []

    for raw in tex_path.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if in_code:
            if line.startswith("\\end{lstlisting}"):
                add_code_block(doc, "\n".join(code_lines), code_caption, chapter_no, code_counter)
                in_code = False
                code_lines = []
                code_caption = ""
            else:
                code_lines.append(raw.rstrip())
            continue
        if not line:
            flush_paragraph()
            continue
        if line.startswith("\\begin{lstlisting}"):
            flush_paragraph()
            cap_match = re.search(r"caption=\{([^}]*)\}", line)
            code_caption = cap_match.group(1) if cap_match else "核心代码片段"
            in_code = True
            continue
        if line.startswith("\\chapter"):
            flush_paragraph()
            args = extract_balanced_args(line, "\\chapter")
            title = clean_latex(args[0]) if args else ""
            if title == "参考文献":
                chapter_no = 8
                add_heading(doc, title, 1)
            elif title.startswith("核心代码"):
                chapter_no = 9
                add_heading(doc, "附录A 核心代码", 1)
            elif title.startswith("完整实验"):
                chapter_no = 10
                add_heading(doc, "附录B 完整实验运行说明", 1)
            elif title.startswith("结果文件"):
                chapter_no = 11
                add_heading(doc, "附录C 结果文件与图表清单", 1)
            elif title.startswith("个人完成"):
                chapter_no = 12
                add_heading(doc, "附录D 个人完成情况与独立性说明", 1)
            else:
                match = re.match(r"^(\d+)_", tex_path.name)
                chapter_no = int(match.group(1)) if match else chapter_no + 1
                add_heading(doc, f"第{chapter_no}章 {title}", 1)
            continue
        if line.startswith("\\chapter*"):
            flush_paragraph()
            args = extract_balanced_args(line, "\\chapter*")
            title = clean_latex(args[0]) if args else ""
            chapter_no = 8
            add_heading(doc, title, 1)
            continue
        if line.startswith("\\section"):
            flush_paragraph()
            args = extract_balanced_args(line, "\\section")
            add_heading(doc, clean_latex(args[0]) if args else "", 2)
            continue
        if line.startswith("\\subsection"):
            flush_paragraph()
            args = extract_balanced_args(line, "\\subsection")
            add_heading(doc, clean_latex(args[0]) if args else "", 3)
            continue
        if line.startswith("\\input"):
            flush_paragraph()
            args = extract_balanced_args(line, "\\input")
            if args:
                table_path = tables_dir / (Path(args[0]).name + ".tex")
                caption, rows = parse_table(table_path)
                add_docx_table(doc, caption, rows, chapter_no, table_counter)
            continue
        if line.startswith("\\maybefigure"):
            flush_paragraph()
            args = extract_balanced_args(line, "\\maybefigure")
            if len(args) >= 2:
                add_figure(doc, figures_dir / args[0], args[1], chapter_no, fig_counter)
            continue
        if line.startswith("\\begin{enumerate}"):
            flush_paragraph()
            list_mode = "number"
            continue
        if line.startswith("\\begin{itemize}"):
            flush_paragraph()
            list_mode = "bullet"
            continue
        if line.startswith("\\end{enumerate}") or line.startswith("\\end{itemize}"):
            flush_paragraph()
            list_mode = None
            continue
        if line.startswith("\\item"):
            flush_paragraph()
            item_text = clean_latex(line.replace("\\item", "", 1))
            style = "List Number" if list_mode == "number" else "List Bullet"
            p = doc.add_paragraph(style=style)
            r = p.add_run(item_text)
            set_run_font(r, size=10.5)
            continue
        if line.startswith("\\addcontentsline") or line.startswith("\\clearpage"):
            continue
        paragraph_lines.append(line)
    flush_paragraph()
    return chapter_no


def collect_doc_text(doc: Document) -> str:
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def try_convert_pdf(docx_path: Path, pdf_path: Path) -> tuple[bool, str]:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice:
        try:
            subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(pdf_path.parent), str(docx_path)],
                check=True,
                capture_output=True,
                text=True,
                timeout=120,
            )
            converted = pdf_path.parent / (docx_path.stem + ".pdf")
            if converted.exists() and converted != pdf_path:
                converted.rename(pdf_path)
            return pdf_path.exists(), "LibreOffice conversion attempted."
        except Exception as exc:  # noqa: BLE001
            return False, f"LibreOffice conversion failed: {exc}"
    return False, "LibreOffice/soffice and Word COM are unavailable in this environment."


def main() -> None:
    repo_root = Path.cwd()
    project = find_project_root()
    overleaf = project / "overleaf_final"
    sections_dir = overleaf / "sections"
    tables_dir = overleaf / "tables"
    figures_dir = overleaf / "figures"
    final_dir = project / "final_submit"
    final_dir.mkdir(exist_ok=True)
    output_docx = final_dir / OUTPUT_DOCX_NAME
    output_pdf = final_dir / OUTPUT_PDF_NAME

    template_path = TEMPLATE_PATH if TEMPLATE_PATH.exists() else repo_root / "期末报告资料" / "课程模板" / "QinTian_experiment.docx"
    doc = Document(str(template_path))
    clear_document_body(doc)
    setup_document(doc)
    add_cover(doc)
    main_tex = (overleaf / "main.tex").read_text(encoding="utf-8")
    add_abstract_and_toc(doc, main_tex)

    fig_counter: dict[int, int] = {}
    table_counter: dict[int, int] = {}
    code_counter: dict[int, int] = {}
    for name in SECTION_ORDER:
        parse_and_add_section(doc, sections_dir / name, figures_dir, tables_dir, fig_counter, table_counter, code_counter)

    doc.save(output_docx)

    generated_doc = Document(str(output_docx))
    text = collect_doc_text(generated_doc)
    banned_hits = [term for term in BANNED_TERMS if term in text]
    latex_hits = [term for term in LATEX_RESIDUES if term in text]
    personal_fields_ok = all(term in text for term in ["学号：", "姓名：", "学院：", "专业年级：", "Institution:", "Grade and major:"])
    pdf_ok, pdf_note = try_convert_pdf(output_docx, output_pdf)

    stats = {
        "template_path": str(template_path),
        "output_docx": str(output_docx),
        "output_pdf": str(output_pdf) if pdf_ok else "",
        "pdf_generated": pdf_ok,
        "pdf_note": pdf_note,
        "stage6_source": str(overleaf),
        "figures_inserted": sum(fig_counter.values()),
        "tables_inserted": sum(table_counter.values()),
        "code_blocks_inserted": sum(code_counter.values()),
        "banned_hits": banned_hits,
        "latex_hits": latex_hits,
        "personal_fields_ok": personal_fields_ok,
        "docx_size": output_docx.stat().st_size,
        "pdf_size": output_pdf.stat().st_size if output_pdf.exists() else 0,
    }
    (project / "results" / "stage7_word_generation_stats.json").write_text(
        json.dumps(stats, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    print(json.dumps(stats, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
